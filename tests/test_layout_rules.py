import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker
from modules.utils import count_columns


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {"font_name": "Palatino Linotype", "font_size": 24, "bold": True, "alignment": "CENTER"},
        "title": {"font_name": "Times New Roman", "font_size": 24, "alignment": "CENTER"},
        "body": {"font_name": "Times New Roman", "font_size": 10},
        "heading": {"font_name": "Times New Roman", "font_size": 10},
        "abstract": {"font_name": "Times New Roman", "font_size": 9},
        "keywords": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {"font_name": "Times New Roman", "font_size": 9},
        "layout": {"page_size": "Letter", "orientation": "PORTRAIT", "columns": 1},
    }


def _save_document(path: Path):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    columns = OxmlElement("w:cols")
    columns.set(qn("w:num"), "2")
    section._sectPr.append(columns)
    for text in [
        "Journal of Informatics and",
        "Web Engineering",
        "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X",
        "A Test Paper Title for Layout Validation",
        "Abstract - This is the abstract.",
        "Keywords - Testing, Layout, Columns, Page, Checker",
        "INTRODUCTION",
        "Body text.",
        "CONCLUSION",
        "Conclusion text.",
        "REFERENCES",
        "[1] Reference text.",
    ]:
        document.add_paragraph(text)
    document.save(path)


class LayoutRulesTest(unittest.TestCase):
    def test_checker_reports_page_size_orientation_and_columns(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "layout_issues.docx"
            _save_document(path)
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("layout", [])
        ]
        self.assertIn("Page size does not match template", descriptions)
        self.assertIn("Page orientation does not match template", descriptions)
        self.assertIn("Column count does not match template", descriptions)

    def test_auto_fix_applies_layout_without_changing_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "layout_fix.docx"
            _save_document(path)
            original_text = [paragraph.text for paragraph in Document(path).paragraphs]
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        section = fixed.sections[0]
        self.assertEqual([paragraph.text for paragraph in fixed.paragraphs], original_text)
        self.assertAlmostEqual(section.page_width.inches, 8.5, places=1)
        self.assertAlmostEqual(section.page_height.inches, 11.0, places=1)
        self.assertEqual(section.orientation, WD_ORIENT.PORTRAIT)
        self.assertEqual(count_columns(fixed), 1)


if __name__ == "__main__":
    unittest.main()
