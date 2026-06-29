import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {"font_name": "Palatino Linotype", "font_size": 24, "bold": True, "alignment": "CENTER"},
        "title": {"font_name": "Times New Roman", "font_size": 24, "alignment": "CENTER"},
        "body": {"font_name": "Times New Roman", "font_size": 10},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "abstract": {"font_name": "Times New Roman", "font_size": 9},
        "keywords": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
            "italic": True,
            "min_count": 5,
            "capitalize_first_letter": True,
        },
        "reference": {"font_name": "Times New Roman", "font_size": 9},
    }


def _save_document(path: Path):
    document = Document()
    for text in [
        "Journal of Informatics and",
        "Web Engineering",
        "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X",
        "A Test Paper Title for Keyword Rules",
        "Abstract - This abstract is long enough for classification.",
    ]:
        document.add_paragraph(text)
    keyword_paragraph = document.add_paragraph()
    run = keyword_paragraph.add_run("Keywords\u2014machine learning, security, detection")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.italic = False
    document.add_paragraph("INTRODUCTION")
    document.add_paragraph("Body text.")
    document.add_paragraph("CONCLUSION")
    document.add_paragraph("Conclusion text.")
    document.add_paragraph("REFERENCES")
    document.add_paragraph("[1] Reference text.")
    document.save(path)


class KeywordRulesTest(unittest.TestCase):
    def test_checker_reports_keyword_count_capitalization_and_italic(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "keyword_issues.docx"
            _save_document(path)
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        self.assertIn("Keyword count is below the template minimum", descriptions)
        self.assertIn("Keyword capitalization does not match template", descriptions)
        self.assertIn("Keywords italic formatting does not match template", descriptions)

    def test_auto_fix_applies_italic_without_rewriting_keywords(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "keyword_fix.docx"
            _save_document(path)
            original = Document(path).paragraphs[5].text
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual(fixed.paragraphs[5].text, original)
        self.assertTrue(all(run.font.italic for run in fixed.paragraphs[5].runs if run.text.strip()))


if __name__ == "__main__":
    unittest.main()
