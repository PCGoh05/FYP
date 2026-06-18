import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from modules.utils import get_paragraph_font_info


class FontExtractionTest(unittest.TestCase):
    def test_english_text_ignores_east_asia_font_when_latin_font_is_in_style(self):
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.style.font.name = "Times New Roman"
        run = paragraph.add_run("The data support the findings of this study.")

        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        r_fonts.set(qn("w:eastAsia"), "Calibri")

        font_info = get_paragraph_font_info(paragraph)

        self.assertEqual(font_info["font_name"], "Times New Roman")


if __name__ == "__main__":
    unittest.main()
