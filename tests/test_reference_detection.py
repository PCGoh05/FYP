import tempfile
import unittest
from pathlib import Path

from docx import Document

from config import DEFAULT_RULES
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {"font_name": "Times New Roman", "font_size": 10, "line_spacing": 1.0},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "references": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {"font_name": "Times New Roman", "font_size": 9, "bold": False},
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


class ReferenceDetectionTest(unittest.TestCase):
    def test_short_ieee_reference_after_references_heading_is_detected(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "short_reference.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()

            reference_issues = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]
            reference_count = result.statistics["reference_count"]

        self.assertEqual(reference_count, 1)
        self.assertNotIn("No references found in document", reference_issues)

    def test_reference_section_stops_before_biography_and_appendix(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "post_reference_sections.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.add_paragraph("BIOGRAPHIES OF AUTHORS")
            document.add_paragraph("Author biography text.")
            document.add_paragraph("APPENDIX A: IMPLEMENTATION CODE")
            document.add_paragraph("import numpy as np")
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()

            reference_texts = [
                classification.text
                for classification in result.classifications
                if classification.paragraph_type.value == "reference"
            ]

        self.assertEqual(reference_texts, ["[1] Reference text."])

    def test_reference_with_publication_source_without_italic_is_reported(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_missing_italic.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph(
                '[1] A. Author, "A useful method," Journal of Informatics and Web Engineering, '
                "vol. 3, no. 1, pp. 1-9, 2026."
            )
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()

            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]

        self.assertIn("Reference publication source may need italic formatting", descriptions)

    def test_reference_with_italic_publication_source_is_not_reported(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_with_italic.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            paragraph = document.add_paragraph()
            paragraph.add_run('[1] A. Author, "A useful method," ')
            italic_run = paragraph.add_run("Journal of Informatics and Web Engineering")
            italic_run.font.italic = True
            paragraph.add_run(", vol. 3, no. 1, pp. 1-9, 2026.")
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()

            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]

        self.assertNotIn("Reference publication source may need italic formatting", descriptions)

    def test_default_rules_enable_jiwe_reference_publication_italic_warning(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "default_rules_reference_missing_italic.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph(
                '[1] A. Author, "A useful method," Journal of Informatics and Web Engineering, '
                "vol. 3, no. 1, pp. 1-9, 2026."
            )
            document.save(path)

            checker = ManuscriptChecker(DEFAULT_RULES).load_manuscript(str(path))
            result = checker.check_all()

            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]

        self.assertIn("Reference publication source may need italic formatting", descriptions)


if __name__ == "__main__":
    unittest.main()
