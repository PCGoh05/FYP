import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {"font_name": "Palatino Linotype", "font_size": 24, "bold": True, "alignment": "CENTER"},
        "title": {"font_name": "Times New Roman", "font_size": 24, "alignment": "CENTER"},
        "author": {"font_name": "Times New Roman", "font_size": 11, "bold": True, "alignment": "CENTER"},
        "affiliation": {"font_name": "Times New Roman", "font_size": 9, "bold": False, "alignment": "CENTER"},
        "corresponding_author": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
            "italic": True,
            "alignment": "CENTER",
            "email_required": True,
            "orcid_required": True,
        },
        "body": {"font_name": "Times New Roman", "font_size": 10},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "abstract": {"font_name": "Times New Roman", "font_size": 9},
        "keywords": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {"font_name": "Times New Roman", "font_size": 9},
    }


def _add(document, text, size=10, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return paragraph


def _save_document(path: Path):
    document = Document()
    _add(document, "Journal of Informatics and", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(document, "Web Engineering", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(document, "A Test Paper Title for Author Validation", 24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add(document, "First Author1, Second Author2", 10, False, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add(document, "1Faculty of Computing, Example University", 10, True, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    _add(document, "*corresponding author: invalid-email; ORCID: 1234", 10, False, False)
    _add(document, "Abstract - This is the abstract.", 9)
    _add(document, "Keywords - Testing, Rules, Author, Format, Checker", 9)
    _add(document, "INTRODUCTION", 10, True)
    _add(document, "Body text.")
    _add(document, "CONCLUSION", 10, True)
    _add(document, "Conclusion text.")
    _add(document, "REFERENCES", 10, True)
    _add(document, "[1] Reference text.", 9)
    document.save(path)


class AuthorInfoRulesTest(unittest.TestCase):
    def test_checker_reports_author_format_and_invalid_contact_details(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "author_issues.docx"
            _save_document(path)
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("author_info", [])
        ]
        self.assertIn("Author name formatting does not match template", descriptions)
        self.assertIn("Affiliation formatting does not match template", descriptions)
        self.assertIn("Corresponding author formatting does not match template", descriptions)
        self.assertIn("Corresponding author email is missing or invalid", descriptions)
        self.assertIn("Corresponding author ORCID is missing or invalid", descriptions)

    def test_auto_fix_formats_author_lines_without_inventing_contact_details(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "author_fix.docx"
            _save_document(path)
            original_text = [paragraph.text for paragraph in Document(path).paragraphs]
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual([paragraph.text for paragraph in fixed.paragraphs], original_text)
        self.assertEqual(fixed.paragraphs[4].alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertEqual(fixed.paragraphs[4].runs[0].font.size.pt, 11)
        self.assertTrue(fixed.paragraphs[4].runs[0].font.bold)
        self.assertEqual(fixed.paragraphs[5].runs[0].font.size.pt, 9)
        self.assertFalse(bool(fixed.paragraphs[5].runs[0].font.bold))
        self.assertTrue(fixed.paragraphs[6].runs[0].font.italic)
        self.assertIn("invalid-email", fixed.paragraphs[6].text)


if __name__ == "__main__":
    unittest.main()
