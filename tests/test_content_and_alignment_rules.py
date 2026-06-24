import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker
from modules.template_extractor import TemplateExtractor


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {
            "font_name": "Times New Roman",
            "font_size": 10,
            "bold": False,
            "line_spacing": 1.0,
            "alignment": "JUSTIFY",
        },
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "subheading": {"font_name": "Times New Roman", "font_size": 10, "bold": False, "italic": True},
        "abstract": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
            "alignment": "JUSTIFY",
            "min_words": 200,
            "max_words": 300,
        },
        "keywords": {"font_name": "Times New Roman", "font_size": 9, "bold": False},
        "reference": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
            "alignment": "JUSTIFY",
            "line_spacing": 1.0,
        },
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


def _add_paragraph(document, text, size, alignment, bold=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    return paragraph


def _save_manuscript(path: Path):
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
    _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
    _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "A Test Paper Title for Rule Validation", 24, WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Abstract - " + "word " * 50, 9, WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(document, "Keywords - checker, template, rules, alignment, abstract", 9, WD_ALIGN_PARAGRAPH.LEFT)
    _add_paragraph(document, "INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    _add_paragraph(
        document,
        "This body paragraph is intentionally left aligned so the checker can report the template mismatch.",
        10,
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    _add_paragraph(document, "CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    _add_paragraph(document, "Conclusion text for the manuscript.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
    _add_paragraph(document, "REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    _add_paragraph(
        document,
        "[1] A. Author, \"Article title,\" Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.",
        9,
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    document.save(path)


class ContentAndAlignmentRulesTest(unittest.TestCase):
    def test_template_extractor_reads_explicit_abstract_word_limit_and_alignment(self):
        document = Document()
        paragraph = document.add_paragraph(
            "Abstract - The abstract must be between 150-250 words and written as one paragraph."
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        rules = TemplateExtractor(document=document).extract_all_rules()

        self.assertEqual(rules["abstract"]["min_words"], 150)
        self.assertEqual(rules["abstract"]["max_words"], 250)
        self.assertEqual(rules["abstract"]["alignment"], "JUSTIFY")

    def test_checker_reports_abstract_word_count_and_alignment_mismatches(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "alignment_issues.docx"
            _save_manuscript(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        body_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        reference_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("references", [])
        ]

        self.assertIn("Abstract word count is outside the template limit", body_descriptions)
        self.assertIn("Abstract alignment does not match template", body_descriptions)
        self.assertIn("Body text alignment does not match template", body_descriptions)
        self.assertIn("Reference alignment does not match template", reference_descriptions)
        self.assertNotIn("Reference line spacing does not match template", reference_descriptions)

    def test_auto_fix_changes_alignment_without_changing_text(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "alignment_fix.docx"
            _save_manuscript(path)
            before = Document(path)
            original_text = [paragraph.text for paragraph in before.paragraphs]

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual([paragraph.text for paragraph in fixed.paragraphs], original_text)
        self.assertEqual(fixed.paragraphs[4].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertEqual(fixed.paragraphs[7].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertEqual(fixed.paragraphs[11].alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)


if __name__ == "__main__":
    unittest.main()
