import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {"_profile": {"name": "JIWE", "required_sections": []}}


def _add_display_equation(document, number=None):
    paragraph = document.add_paragraph()
    math_paragraph = OxmlElement("m:oMathPara")
    math_paragraph.append(OxmlElement("m:oMath"))
    paragraph._p.append(math_paragraph)
    if number is not None:
        paragraph.add_run(f" ({number})")


def _equation_descriptions(document):
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "equations.docx"
        document.save(path)
        result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
    return [
        issue.description
        for issue in result.issues_by_category.get("other", [])
    ]


class EquationNumberingTest(unittest.TestCase):
    def test_checker_accepts_numbered_equation_with_matching_citation(self):
        document = Document()
        _add_display_equation(document, 1)
        document.add_paragraph("Equation (1) defines the objective function.")

        descriptions = _equation_descriptions(document)

        self.assertNotIn("Display equation may be missing a number", descriptions)
        self.assertNotIn("Equation numbering is not continuous", descriptions)
        self.assertNotIn("Equation citation has no matching equation", descriptions)

    def test_checker_reports_missing_numbers_gaps_and_unmatched_citations(self):
        document = Document()
        _add_display_equation(document, 1)
        _add_display_equation(document, 3)
        _add_display_equation(document)
        document.add_paragraph("Equation (2) defines the missing intermediate result.")

        descriptions = _equation_descriptions(document)

        self.assertIn("Display equation may be missing a number", descriptions)
        self.assertIn("Equation numbering is not continuous", descriptions)
        self.assertIn("Equation citation has no matching equation", descriptions)


if __name__ == "__main__":
    unittest.main()
