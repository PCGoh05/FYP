import tempfile
import unittest
from pathlib import Path

from docx import Document

from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE", "required_sections": []},
        "reference": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
        },
    }


def _reference_descriptions(document):
    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "references.docx"
        document.save(path)
        result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
    return [
        issue.description
        for issue in result.issues_by_category.get("references", [])
    ]


class CitationReferenceConsistencyTest(unittest.TestCase):
    def test_checker_accepts_continuous_references_with_matching_citations(self):
        document = Document()
        document.add_paragraph("Prior studies support the method [1], [2].")
        document.add_paragraph("REFERENCES")
        document.add_paragraph("[1] A. Author, First article, Journal, 2025.")
        document.add_paragraph("[2] B. Author, Second article, Journal, 2026.")

        descriptions = _reference_descriptions(document)

        self.assertNotIn("Reference numbering is not continuous", descriptions)
        self.assertNotIn("In-text citation has no matching reference", descriptions)
        self.assertNotIn("Reference is not cited in the manuscript", descriptions)

    def test_checker_reconstructs_split_run_citations_and_ignores_zero_based_ranges(self):
        document = Document()
        paragraph = document.add_paragraph("Prior studies support the method ")
        paragraph.add_run("[")
        paragraph.add_run("1")
        paragraph.add_run("]")
        paragraph.add_run(", while values are clipped to [0,1].")
        document.add_paragraph("REFERENCES")
        document.add_paragraph("[1] A. Author, First article, Journal, 2025.")

        descriptions = _reference_descriptions(document)

        self.assertNotIn("In-text citation has no matching reference", descriptions)
        self.assertNotIn("Reference is not cited in the manuscript", descriptions)

    def test_checker_reports_numbering_gaps_missing_references_and_uncited_entries(self):
        document = Document()
        document.add_paragraph("Prior studies support the method [1] and [3].")
        document.add_paragraph("REFERENCES")
        document.add_paragraph("[1] A. Author, First article, Journal, 2025.")
        document.add_paragraph("[2] B. Author, Second article, Journal, 2025.")
        document.add_paragraph("[4] C. Author, Fourth article, Journal, 2026.")

        descriptions = _reference_descriptions(document)

        self.assertIn("Reference numbering is not continuous", descriptions)
        self.assertIn("In-text citation has no matching reference", descriptions)
        self.assertIn("Reference is not cited in the manuscript", descriptions)


if __name__ == "__main__":
    unittest.main()
