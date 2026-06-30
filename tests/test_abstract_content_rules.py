import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt

from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {"font_name": "Palatino Linotype", "font_size": 24, "bold": True, "alignment": "CENTER"},
        "title": {"font_name": "Times New Roman", "font_size": 24, "alignment": "CENTER"},
        "body": {"font_name": "Times New Roman", "font_size": 10},
        "heading": {"font_name": "Times New Roman", "font_size": 10},
        "abstract": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "min_words": 200,
            "max_words": 300,
            "one_paragraph": True,
            "prohibit_equations": True,
            "prohibit_tables": True,
            "prohibit_citations": True,
        },
        "keywords": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {"font_name": "Times New Roman", "font_size": 9},
    }


class AbstractContentRulesTest(unittest.TestCase):
    def test_checker_reports_multiple_paragraphs_citation_equation_and_table_material(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "abstract_content_issues.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X",
                "A Test Paper Title for Abstract Validation",
            ]:
                document.add_paragraph(text)
            first = document.add_paragraph("Abstract - This study follows earlier work [1].")
            first.runs[0].font.size = Pt(9)
            second = document.add_paragraph("The result is x = y and Table 1 summarizes the findings.")
            second.runs[0].font.size = Pt(9)
            math = OxmlElement("m:oMath")
            second._p.append(math)
            document.add_paragraph("Keywords - Testing, Abstract, Content, Rules, Checker")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        self.assertIn("Abstract must be written as one paragraph", descriptions)
        self.assertIn("Abstract contains citation or reference markers", descriptions)
        self.assertIn("Abstract contains equation-like content", descriptions)
        self.assertIn("Abstract contains table or figure material", descriptions)

    def test_checker_reports_abstract_over_word_limit(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "long_abstract.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X",
                "A Test Paper Title for Abstract Validation",
            ]:
                document.add_paragraph(text)
            document.add_paragraph("Abstract - " + "word " * 301)
            document.add_paragraph("Keywords - Testing, Abstract, Content, Rules, Checker")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        issues = result.issues_by_category.get("body_text", [])
        matching = [
            issue for issue in issues
            if issue.description == "Abstract word count is outside the template limit"
        ]
        self.assertEqual(len(matching), 1)
        self.assertEqual(matching[0].current_value, "301 words")
        self.assertEqual(matching[0].expected_value, "200-300 words")

    def test_checker_reports_statistical_formula_like_abstract_content(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "abstract_formula.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X",
                "A Test Paper Title for Abstract Validation",
            ]:
                document.add_paragraph(text)
            document.add_paragraph(
                "Abstract - The model showed significance with p < .05 and F(1, 99) = 4.50."
            )
            document.add_paragraph("Keywords - Testing, Abstract, Content, Rules, Checker")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        self.assertIn("Abstract contains equation-like content", descriptions)


if __name__ == "__main__":
    unittest.main()
