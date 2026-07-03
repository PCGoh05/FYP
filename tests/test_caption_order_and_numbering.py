import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE", "required_sections": []},
        "caption": {"font_name": "Times New Roman", "font_size": 10, "italic": False},
    }


def _add_drawing_paragraph(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:drawing"))


def _add_multi_drawing_paragraph(document, count=2, text=""):
    paragraph = document.add_paragraph(text)
    run = paragraph.add_run()
    for _ in range(count):
        run._r.append(OxmlElement("w:drawing"))


def _add_pict_paragraph(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:pict"))


class CaptionOrderAndNumberingTest(unittest.TestCase):
    def test_checker_accepts_correct_caption_positions_and_numbering(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "valid_captions.docx"
            document = Document()

            _add_drawing_paragraph(document)
            document.add_paragraph("Figure 1: Caption below the image")

            document.add_paragraph("Table 1: Caption above the table")
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table data"
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for category in ("tables", "figures")
            for issue in result.issues_by_category.get(category, [])
        ]

        self.assertNotIn("Table caption should appear above the table", descriptions)
        self.assertNotIn("Table numbering is not continuous", descriptions)
        self.assertNotIn("Figure caption should appear below the figure", descriptions)
        self.assertNotIn("Figure numbering is not continuous", descriptions)

    def test_checker_reports_reversed_caption_positions_and_numbering_gaps(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_issues.docx"
            document = Document()

            document.add_paragraph("Figure 1: Caption placed before the image")
            _add_drawing_paragraph(document)

            _add_drawing_paragraph(document)
            document.add_paragraph("Figure 3: Caption with a numbering gap")

            first_table = document.add_table(rows=1, cols=1)
            first_table.cell(0, 0).text = "First table data"
            document.add_paragraph("Table 1: Caption placed after the table")

            document.add_paragraph("Table 3: Caption with a numbering gap")
            second_table = document.add_table(rows=1, cols=1)
            second_table.cell(0, 0).text = "Second table data"
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        table_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("tables", [])
        ]
        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]

        self.assertIn("Table caption should appear above the table", table_descriptions)
        self.assertIn("Table numbering is not continuous", table_descriptions)
        self.assertIn("Figure caption should appear below the figure", figure_descriptions)
        self.assertIn("Figure numbering is not continuous", figure_descriptions)

    def test_checker_reports_table_caption_font_and_size_mismatches(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "table_caption_format.docx"
            document = Document()
            caption = document.add_paragraph()
            run = caption.add_run("Table 1: Caption above the table")
            run.font.name = "Arial"
            run.font.size = Pt(20)
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table data"
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        table_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("tables", [])
        ]
        self.assertIn("Table caption font does not match template", table_descriptions)
        self.assertIn("Table caption size does not match template", table_descriptions)

    def test_checker_reports_caption_italic_mismatches(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_italic_format.docx"
            document = Document()

            _add_drawing_paragraph(document)
            figure_caption = document.add_paragraph()
            figure_run = figure_caption.add_run("Figure 1: Caption below the image")
            figure_run.font.name = "Times New Roman"
            figure_run.font.size = Pt(10)
            figure_run.font.italic = True

            table_caption = document.add_paragraph()
            table_run = table_caption.add_run("Table 1: Caption above the table")
            table_run.font.name = "Times New Roman"
            table_run.font.size = Pt(10)
            table_run.font.italic = True
            table = document.add_table(rows=1, cols=1)
            table.cell(0, 0).text = "Table data"
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        table_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("tables", [])
        ]
        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertIn("Table caption italic formatting does not match template", table_descriptions)
        self.assertIn("Figure caption italic formatting does not match template", figure_descriptions)

    def test_auto_fix_does_not_reformat_caption_for_position_issue_only(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_position_only.docx"
            document = Document()
            section = document.sections[0]
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

            document.add_paragraph("Figure 1: Caption placed before the image")
            _add_drawing_paragraph(document)
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertIn("Figure caption should appear below the figure", descriptions)
        self.assertEqual([], fixer.get_change_records())

    def test_checker_reports_missing_caption_for_body_image_before_any_caption(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "body_image_without_caption.docx"
            document = Document()
            document.add_paragraph("1. INTRODUCTION")
            _add_drawing_paragraph(document)
            document.add_paragraph("Body text after the image.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertIn("Some figures may be missing captions", descriptions)

    def test_checker_ignores_front_matter_images_before_introduction(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "front_matter_logo.docx"
            document = Document()
            _add_drawing_paragraph(document)
            document.add_paragraph("1. INTRODUCTION")
            document.add_paragraph("Body text with no figures.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertNotIn("Some figures may be missing captions", descriptions)

    def test_checker_detects_legacy_pict_for_caption_order(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "legacy_pict_caption_order.docx"
            document = Document()
            document.add_paragraph("Figure 1: Caption placed before the image")
            _add_pict_paragraph(document)
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertIn("Figure caption should appear below the figure", descriptions)

    def test_checker_does_not_treat_body_figure_mentions_as_captions(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "body_figure_mentions.docx"
            document = Document()
            document.add_paragraph("1. INTRODUCTION")
            document.add_paragraph("Figure 1 and Figure 2 show the proposed workflow in detail.")
            _add_drawing_paragraph(document)
            document.add_paragraph("Figure 1. Proposed Workflow")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertNotIn("Figure numbering is not continuous", figure_descriptions)
        self.assertNotIn("Figure caption should appear below the figure", figure_descriptions)

    def test_checker_accepts_caption_without_period_when_title_starts_after_number(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_without_period.docx"
            document = Document()
            document.add_paragraph("1. INTRODUCTION")
            _add_drawing_paragraph(document)
            document.add_paragraph("Figure 1 Proposed Workflow")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertNotIn("Some figures may be missing captions", figure_descriptions)

    def test_checker_counts_multi_drawing_block_as_one_figure(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "multi_drawing_one_figure.docx"
            document = Document()
            document.add_paragraph("1. INTRODUCTION")
            _add_multi_drawing_paragraph(document, count=3)
            document.add_paragraph("Figure 1. Combined Result Panels")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertNotIn("Some figures may be missing captions", figure_descriptions)

    def test_checker_does_not_count_equation_drawing_as_missing_figure(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "equation_drawing.docx"
            document = Document()
            document.add_paragraph("1. INTRODUCTION")
            _add_multi_drawing_paragraph(document, count=2, text="Equation (7) defines the model.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        figure_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertNotIn("Some figures may be missing captions", figure_descriptions)


if __name__ == "__main__":
    unittest.main()
