import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from modules.template_extractor import TemplateExtractor
from modules.profile_loader import ProfileLoader
from config import DEFAULT_RULES


def _set_document_default_line_spacing(document, line_value):
    styles = document.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    p_pr_default = doc_defaults.find(qn("w:pPrDefault"))
    p_pr = p_pr_default.find(qn("w:pPr"))
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        p_pr.append(spacing)
    spacing.set(qn("w:line"), str(line_value))
    spacing.set(qn("w:lineRule"), "auto")


def _append_reference_content_control(document, text):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    justification = OxmlElement("w:jc")
    justification.set(qn("w:val"), "both")
    paragraph_properties.append(justification)
    paragraph.append(paragraph_properties)

    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Times New Roman")
    fonts.set(qn("w:hAnsi"), "Times New Roman")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    run_properties.append(fonts)
    run_properties.append(size)
    run.append(run_properties)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    content.append(paragraph)
    sdt.append(content)
    document.element.body.insert(len(document.element.body) - 1, sdt)


class ReferenceTemplateExtractionTest(unittest.TestCase):
    def test_jiwe_profile_reference_spacing_matches_template_default(self):
        rules = ProfileLoader().default_rules(ProfileLoader().load("jiwe"))

        self.assertEqual(rules["reference"]["line_spacing"], 1.15)
        self.assertEqual(DEFAULT_RULES["reference"]["line_spacing"], 1.15)

    def test_reference_examples_in_content_controls_inherit_document_spacing(self):
        document = Document()
        _set_document_default_line_spacing(document, 276)
        document.add_paragraph("Journal of Informatics and Web Engineering")
        document.add_paragraph("REFERENCES (10-Font size, Times New Roman)")
        _append_reference_content_control(
            document,
            '[1] A. Author, "Article title," Journal of Testing, 2026.',
        )

        rules = TemplateExtractor(document=document).extract_all_rules()

        self.assertEqual(rules["reference"]["font_name"], "Times New Roman")
        self.assertEqual(rules["reference"]["font_size"], 9)
        self.assertEqual(rules["reference"]["alignment"], "JUSTIFY")
        self.assertEqual(rules["reference"]["line_spacing"], 1.15)
        self.assertIsNone(rules["reference"]["left_indent"])
        self.assertAlmostEqual(rules["reference"]["hanging_indent"], 0.4444444444444444)
        self.assertEqual(
            rules["_provenance"]["reference.line_spacing"]["source"],
            "extracted",
        )


if __name__ == "__main__":
    unittest.main()
