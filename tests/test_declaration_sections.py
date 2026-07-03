import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker
from modules.profile_loader import ProfileLoader


def _rules():
    return {
        "_profile": {
            "name": "JIWE",
            "required_sections": ["abstract", "keywords", "introduction", "conclusion", "references"],
            "required_declarations": [
                "acknowledgement",
                "funding_statement",
                "author_contributions",
                "conflict_of_interests",
                "ethics_statements",
                "data_availability",
            ],
        },
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {"font_name": "Palatino Linotype", "font_size": 24, "bold": True, "alignment": "CENTER"},
        "title": {"font_name": "Times New Roman", "font_size": 24, "alignment": "CENTER"},
        "body": {"font_name": "Times New Roman", "font_size": 10},
        "heading": {"font_name": "Times New Roman", "font_size": 10},
        "abstract": {"font_name": "Times New Roman", "font_size": 9},
        "keywords": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {"font_name": "Times New Roman", "font_size": 9},
    }


def _profile_rules():
    loader = ProfileLoader()
    profile = loader.load("jiwe")
    rules = loader.default_rules(profile)
    rules["_profile"] = {
        "name": profile.get("name", "JIWE"),
        "required_sections": profile.get("required_sections", []),
        "required_declarations": profile.get("required_declarations", []),
        "declaration_templates": profile.get("declaration_templates", {}),
        "classification_patterns": profile.get("classification_patterns", {}),
    }
    return rules


class DeclarationSectionsTest(unittest.TestCase):
    def test_body_sentence_does_not_override_real_conclusion_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "section_evidence.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "A Test Paper Title",
                "Abstract - This is the abstract.",
                "Keywords - Testing, Rules, Sections, Format, Checker",
                "INTRODUCTION",
                (
                    "The rest of this paper is organized as follows: Section 2 describes the method, "
                    "and Section 6 presents the conclusion."
                ),
                "CONCLUSION",
                "Conclusion text.",
                "REFERENCES",
                "[1] Reference text.",
            ]:
                document.add_paragraph(text)
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        self.assertEqual(result.document_structure["sections"]["conclusion"]["index"], 7)
        self.assertEqual(result.document_structure["sections"]["conclusion"]["format_status"], "valid")
        weak_section_issues = [
            issue for issue in result.issues_by_category.get("structure", [])
            if issue.description == "Section was found but its heading role is not confidently detected"
            and issue.location == "Conclusion Section"
        ]
        self.assertEqual(weak_section_issues, [])

    def test_checker_reports_missing_required_jiwe_declarations(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "missing_declarations.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "A Test Paper Title",
                "Abstract - This is the abstract.",
                "Keywords - Testing, Rules, Sections, Format, Checker",
                "INTRODUCTION",
                "Body text.",
                "CONCLUSION",
                "Conclusion text.",
                "REFERENCES",
                "[1] Reference text.",
            ]:
                document.add_paragraph(text)
            document.save(path)
            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("structure", [])
        ]
        self.assertIn("Missing required declaration section: Funding Statement", descriptions)
        self.assertIn("Missing required declaration section: Author Contributions", descriptions)
        self.assertIn("Missing required declaration section: Conflict Of Interests", descriptions)
        self.assertIn("Missing required declaration section: Data Availability", descriptions)
        self.assertIn("Missing required declaration section: Acknowledgement", descriptions)
        self.assertIn("Missing required declaration section: Ethics Statements", descriptions)

    def test_auto_fix_inserts_jiwe_declaration_template_text_before_references(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "missing_declarations_fix.docx"
            document = Document()
            for text in [
                "Journal of Informatics and",
                "Web Engineering",
                "A Test Paper Title",
                "Abstract - This is the abstract.",
                "Keywords - Testing, Rules, Sections, Format, Checker",
                "INTRODUCTION",
                "Body text.",
                "CONCLUSION",
                "Conclusion text.",
                "REFERENCES",
                "[1] Reference text.",
            ]:
                document.add_paragraph(text)
            document.save(path)

            rules = _profile_rules()
            result = ManuscriptChecker(rules).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(rules, result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        fixed_text = [paragraph.text for paragraph in fixed.paragraphs]
        references_index = fixed_text.index("REFERENCES")
        acknowledgement_index = fixed_text.index("ACKNOWLEDGEMENT")
        funding_index = fixed_text.index("FUNDING STATEMENT")
        data_index = fixed_text.index("DATA AVAILABILITY")

        self.assertLess(acknowledgement_index, references_index)
        self.assertLess(funding_index, references_index)
        self.assertLess(data_index, references_index)
        self.assertIn(
            "The authors received no funding from any party for the research and publication of this article.",
            fixed_text,
        )
        self.assertIn(
            "Data availability is not applicable to this paper as no new data were created or analyzed in this study.",
            fixed_text,
        )
        change_properties = [change.property_name for change in fixer.changes]
        self.assertIn("declaration_template", change_properties)


if __name__ == "__main__":
    unittest.main()
