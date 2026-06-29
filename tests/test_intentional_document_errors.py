import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {"font_name": "Times New Roman", "font_size": 10, "bold": False, "line_spacing": 1.0},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "subheading": {"font_name": "Times New Roman", "font_size": 10, "bold": False, "italic": True},
        "abstract": {"font_name": "Times New Roman", "font_size": 9, "bold": False},
        "keywords": {"font_name": "Times New Roman", "font_size": 9, "bold": False},
        "reference": {"font_name": "Times New Roman", "font_size": 9, "bold": False},
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


def _add_paragraph(document, text, font_name="Times New Roman", size=10, bold=None, italic=None, alignment=None):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return paragraph


def _save_document(path: Path, intro_text="INTRODUCTION", body_size=10, all_bold=False):
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    _add_paragraph(document, "Journal of Informatics and", "Palatino Linotype", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Web Engineering", "Palatino Linotype", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Vol. 5 No. 2 (June 2026)\teISSN: 2821-370X", "Palatino Linotype", 10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "A Test Paper Title for Format Validation", size=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Abstract - This is the abstract.", size=9, bold=all_bold)
    _add_paragraph(document, "Keywords - checking, template", size=9, bold=all_bold)
    if intro_text:
        _add_paragraph(document, intro_text, size=10, bold=True)
    _add_paragraph(document, "This body paragraph should be normal text.", size=body_size, bold=all_bold)
    _add_paragraph(document, "Another normal body paragraph for testing.", size=body_size, bold=all_bold)
    _add_paragraph(document, "CONCLUSION", size=10, bold=True)
    _add_paragraph(document, "Conclusion text.", size=body_size, bold=all_bold)
    _add_paragraph(document, "REFERENCES", size=10, bold=True)
    _add_paragraph(document, "[1] Reference text.", size=9, bold=all_bold)
    document.save(path)


def _save_bulk_bad_body_document(path: Path):
    document = Document()
    section = document.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    _add_paragraph(document, "Journal of Informatics and", "Palatino Linotype", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Web Engineering", "Palatino Linotype", 24, True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "A Bulk Formatting Regression Paper", size=24, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Abstract - This is the abstract.", font_name="Arial", size=20, bold=True)
    _add_paragraph(document, "Keywords - checking, template, rules, formatting, validation", font_name="Arial", size=20, bold=True)
    _add_paragraph(document, "INTRODUCTION", font_name="Arial", size=20, bold=True)
    for index in range(18):
        _add_paragraph(
            document,
            f"This body paragraph {index} uses deliberately wrong bulk formatting.",
            font_name="Arial",
            size=20,
            bold=True,
        )
    _add_paragraph(document, "CONCLUSION", font_name="Arial", size=20, bold=True)
    _add_paragraph(document, "Conclusion text.", font_name="Arial", size=20, bold=True)
    _add_paragraph(document, "REFERENCES", font_name="Arial", size=20, bold=True)
    _add_paragraph(document, "[1] Reference text.", font_name="Arial", size=20, bold=True)
    document.save(path)


class IntentionalDocumentErrorsTest(unittest.TestCase):
    def _check(self, **kwargs):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "intentional_error.docx"
            _save_document(path, **kwargs)
            return ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

    def test_missing_required_heading_is_reported(self):
        result = self._check(intro_text=None)
        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("structure", [])
        ]
        self.assertIn("Missing required section: Introduction", descriptions)

    def test_misspelled_required_heading_is_reported(self):
        result = self._check(intro_text="INTRODUTION")
        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("structure", [])
        ]
        self.assertIn("Possible misspelled required section heading", descriptions)

    def test_large_body_font_size_is_reported(self):
        result = self._check(body_size=20)
        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        self.assertIn("Body text font size does not match template", descriptions)

    def test_all_bold_body_and_reference_text_is_reported(self):
        result = self._check(all_bold=True)
        body_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        reference_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("references", [])
        ]
        self.assertIn("Body text bold formatting does not match template", body_descriptions)
        self.assertIn("Reference bold formatting does not match template", reference_descriptions)

    def test_auto_fix_corrects_bulk_body_font_size_and_bold_after_issue_cap(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "bulk_bad_body.docx"
            _save_bulk_bad_body_document(path)
            before = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()

            fixer = AutoFixer(_rules(), before.classifications, before.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()

            after = ManuscriptChecker(_rules()).load_manuscript(
                BytesIO(fixer.get_fixed_document_bytes())
            ).check_all()

        body_descriptions = [
            issue.description
            for issue in after.issues_by_category.get("body_text", [])
        ]
        self.assertNotIn("Body text font does not match template", body_descriptions)
        self.assertNotIn("Body text font size does not match template", body_descriptions)
        self.assertNotIn("Body text bold formatting does not match template", body_descriptions)


if __name__ == "__main__":
    unittest.main()
