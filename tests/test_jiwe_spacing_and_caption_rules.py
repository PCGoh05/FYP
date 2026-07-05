import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker
from modules.profile_loader import ProfileLoader
from modules.template_extractor import TemplateExtractor
from modules.utils import paragraph_has_manual_line_breaks, to_journal_caption_title_case


def _jiwe_rules():
    return ProfileLoader().default_rules(ProfileLoader().load("jiwe"))


def _add_paragraph(document, text, size=10, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False):
    paragraph = document.add_paragraph()
    paragraph.alignment = alignment
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.paragraph_format.space_after = Pt(7.5)
    return paragraph


def _docx_bytes(document):
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


def _blank_count_before(document, text):
    target_index = next(
        index for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == text
    )
    count = 0
    index = target_index - 1
    while index >= 0 and not document.paragraphs[index].text.strip():
        count += 1
        index -= 1
    return count


def _add_drawing_paragraph(document):
    from docx.oxml import OxmlElement

    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    run._r.append(OxmlElement("w:drawing"))
    return paragraph


def _save_spacing_issue_document(path: Path):
    document = Document()
    _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
    _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
    _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "A Test Paper Title for Spacing Validation", 24, WD_ALIGN_PARAGRAPH.CENTER)
    _add_paragraph(document, "Abstract - " + "word " * 210, 9)
    _add_paragraph(document, "Keywords - Template, Checking, Rules, References, Formatting", 9)
    _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    body = _add_paragraph(
        document,
        "This body paragraph cites reference [1] and intentionally has wrong paragraph spacing after.",
    )
    body.paragraph_format.space_after = Pt(0)
    _add_paragraph(document, "4. CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    _add_paragraph(document, "Conclusion text for the manuscript.")
    _add_paragraph(document, "5. REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
    reference = _add_paragraph(
        document,
        '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
        9,
    )
    reference.paragraph_format.line_spacing = 1.0
    reference.paragraph_format.space_after = Pt(0)
    reference.paragraph_format.first_line_indent = Inches(-0.5)
    document.save(path)


class JIWESpacingAndCaptionRulesTest(unittest.TestCase):
    def test_jiwe_profile_contains_template_spacing_and_caption_case_rules(self):
        rules = _jiwe_rules()

        self.assertEqual(rules["body"]["space_after"], 7.5)
        self.assertEqual(rules["caption"]["space_after"], 7.5)
        self.assertFalse(rules["caption"]["bold"])
        self.assertTrue(rules["caption"]["title_case"])
        self.assertEqual(rules["reference"]["line_spacing"], 1.15)
        self.assertEqual(rules["reference"]["space_after"], 10.0)
        self.assertIsNone(rules["reference"]["left_indent"])
        self.assertAlmostEqual(rules["reference"]["hanging_indent"], 0.4444444444444444)
        self.assertEqual(rules["heading"]["alignment"], "LEFT")
        self.assertEqual(rules["heading"]["line_spacing"], 1.0)
        self.assertIsNone(rules["heading"]["space_before"])
        self.assertEqual(rules["heading"]["space_after"], 7.5)
        self.assertEqual(rules["heading"]["blank_before_max"], 1)
        self.assertEqual(rules["introduction_heading"]["space_before"], 15.0)
        self.assertEqual(rules["subheading"]["alignment"], "LEFT")
        self.assertEqual(rules["subheading"]["line_spacing"], 1.0)
        self.assertEqual(rules["subheading"]["space_before"], 0.0)
        self.assertEqual(rules["subheading"]["space_after"], 7.5)
        self.assertEqual(rules["subheading"]["blank_before"], 1)
        self.assertEqual(rules["biography_heading"]["font_size"], 10.5)
        self.assertEqual(rules["biography_heading"]["line_spacing"], 1.15)
        self.assertEqual(rules["biography_heading"]["space_after"], 10.0)
        self.assertTrue(rules["reference"]["number_tab_required"])

    def test_jiwe_biography_heading_uses_template_style(self):
        document = Document()
        _add_paragraph(document, "5. REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
        _add_paragraph(
            document,
            '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
            9,
        )
        document.add_paragraph("")
        biography = document.add_paragraph()
        biography.alignment = WD_ALIGN_PARAGRAPH.LEFT
        biography.paragraph_format.line_spacing = 1.15
        biography.paragraph_format.space_after = Pt(10)
        run = biography.add_run("BIOGRAPHIES OF AUTHORS")
        run.font.name = "Times New Roman"
        run.font.size = Pt(10.5)
        run.font.bold = True

        result = ManuscriptChecker(_jiwe_rules()).load_manuscript(_docx_bytes(document)).check_all()

        biography_issues = [
            issue.description
            for issue in result.issues_by_category.get("headings", [])
            if "BIOGRAPHIES" in issue.text_preview
        ]
        self.assertEqual(biography_issues, [])

    def test_jiwe_subheading_blank_paragraphs_are_detected_and_fixed(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "subheading_blanks.docx"
            document = Document()
            _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "This body paragraph should be separated from the next numbered subheading.")
            subheading = _add_paragraph(document, "1.1 Missing Blank Before", 10, WD_ALIGN_PARAGRAPH.LEFT)
            for run in subheading.runs:
                run.font.italic = True
            _add_paragraph(document, "This body paragraph is followed by too many blank paragraphs.")
            for _ in range(4):
                blank = document.add_paragraph("")
                blank.paragraph_format.line_spacing = 1.15
                blank.paragraph_format.space_before = Pt(4)
                blank.paragraph_format.space_after = Pt(4)
            subheading = _add_paragraph(document, "1.2 Excess Blank Before", 10, WD_ALIGN_PARAGRAPH.LEFT)
            for run in subheading.runs:
                run.font.italic = True
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            heading_descriptions = [
                issue.description
                for issue in result.issues_by_category.get("headings", [])
            ]
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertIn("Heading is missing required blank paragraph before it", heading_descriptions)
        self.assertIn("Heading has too many blank paragraphs before it", heading_descriptions)
        self.assertEqual(_blank_count_before(fixed, "1.1 Missing Blank Before"), 1)
        self.assertEqual(_blank_count_before(fixed, "1.2 Excess Blank Before"), 1)

    def test_jiwe_main_heading_excess_blank_paragraphs_are_collapsed(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "main_heading_blanks.docx"
            document = Document()
            _add_paragraph(document, "Figure 1. System Architecture Diagram", 10, WD_ALIGN_PARAGRAPH.CENTER)
            for _ in range(4):
                blank = document.add_paragraph("")
                blank.paragraph_format.line_spacing = 1.0
                blank.paragraph_format.space_after = Pt(7.5)
            _add_paragraph(document, "CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            heading_descriptions = [
                issue.description
                for issue in result.issues_by_category.get("headings", [])
            ]
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertIn("Heading has too many blank paragraphs before it", heading_descriptions)
        self.assertEqual(_blank_count_before(fixed, "CONCLUSION"), 1)

    def test_jiwe_introduction_heading_space_before_matches_template(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "introduction_heading_space_before.docx"
            document = Document()
            introduction = _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            introduction.paragraph_format.space_before = Pt(0)
            conclusion = _add_paragraph(document, "CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            conclusion.paragraph_format.space_before = Pt(0)
            document.save(path)

            rules = _jiwe_rules()
            before = ManuscriptChecker(rules).load_manuscript(str(path)).check_all()
            spacing_before_previews = [
                issue.text_preview
                for issue in before.issues_by_category.get("headings", [])
                if issue.description == "Heading spacing before does not match template"
            ]
            fixer = AutoFixer(rules, before.classifications, before.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        fixed_heading = next(paragraph for paragraph in fixed.paragraphs if paragraph.text == "1. INTRODUCTION")
        fixed_conclusion = next(paragraph for paragraph in fixed.paragraphs if paragraph.text == "CONCLUSION")
        self.assertIn("1. INTRODUCTION", spacing_before_previews)
        self.assertNotIn("CONCLUSION", spacing_before_previews)
        self.assertAlmostEqual(fixed_heading.paragraph_format.space_before.pt, 15.0)
        self.assertAlmostEqual(fixed_conclusion.paragraph_format.space_before.pt, 0.0)

    def test_jiwe_introduction_front_matter_blank_paragraphs_are_preserved(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "introduction_front_matter_blanks.docx"
            document = Document()
            _add_paragraph(document, "Published: 16 June 2026", 10, WD_ALIGN_PARAGRAPH.RIGHT)
            for _ in range(2):
                blank = document.add_paragraph("")
                blank.paragraph_format.line_spacing = 1.0
                blank.paragraph_format.space_after = Pt(7.5)
            _add_paragraph(document, "INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            document.save(path)

            rules = _jiwe_rules()
            before = ManuscriptChecker(rules).load_manuscript(str(path)).check_all()
            intro_blank_issues = [
                issue.description
                for issue in before.issues_by_category.get("headings", [])
                if issue.text_preview == "INTRODUCTION"
                and issue.description == "Heading has too many blank paragraphs before it"
            ]
            fixer = AutoFixer(rules, before.classifications, before.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual(intro_blank_issues, [])
        self.assertEqual(_blank_count_before(fixed, "INTRODUCTION"), 2)

    def test_jiwe_heading_spacing_is_detected_and_fixed(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "heading_spacing.docx"
            document = Document()
            _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "A Test Paper Title for Heading Spacing", 24, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "Abstract - " + "word " * 210, 9)
            _add_paragraph(document, "Keywords - Template, Checking, Rules, References, Formatting", 9)
            _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            subheading = _add_paragraph(document, "1.1 Experimental Setup", 10, WD_ALIGN_PARAGRAPH.LEFT)
            subheading.paragraph_format.line_spacing = 1.15
            subheading.paragraph_format.space_before = Pt(4)
            subheading.paragraph_format.space_after = Pt(4)
            for run in subheading.runs:
                run.font.italic = True
            _add_paragraph(document, "This body paragraph is correctly justified.")
            _add_paragraph(document, "4. CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "Conclusion text for the manuscript.")
            _add_paragraph(document, "5. REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(
                document,
                '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
                9,
            )
            document.save(path)

            rules = _jiwe_rules()
            before = ManuscriptChecker(rules).load_manuscript(str(path)).check_all()
            heading_descriptions = [
                issue.description
                for issue in before.issues_by_category.get("headings", [])
            ]
            fixer = AutoFixer(rules, before.classifications, before.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        fixed_subheading = next(paragraph for paragraph in fixed.paragraphs if paragraph.text == "1.1 Experimental Setup")
        self.assertIn("Heading line spacing does not match template", heading_descriptions)
        self.assertIn("Heading spacing before does not match template", heading_descriptions)
        self.assertIn("Heading spacing after does not match template", heading_descriptions)
        self.assertAlmostEqual(fixed_subheading.paragraph_format.line_spacing, 1.0)
        self.assertAlmostEqual(fixed_subheading.paragraph_format.space_before.pt, 0.0)
        self.assertAlmostEqual(fixed_subheading.paragraph_format.space_after.pt, 7.5)

    def test_template_extraction_includes_subheading_spacing(self):
        document = Document()
        _add_paragraph(document, "3. RESEARCH METHODOLOGY", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
        subheading = _add_paragraph(document, "3.1 Research approach", 10, WD_ALIGN_PARAGRAPH.LEFT)
        for run in subheading.runs:
            run.font.italic = True

        rules = TemplateExtractor(document=document).extract_all_rules()

        self.assertEqual(rules["subheading"]["line_spacing"], 1.0)
        self.assertEqual(rules["subheading"]["space_before"], 0.0)
        self.assertEqual(rules["subheading"]["space_after"], 7.5)

    def test_jiwe_heading_alignment_is_detected_and_fixed(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "heading_alignment.docx"
            document = Document()
            _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "A Test Paper Title for Heading Alignment", 24, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "Abstract - " + "word " * 210, 9)
            _add_paragraph(document, "Keywords - Template, Checking, Rules, References, Formatting", 9)
            _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.CENTER, True)
            subheading = _add_paragraph(document, "1.1 Experimental Setup", 10, WD_ALIGN_PARAGRAPH.CENTER)
            for run in subheading.runs:
                run.font.italic = True
            _add_paragraph(document, "This body paragraph is correctly justified.")
            _add_paragraph(document, "4. CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "Conclusion text for the manuscript.")
            _add_paragraph(document, "5. REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(
                document,
                '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
                9,
            )
            document.save(path)

            rules = _jiwe_rules()
            before = ManuscriptChecker(rules).load_manuscript(str(path)).check_all()
            heading_descriptions = [
                issue.description
                for issue in before.issues_by_category.get("headings", [])
            ]

            fixer = AutoFixer(rules, before.classifications, before.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertIn("Heading alignment does not match template", heading_descriptions)
        fixed_headings = {
            paragraph.text: paragraph.alignment
            for paragraph in fixed.paragraphs
            if paragraph.text in {"1. INTRODUCTION", "1.1 Experimental Setup"}
        }
        self.assertEqual(fixed_headings["1. INTRODUCTION"], WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(fixed_headings["1.1 Experimental Setup"], WD_ALIGN_PARAGRAPH.LEFT)

    def test_jiwe_declaration_left_content_is_not_body_alignment_mismatch(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "declaration_alignment.docx"
            document = Document()
            _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "A Test Paper Title for Declaration Alignment", 24, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "Abstract - " + "word " * 210, 9)
            _add_paragraph(document, "Keywords - Template, Checking, Rules, References, Formatting", 9)
            _add_paragraph(document, "1. INTRODUCTION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "This body paragraph is correctly justified.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(document, "4. CONCLUSION", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "Conclusion text for the manuscript.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(document, "FUNDING STATEMENT", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(
                document,
                "The authors received no funding from any party for the research and publication of this article.",
                10,
                WD_ALIGN_PARAGRAPH.LEFT,
            )
            _add_paragraph(document, "CONFLICT OF INTERESTS", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "No conflict of interests were disclosed.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(document, "5. REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(
                document,
                '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
                9,
            )
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()

        body_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        self.assertNotIn("Body text alignment does not match template", body_descriptions)

    def test_jiwe_inline_font_instruction_does_not_break_heading_capitalization(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "heading_instruction.docx"
            document = Document()
            _add_paragraph(document, "Journal of Informatics and", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Web Engineering", 24, WD_ALIGN_PARAGRAPH.CENTER, True)
            _add_paragraph(document, "Vol. 5 No. 2 (June 2026) eISSN: 2821-370X", 10, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "A Test Paper Title for Heading Instruction", 24, WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(document, "Abstract - " + "word " * 210, 9)
            _add_paragraph(document, "Keywords - Template, Checking, Rules, References, Formatting", 9)
            _add_paragraph(document, "INTRODUCTION (10-Font size, Times New Roman)", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "This body paragraph is correctly justified.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(document, "CONCLUSION (10-Font size, Times New Roman)", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(document, "Conclusion text for the manuscript.", 10, WD_ALIGN_PARAGRAPH.JUSTIFY)
            _add_paragraph(document, "REFERENCES", 10, WD_ALIGN_PARAGRAPH.LEFT, True)
            _add_paragraph(
                document,
                '[1]\tA. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.',
                9,
            )
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()

        heading_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("headings", [])
        ]
        self.assertNotIn("Heading capitalization does not match template", heading_descriptions)

    def test_checker_reports_body_spacing_and_reference_indent_mismatches(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spacing_issues.docx"
            _save_spacing_issue_document(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()

        body_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("body_text", [])
        ]
        reference_descriptions = [
            issue.description
            for issue in result.issues_by_category.get("references", [])
        ]

        self.assertIn("Body paragraph spacing after does not match template", body_descriptions)
        self.assertIn("Reference paragraph spacing after does not match template", reference_descriptions)
        self.assertIn("Reference hanging indent does not match template", reference_descriptions)

    def test_auto_fix_applies_body_spacing_and_reference_indent(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spacing_fix.docx"
            _save_spacing_issue_document(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        body = fixed.paragraphs[7]
        reference = fixed.paragraphs[11]
        self.assertAlmostEqual(body.paragraph_format.space_after.pt, 7.5)
        self.assertAlmostEqual(reference.paragraph_format.line_spacing, 1.15)
        self.assertAlmostEqual(reference.paragraph_format.space_after.pt, 10.0)
        self.assertAlmostEqual(abs(reference.paragraph_format.first_line_indent.inches), 0.44, places=2)
        self.assertEqual(reference._p.pPr.ind.get(qn("w:hanging")), "640")

    def test_reference_explicit_left_with_matching_hanging_is_accepted(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_citation_manager_indent.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            reference = document.paragraphs[11]
            reference.text = '[1] A. Author, "Article title," Journal of Testing, vol. 1, no. 1, pp. 1-5, 2026.'
            reference.paragraph_format.left_indent = Inches(0.4444444444444444)
            reference.paragraph_format.first_line_indent = Inches(-0.4444444444444444)
            reference.paragraph_format.line_spacing = 1.15
            reference.paragraph_format.space_after = Pt(10)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            reference_descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        fixed_reference = fixed.paragraphs[11]
        self.assertNotIn("Reference left indent does not match template", reference_descriptions)
        self.assertIn("Reference number should be followed by a tab", reference_descriptions)
        self.assertAlmostEqual(fixed_reference.paragraph_format.left_indent.inches, 0.4444444444444444)
        self.assertAlmostEqual(fixed_reference.paragraph_format.first_line_indent.inches, -0.4444444444444444)
        self.assertTrue(fixed_reference.text.startswith("[1]\t"))

    def test_highlighted_reference_layout_fix_marks_number_only(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_highlight_marker.docx"
            _save_spacing_issue_document(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            highlighted = Document(BytesIO(fixer.get_highlighted_document_bytes()))

        reference = highlighted.paragraphs[11]
        highlighted_runs = [
            run.text.strip()
            for run in reference.runs
            if run.text.strip() and run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        ]
        unhighlighted_runs = [
            run.text
            for run in reference.runs
            if run.text.strip() and run.font.highlight_color != WD_COLOR_INDEX.YELLOW
        ]
        raw_highlighted_runs = [
            run.text
            for run in reference.runs
            if run.text and run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        ]
        self.assertEqual(highlighted_runs, ["[1]"])
        self.assertEqual(raw_highlighted_runs, ["[1]"])
        self.assertTrue(any("A. Author" in text for text in unhighlighted_runs), reference.text)

    def test_auto_fix_replaces_body_manual_line_breaks_before_justifying(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "manual_line_breaks.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            body = document.paragraphs[7]
            body.clear()
            run = body.add_run("This sentence should wrap naturally,")
            run.add_break(WD_BREAK.LINE)
            body.add_run("not be forced onto a stretched justified line.")
            body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("body_text", [])
            ]
            self.assertIn(
                "Body paragraph contains manual line breaks that can stretch justified text",
                descriptions,
            )

            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertFalse(paragraph_has_manual_line_breaks(fixed.paragraphs[7]))
        self.assertIn("naturally, not be forced", fixed.paragraphs[7].text)

    def test_auto_fix_clears_extra_reference_left_indent(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_left_indent_fix.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            reference = document.paragraphs[11]
            reference.paragraph_format.left_indent = Inches(0.89)
            reference.paragraph_format.first_line_indent = Inches(-0.44)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]
            self.assertIn("Reference left indent does not match template", descriptions)

            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        fixed_reference = fixed.paragraphs[11]
        self.assertIsNone(fixed_reference.paragraph_format.left_indent)
        indentation = fixed_reference._p.pPr.ind
        self.assertIsNone(indentation.get(qn("w:left")))
        self.assertAlmostEqual(abs(fixed_reference.paragraph_format.first_line_indent.inches), 0.44, places=2)

    def test_checker_reports_reference_number_space_instead_of_tab(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_number_tab.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            reference = document.paragraphs[11]
            reference.runs[0].text = reference.runs[0].text.replace("[1]\t", "[1] ", 1)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("references", [])
        ]
        self.assertIn("Reference number should be followed by a tab", descriptions)

    def test_auto_fix_replaces_reference_number_spaces_with_tab(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_number_tab_fix.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            reference = document.paragraphs[11]
            reference.runs[0].text = reference.runs[0].text.replace("[1]\t", "[1] ", 1)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertTrue(fixed.paragraphs[11].text.startswith("[1]\t"))

    def test_auto_fix_collapses_duplicate_tabs_after_reference_number(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "reference_number_duplicate_tab_fix.docx"
            _save_spacing_issue_document(path)
            document = Document(path)
            reference = document.paragraphs[11]
            reference.runs[0].text = "[1]\t"
            reference.add_run("\t")
            reference.add_run('A. Author, "Article title," Journal of Testing, 2026.')
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            descriptions = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertIn("Reference number should be followed by a tab", descriptions)
        self.assertTrue(fixed.paragraphs[11].text.startswith("[1]\t"))
        self.assertFalse(fixed.paragraphs[11].text.startswith("[1]\t\t"))

    def test_checker_reports_caption_title_case_mismatch(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_case.docx"
            document = Document()
            _add_drawing_paragraph(document)
            caption = _add_paragraph(
                document,
                "Figure 3. top 20 features ranked by their chi-squared scores with the target variable",
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()

        descriptions = [
            issue.description
            for issue in result.issues_by_category.get("figures", [])
        ]
        self.assertIn("Figure caption capitalization does not match template", descriptions)

    def test_auto_fix_applies_caption_title_case(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_case_fix.docx"
            document = Document()
            _add_drawing_paragraph(document)
            _add_paragraph(
                document,
                "Figure 3. top 20 features ranked by their chi-squared scores with the target variable",
            )
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual(
            fixed.paragraphs[1].text,
            "Figure 3. Top 20 Features Ranked by Their Chi-Squared Scores with the Target Variable",
        )

    def test_caption_title_case_keeps_joining_words_lowercase(self):
        self.assertEqual(
            to_journal_caption_title_case(
                "Table 2. bootstrap metrics for model performance with features from training data via resampling"
            ),
            "Table 2. Bootstrap Metrics for Model Performance with Features from Training Data via Resampling",
        )

    def test_caption_title_case_keeps_hyphenated_joining_words_lowercase(self):
        self.assertEqual(
            to_journal_caption_title_case(
                "Figure 10. real-time attack alerts for state-of-the-art methods"
            ),
            "Figure 10. Real-Time Attack Alerts for State-of-the-Art Methods",
        )

    def test_caption_title_case_capitalizes_word_after_sentence_boundary(self):
        self.assertEqual(
            to_journal_caption_title_case(
                "Figure 2. distribution of learning gains. the treatment distribution with control group"
            ),
            "Figure 2. Distribution of Learning Gains. The Treatment Distribution with Control Group",
        )

    def test_caption_auto_fix_removes_bold_and_preserves_title_case_runs(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_case_mixed_bold.docx"
            document = Document()
            _add_drawing_paragraph(document)
            caption = document.add_paragraph()
            label = caption.add_run("Figure 1. ")
            label.font.name = "Times New Roman"
            label.font.size = Pt(10)
            label.font.bold = True
            body = caption.add_run("dual-phase evidence chain")
            body.font.name = "Times New Roman"
            body.font.size = Pt(10)
            body.font.bold = False
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            descriptions = [
                issue.description
                for issues in result.issues_by_category.values()
                for issue in issues
            ]
            self.assertIn("Figure caption bold formatting does not match template", descriptions)

            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        visible_runs = [
            (run.text, run.font.bold)
            for run in fixed.paragraphs[1].runs
            if run.text.strip()
        ]
        self.assertEqual(fixed.paragraphs[1].text, "Figure 1. Dual-Phase Evidence Chain")
        self.assertTrue(visible_runs)
        self.assertTrue(all(bold is False for _, bold in visible_runs), visible_runs)

    def test_table_caption_auto_fix_removes_bold(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "table_caption_bold.docx"
            document = Document()
            caption = _add_paragraph(
                document,
                "Table 2. Bootstrap Metrics for Model Performance",
                bold=True,
            )
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_table(rows=2, cols=2)
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            descriptions = [
                issue.description
                for issues in result.issues_by_category.values()
                for issue in issues
            ]
            self.assertIn("Table caption bold formatting does not match template", descriptions)

            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        visible_runs = [
            (run.text, run.font.bold)
            for run in fixed.paragraphs[0].runs
            if run.text.strip()
        ]
        self.assertTrue(visible_runs)
        self.assertTrue(all(bold is False for _, bold in visible_runs), visible_runs)

    def test_caption_title_case_preserves_acronyms(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_acronyms.docx"
            document = Document()
            _add_drawing_paragraph(document)
            _add_paragraph(
                document,
                "Figure 1. AI-Based OWASP WCSS rPPG and ROI Results",
            )
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual(
            fixed.paragraphs[1].text,
            "Figure 1. AI-Based OWASP WCSS rPPG and ROI Results",
        )

    def test_caption_title_case_preserves_numeric_expressions_and_vs(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "caption_numeric_expression.docx"
            document = Document()
            _add_drawing_paragraph(document)
            _add_paragraph(
                document,
                "Figure 1. Visualization of the 36x36 ROI preprocessing pipeline vs. baseline",
            )
            document.save(path)

            result = ManuscriptChecker(_jiwe_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_jiwe_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))

        self.assertEqual(
            fixed.paragraphs[1].text,
            "Figure 1. Visualization of the 36x36 ROI Preprocessing Pipeline vs. Baseline",
        )


if __name__ == "__main__":
    unittest.main()
