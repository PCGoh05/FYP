import tempfile
import unittest
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {"font_name": "Times New Roman", "font_size": 10, "line_spacing": 1.0},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "references": {"font_name": "Times New Roman", "font_size": 9},
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


class TitleFormattingTest(unittest.TestCase):
    def test_missing_explicit_italic_is_not_reported_when_title_should_be_not_italic(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "title_not_italic.docx"
            document = Document()
            document.add_paragraph("Journal of Informatics and")
            document.add_paragraph("Web Engineering")
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = title.add_run("A Test Paper Title for Format Validation")
            run.font.name = "Times New Roman"
            run.font.size = Pt(24)
            document.add_paragraph("Abstract - This is the abstract.")
            document.add_paragraph("Keywords - checking, template")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()

            title_issues = [
                issue.description
                for issue in result.issues_by_category.get("title", [])
            ]

        self.assertNotIn(
            "Title italic formatting does not match template",
            title_issues,
        )


if __name__ == "__main__":
    unittest.main()
