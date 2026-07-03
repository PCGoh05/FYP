import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker
from modules.utils import get_paragraph_font_info, get_sdt_reference_paragraphs


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {"font_name": "Times New Roman", "font_size": 10, "line_spacing": 1.0},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "subheading": {"font_name": "Times New Roman", "font_size": 10, "bold": False, "italic": True},
        "references": {"font_name": "Times New Roman", "font_size": 9},
        "reference": {
            "font_name": "Times New Roman",
            "font_size": 9,
            "bold": False,
            "alignment": "JUSTIFY",
            "line_spacing": 1.0,
        },
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


def _add_minimal_front_matter(document):
    document.add_paragraph("Journal of Informatics and")
    document.add_paragraph("Web Engineering")
    document.add_paragraph("Vol. 5 No. 2 (June 2026)\teISSN: 2821-370X")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("A Test Paper Title for Format Validation")
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(24)
    document.add_paragraph("Abstract - This is the abstract.")
    document.add_paragraph("Keywords - checking, template")
    document.add_paragraph("1. INTRODUCTION")
    document.add_paragraph("Body text.")


def _append_sdt_paragraph(
    document,
    text,
    font_name="Calibri",
    font_size=11,
    alignment="CENTER",
):
    sdt = OxmlElement("w:sdt")
    content = OxmlElement("w:sdtContent")
    paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    justification = OxmlElement("w:jc")
    justification.set(qn("w:val"), alignment.lower())
    paragraph_properties.append(justification)
    paragraph.append(paragraph_properties)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), str(int(font_size * 2)))
    run_properties.append(fonts)
    run_properties.append(size)
    run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.set(qn("xml:space"), "preserve")
    text_node.text = text
    run.append(text_node)
    paragraph.append(run)
    content.append(paragraph)
    sdt.append(content)
    body = document.element.body
    body.insert(max(len(body) - 1, 0), sdt)


class JiweEdgeCasesTest(unittest.TestCase):
    def test_numbered_subheading_uses_italic_not_bold_rule(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "subheading.docx"
            document = Document()
            _add_minimal_front_matter(document)
            subheading = document.add_paragraph()
            run = subheading.add_run("4.1. Experimental Setup")
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.font.italic = True
            document.add_paragraph("The experiment text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            heading_issues = [
                issue.description
                for issue in result.issues_by_category.get("headings", [])
                if "4.1" in issue.location or "4.1" in issue.text_preview
            ]

        self.assertNotIn("Heading bold formatting does not match template", heading_issues)

    def test_references_inside_word_content_controls_are_detected(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sdt_references.docx"
            document = Document()
            _add_minimal_front_matter(document)
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            _append_sdt_paragraph(document, "[1]")
            _append_sdt_paragraph(document, "A. Author, Article title, Journal, 2025, doi: 10.1000/example.")
            document.add_paragraph("BIOGRAPHIES OF AUTHORS")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            reference_issues = [
                issue.description
                for issue in result.issues_by_category.get("references", [])
            ]

        self.assertNotIn("No references found in document", reference_issues)
        self.assertIn("Reference font does not match template", reference_issues)
        self.assertIn("Reference font size does not match template", reference_issues)
        self.assertIn("Reference alignment does not match template", reference_issues)

    def test_auto_fixer_formats_references_inside_word_content_controls(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sdt_reference_fix.docx"
            document = Document()
            _add_minimal_front_matter(document)
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            _append_sdt_paragraph(document, "[1]")
            _append_sdt_paragraph(document, "A. Author, Article title, Journal of Testing, 2025.")
            document.add_paragraph("BIOGRAPHIES OF AUTHORS")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()
            fixed = Document(BytesIO(fixer.get_fixed_document_bytes()))
            references = get_sdt_reference_paragraphs(fixed)

        self.assertEqual([paragraph.text for paragraph in references], [
            "[1]",
            "A. Author, Article title, Journal of Testing, 2025.",
        ])
        for paragraph in references:
            font_info = get_paragraph_font_info(paragraph)
            self.assertEqual(font_info["font_name"], "Times New Roman")
            self.assertEqual(font_info["font_size"], 9)
            self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)

    def test_highlighted_document_marks_reference_content_control_changes(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "sdt_reference_highlight.docx"
            document = Document()
            _add_minimal_front_matter(document)
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            _append_sdt_paragraph(document, "[1]")
            _append_sdt_paragraph(document, "A. Author, Article title, Journal of Testing, 2025.")
            document.add_paragraph("BIOGRAPHIES OF AUTHORS")
            document.save(path)

            result = ManuscriptChecker(_rules()).load_manuscript(str(path)).check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()

            highlighted = Document(BytesIO(fixer.get_highlighted_document_bytes()))
            references = get_sdt_reference_paragraphs(highlighted)

        self.assertEqual(len(references), 2)
        highlighted_runs = [
            run
            for paragraph in references
            for run in paragraph.runs
            if run.text.strip() and run.font.highlight_color == WD_COLOR_INDEX.YELLOW
        ]
        self.assertTrue(highlighted_runs)


if __name__ == "__main__":
    unittest.main()
