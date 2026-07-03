import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.shared import RGBColor

from modules.auto_fixer import AutoFixer
from modules.manuscript_checker import ManuscriptChecker


def _rules():
    return {
        "_profile": {"name": "JIWE"},
        "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
        "journal_header": {
            "font_name": "Palatino Linotype",
            "font_size": 24,
            "bold": True,
            "alignment": "CENTER",
        },
        "title": {
            "font_name": "Times New Roman",
            "font_size": 24,
            "bold": None,
            "italic": False,
            "alignment": "CENTER",
        },
        "body": {"font_name": "Times New Roman", "font_size": 10, "line_spacing": 1.0},
        "heading": {"font_name": "Times New Roman", "font_size": 10, "bold": True},
        "references": {"font_name": "Times New Roman", "font_size": 9},
        "caption": {"font_name": "Times New Roman", "font_size": 10},
    }


def _save_unstable_docx(path: Path):
    document = Document()
    header = document.sections[0].header.paragraphs[0]
    header.text = (
        "Journal of Informatics and Web Engineering "
        "\t\t\t\t             Vol. 3 No. 3 (January 2026)"
    )

    p = document.add_paragraph("Journal of Informatics and")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run()
    p = document.add_paragraph("\t\tWeb Engineering")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
    document.add_paragraph("Practicality Study of Hybrid Voice Generation Model Based on Meta-LoRA and VoxCPM")
    document.add_paragraph("Abstract - This is an abstract.")
    document.add_paragraph("Keywords - template, checking")
    document.add_paragraph("INTRODUCTION")
    document.add_paragraph("Body text.")
    document.add_paragraph("CONCLUSION")
    document.add_paragraph("Conclusion text.")
    document.add_paragraph("REFERENCES")
    document.add_paragraph("[1] Reference text.")
    document.save(path)


class HeaderLayoutStabilityTest(unittest.TestCase):
    def test_checker_flags_manual_tabs_that_can_shift_template_layout(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "unstable.docx"
            _save_unstable_docx(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            descriptions = [
                issue.description
                for issues in result.issues_by_category.values()
                for issue in issues
            ]

        self.assertIn(
            "Journal header contains manual tab indentation that can shift layout",
            descriptions,
        )
        self.assertIn(
            "Page header uses multiple manual tabs/spaces that can wrap in Word",
            descriptions,
        )

    def test_auto_fixer_normalizes_manual_tabs_in_journal_and_page_header(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "unstable.docx"
            _save_unstable_docx(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixed_doc, changes = fixer.fix_all()

            self.assertEqual(fixed_doc.paragraphs[1].text, "Web Engineering")
            self.assertEqual(
                fixed_doc.sections[0].header.paragraphs[0].text,
                "Journal of Informatics and Web Engineering\tVol. 3 No. 3 (January 2026)",
            )
            self.assertTrue(
                any(change.property_name == "manual_tabs" for change in changes)
            )
            header_changes = [
                change
                for change in changes
                if change.change_type == "page_header" and change.property_name == "manual_tabs"
            ]
            self.assertTrue(header_changes)
            self.assertEqual(
                header_changes[0].current_value,
                "Manual tabs/spaces between left and right header text",
            )
            self.assertEqual(
                header_changes[0].target_value,
                "Single right-aligned tab stop; visible header text unchanged",
            )

    def test_auto_fixer_trims_journal_header_spaces_without_losing_run_formatting(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "spaced_header.docx"
            document = Document()
            header = document.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = header.add_run(" Journal of Informatics and ")
            run.font.name = "Palatino Linotype"
            run.font.size = Pt(24)
            run.font.bold = True
            second = document.add_paragraph("Web Engineering")
            second.alignment = WD_ALIGN_PARAGRAPH.CENTER
            second.runs[0].font.name = "Palatino Linotype"
            second.runs[0].font.size = Pt(24)
            second.runs[0].font.bold = True
            document.add_paragraph("Vol. 3 No. 3 (January 2026)\teISSN: 2821-370X")
            document.add_paragraph("A Test Paper Title for Format Validation")
            document.add_paragraph("Abstract - This is an abstract.")
            document.add_paragraph("Keywords - template, checking")
            document.add_paragraph("INTRODUCTION")
            document.add_paragraph("Body text.")
            document.add_paragraph("CONCLUSION")
            document.add_paragraph("Conclusion text.")
            document.add_paragraph("REFERENCES")
            document.add_paragraph("[1] Reference text.")
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixed_doc, _ = fixer.fix_all()

            fixed_header = fixed_doc.paragraphs[0]
            self.assertEqual(fixed_header.text, "Journal of Informatics and")
            visible_runs = [run for run in fixed_header.runs if run.text.strip()]
            self.assertEqual(len(visible_runs), 1)
            self.assertEqual(visible_runs[0].font.name, "Palatino Linotype")
            self.assertEqual(visible_runs[0].font.size.pt, 24)
            self.assertTrue(visible_runs[0].font.bold)

    def test_auto_fixer_preserves_stable_single_tab_page_headers(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "mixed_headers.docx"
            _save_unstable_docx(path)

            document = Document(path)
            document.sections[0].first_page_header.paragraphs[0].text = "Left\tMiddle\tRight"
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixed_doc, changes = fixer.fix_all()

            self.assertEqual(
                fixed_doc.sections[0].header.paragraphs[0].text,
                "Journal of Informatics and Web Engineering\tVol. 3 No. 3 (January 2026)",
            )
            self.assertEqual(
                fixed_doc.sections[0].first_page_header.paragraphs[0].text,
                "Left\tMiddle\tRight",
            )
            changed_locations = [
                change.location for change in changes if change.property_name == "manual_tabs"
            ]
            self.assertNotIn("First Page Header (Section 1)", changed_locations)

    def test_auto_fixer_normalizes_page_header_without_losing_run_formatting(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "styled_header.docx"
            _save_unstable_docx(path)
            document = Document(path)
            header = document.sections[0].header.paragraphs[0]
            header.clear()
            run = header.add_run(
                "Journal of Informatics and Web Engineering "
                "\t\t\t\t             Vol. 3 No. 3 (January 2026)"
            )
            run.font.name = "Times New Roman"
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(128, 128, 128)
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixed_doc, _ = fixer.fix_all()

            fixed_header = fixed_doc.sections[0].header.paragraphs[0]
            visible_runs = [run for run in fixed_header.runs if run.text.strip()]
            self.assertEqual(
                fixed_header.text,
                "Journal of Informatics and Web Engineering\tVol. 3 No. 3 (January 2026)",
            )
            self.assertEqual(len(visible_runs), 1)
            self.assertEqual(visible_runs[0].font.name, "Times New Roman")
            self.assertEqual(visible_runs[0].font.size.pt, 9)
            self.assertTrue(visible_runs[0].font.italic)

    def test_auto_fixer_preserves_header_line_drawing_when_normalizing_tabs(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "header_line.docx"
            _save_unstable_docx(path)
            document = Document(path)
            header = document.sections[0].header.paragraphs[0]
            drawing_run = header.add_run()
            drawing_run._r.append(OxmlElement("w:drawing"))
            document.save(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixed_doc, _ = fixer.fix_all()

            fixed_header = fixed_doc.sections[0].header.paragraphs[0]
            self.assertIn(
                "Journal of Informatics and Web Engineering\tVol. 3 No. 3 (January 2026)",
                fixed_header.text,
            )
            self.assertTrue(
                any("<w:drawing" in run._r.xml for run in fixed_header.runs)
            )

    def test_highlighted_document_does_not_mark_hidden_page_header_spacing(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "unstable.docx"
            _save_unstable_docx(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()

            highlighted_doc = Document(BytesIO(fixer.get_highlighted_document_bytes()))
            header_runs = [
                run
                for run in highlighted_doc.sections[0].header.paragraphs[0].runs
                if run.text.strip()
            ]

            self.assertTrue(header_runs)
            self.assertIsNone(header_runs[0].font.highlight_color)
            self.assertEqual(
                highlighted_doc.sections[0].header.paragraphs[0].text,
                "Journal of Informatics and Web Engineering\tVol. 3 No. 3 (January 2026)",
            )

    def test_highlighted_document_preserves_original_body_without_inserted_summary(self):
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "unstable.docx"
            _save_unstable_docx(path)

            checker = ManuscriptChecker(_rules()).load_manuscript(str(path))
            result = checker.check_all()
            fixer = AutoFixer(_rules(), result.classifications, result.issues_by_category)
            fixer.load_manuscript(str(path))
            fixer.fix_all()

            highlighted_doc = Document(BytesIO(fixer.get_highlighted_document_bytes()))
            paragraph_texts = [paragraph.text for paragraph in highlighted_doc.paragraphs]

            self.assertNotIn("Highlighted Formatting Changes", paragraph_texts)
            self.assertEqual(paragraph_texts[0], "Journal of Informatics and")


if __name__ == "__main__":
    unittest.main()
