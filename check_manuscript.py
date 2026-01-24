"""
Diagnostic script to check manuscript for strikethrough and revision marks
"""
import os
from docx import Document
from docx.oxml.ns import qn

def check_document_for_issues(doc_path):
    """Check document for strikethrough, revisions, and other formatting issues"""
    
    if not os.path.exists(doc_path):
        print(f"File not found: {doc_path}")
        return
    
    doc = Document(doc_path)
    
    print("=" * 80)
    print(f"DOCUMENT ANALYSIS: {doc_path}")
    print("=" * 80)
    
    # Check for strikethrough in paragraphs
    strike_found = []
    underline_found = []
    
    for i, para in enumerate(doc.paragraphs[:20]):  # Check first 20 paragraphs
        text = para.text.strip()
        if not text:
            continue
        
        for run in para.runs:
            if run.font.strike:
                strike_found.append({
                    'para': i,
                    'text': run.text[:50],
                    'strike': run.font.strike
                })
            
            if run.font.underline:
                underline_found.append({
                    'para': i,
                    'text': run.text[:50],
                    'underline': str(run.font.underline)
                })
    
    # Check for Track Changes / Revisions
    print("\n1. STRIKETHROUGH FOUND:")
    if strike_found:
        for item in strike_found:
            print(f"   Para {item['para']}: '{item['text']}' - Strike: {item['strike']}")
    else:
        print("   None found")
    
    print("\n2. UNDERLINE FOUND:")
    if underline_found:
        for item in underline_found:
            print(f"   Para {item['para']}: '{item['text']}' - Underline: {item['underline']}")
    else:
        print("   None found")
    
    # Check XML for revision marks (Track Changes)
    print("\n3. CHECKING FOR TRACK CHANGES (Revisions)...")
    body = doc._body._body
    
    # Look for w:del (deletions) and w:ins (insertions)
    deletions = body.findall('.//' + qn('w:del'))
    insertions = body.findall('.//' + qn('w:ins'))
    
    print(f"   Deletions (w:del): {len(deletions)}")
    print(f"   Insertions (w:ins): {len(insertions)}")
    
    if deletions:
        print("\n   Deletion examples:")
        for d in deletions[:5]:
            print(f"      - {d.text[:70] if d.text else '(no text)'}")
    
    print("\n4. FIRST 10 PARAGRAPHS PREVIEW:")
    for i, para in enumerate(doc.paragraphs[:10]):
        text = para.text.strip()
        if text:
            preview = text[:60] + "..." if len(text) > 60 else text
            print(f"   [{i}] {preview}")

# Check sample manuscript
if __name__ == "__main__":
    import sys
    
    # Try different possible locations
    paths_to_check = [
        r"sample\JIWE_Template.docx",
        r"sample\sample_manuscript.docx",
    ]
    
    # Also check sample folder for any docx files
    if os.path.exists("sample"):
        for f in os.listdir("sample"):
            if f.endswith(".docx") and f not in [os.path.basename(p) for p in paths_to_check]:
                paths_to_check.append(f"sample\\{f}")
    
    for path in paths_to_check:
        if os.path.exists(path):
            check_document_for_issues(path)
            print("\n" + "=" * 80 + "\n")
