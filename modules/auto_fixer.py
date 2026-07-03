"""
Auto-Fixer Module
Automatically fixes formatting issues while preserving special formatting
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from copy import deepcopy
import re
from io import BytesIO

from .utils import (
    load_document, get_paragraph_text, get_paragraph_alignment, truncate_text,
    is_font_equivalent, get_run_font_info, get_sdt_reference_paragraphs,
    classify_author_info_role, get_space_after_pt, get_direct_left_indent_inches, get_hanging_indent_inches,
    to_journal_caption_title_case, paragraph_has_manual_line_breaks,
    replace_manual_line_breaks_with_spaces
)
from .paragraph_classifier import (
    ParagraphType, ClassifiedParagraph
)
from config import ALIGNMENT_REVERSE_MAP


@dataclass
class ChangeRecord:
    """Record of a change made during auto-fix"""
    paragraph_index: int
    location: str
    change_type: str
    before: str
    after: str
    text_preview: str
    property_name: str = "formatting"
    current_value: str = ""
    target_value: str = ""
    paragraph_type: str = ""
    evidence: str = ""


@dataclass
class PostFixValidationResult:
    """Result of checking the corrected document after auto-fix."""
    is_safe: bool
    before_issues: int
    after_issues: int
    before_score: float
    after_score: float
    issue_delta: int
    score_delta: float
    new_or_increased_categories: Dict[str, int]
    message: str


def _category_issue_counts(result: Any) -> Dict[str, int]:
    """Return issue counts per category from a checker result."""
    issues_by_category = getattr(result, "issues_by_category", {}) or {}
    return {
        category: len(issues)
        for category, issues in issues_by_category.items()
        if issues
    }


def validate_post_fix_result(before_result: Any, after_result: Any) -> PostFixValidationResult:
    """Compare pre-fix and post-fix checker results."""
    before_issues = int(getattr(before_result, "total_issues", 0) or 0)
    after_issues = int(getattr(after_result, "total_issues", 0) or 0)
    before_score = float(getattr(before_result, "compliance_score", 0.0) or 0.0)
    after_score = float(getattr(after_result, "compliance_score", 0.0) or 0.0)
    issue_delta = after_issues - before_issues
    score_delta = round(after_score - before_score, 2)

    before_counts = _category_issue_counts(before_result)
    after_counts = _category_issue_counts(after_result)
    increased_categories = {
        category: count - before_counts.get(category, 0)
        for category, count in after_counts.items()
        if count > before_counts.get(category, 0)
    }

    score_dropped = after_score < before_score
    is_safe = issue_delta <= 0 and not score_dropped
    if is_safe:
        message = "Post-fix validation did not increase detected issues."
    elif issue_delta > 0:
        message = "Post-fix validation found more issues after auto-fix. Review the corrected document before using it."
    else:
        message = "Post-fix validation found a lower compliance index after auto-fix. Review the corrected document before using it."

    return PostFixValidationResult(
        is_safe=is_safe,
        before_issues=before_issues,
        after_issues=after_issues,
        before_score=before_score,
        after_score=after_score,
        issue_delta=issue_delta,
        score_delta=score_delta,
        new_or_increased_categories=increased_categories,
        message=message,
    )


def validate_fixed_document(rules: Dict[str, Any], fixed_doc_bytes: bytes):
    """Run the checker again on a corrected document."""
    from .manuscript_checker import ManuscriptChecker

    checker = ManuscriptChecker(rules, None)
    checker.load_manuscript(BytesIO(fixed_doc_bytes))
    return checker.check_all()


def summarize_remaining_issues(result: Any) -> List[Dict[str, Any]]:
    """Create a compact table summary of issues remaining after auto-fix."""
    issues_by_category = getattr(result, "issues_by_category", {}) or {}
    rows = []
    for category, issues in issues_by_category.items():
        if not issues:
            continue
        first_issue = issues[0]
        rows.append({
            "Category": category.replace("_", " ").title(),
            "Count": len(issues),
            "First Location": getattr(first_issue, "location", ""),
            "First Issue": getattr(first_issue, "description", ""),
        })
    return sorted(rows, key=lambda row: (-row["Count"], row["Category"]))


class AutoFixer:
    """
    Automatically fixes formatting issues in a manuscript
    while preserving special formatting (italic, underline, subscript, etc.)
    """
    
    def __init__(
        self,
        rules: Dict[str, Any],
        classifications: List[ClassifiedParagraph],
        issues_by_category: Optional[Dict[str, List[Any]]] = None,
    ):
        """
        Initialize auto-fixer with template rules and paragraph classifications
        
        Args:
            rules: Template formatting rules
            classifications: List of classified paragraphs from ParagraphClassifier
        """
        self.rules = rules
        self.classifications = classifications
        self.issues_by_category = issues_by_category or {}
        self.changes: List[ChangeRecord] = []
        self.document = None
        self.original_document = None
        self._classification_map: Dict[int, ClassifiedParagraph] = {}
        self._issue_map: Dict[int, set] = {}
        self._issue_property_map: Dict[int, set] = {}
        self._global_issue_categories: set = set()
        self._global_issue_property_map: Dict[str, set] = {}
        self._has_explicit_issues = bool(issues_by_category)
        self._build_issue_map()
    
    def load_manuscript(self, file_path_or_bytes):
        """Load the manuscript to fix"""
        self.document = load_document(file_path_or_bytes)
        self.original_document = deepcopy(self.document)
        # Create index map for quick lookup
        self._classification_map = {cp.index: cp for cp in self.classifications}
        return self

    def _build_issue_map(self):
        """Build paragraph-to-category and paragraph-to-property maps from detected issues."""
        self._issue_map = {}
        self._issue_property_map = {}
        self._global_issue_categories = set()
        self._global_issue_property_map = {}
        for category, issues in self.issues_by_category.items():
            for issue in issues:
                para_index = getattr(issue, "paragraph_index", -1)
                if para_index is None or para_index < 0:
                    self._global_issue_categories.add(category)
                    properties = self._infer_global_issue_properties(category, issue)
                    if properties:
                        self._global_issue_property_map.setdefault(category, set()).update(properties)
                    continue
                self._issue_map.setdefault(para_index, set()).add(category)
                properties = self._infer_issue_properties(issue)
                if properties:
                    self._issue_property_map.setdefault(para_index, set()).update(properties)

    def _infer_issue_properties(self, issue: Any) -> set:
        """Infer one or more affected formatting properties from a checker issue."""
        property_name = self._infer_issue_property(issue)
        if property_name:
            return {property_name}

        description = (getattr(issue, "description", "") or "").lower()
        current_value = (getattr(issue, "current_value", "") or "").lower()
        if "formatting does not match template" not in description:
            return set()

        properties = set()
        if "font " in current_value:
            properties.add("font_name")
        if "size " in current_value:
            properties.add("font_size")
        if "bold" in current_value:
            properties.add("bold")
        if "italic" in current_value:
            properties.add("italic")
        if "alignment" in current_value:
            properties.add("alignment")
        return properties

    def _infer_issue_property(self, issue: Any) -> Optional[str]:
        """Infer the affected formatting property from a checker issue."""
        description = (getattr(issue, "description", "") or "").lower()
        location = (getattr(issue, "location", "") or "").lower()

        if "number" in description or "heading number" in location:
            if "bold" in description:
                return "number_bold"
            if "font size" in description or "size" in description:
                return "number_font_size"
            if "font" in description:
                return "number_font_name"

        if "line spacing" in description:
            return "line_spacing"
        if "spacing after" in description or "after spacing" in description:
            return "space_after"
        if "left indent" in description:
            return "left_indent"
        if "hanging indent" in description:
            return "hanging_indent"
        if "manual line break" in description or "manual line breaks" in description:
            return "manual_line_breaks"
        if "manual tab" in description or "manual tabs" in description:
            return "manual_tabs"
        if "capitalization" in description or "all capital" in description:
            return "capitalization"
        if "alignment" in description:
            return "alignment"
        if "font size" in description or " size " in f" {description} ":
            return "font_size"
        if "font" in description:
            return "font_name"
        if "bold" in description:
            return "bold"
        if "italic" in description:
            return "italic"
        return None

    def _infer_global_issue_properties(self, category: str, issue: Any) -> set:
        """Infer properties represented by an aggregate category-level issue."""
        properties = self._infer_issue_properties(issue)
        if properties:
            return properties

        description = (getattr(issue, "description", "") or "").lower()
        if category == "body_text" and "body text formatting issues" in description:
            return {"font_name", "font_size", "bold"}
        if category == "body_text" and "manual line break" in description:
            return {"manual_line_breaks"}
        if category == "line_spacing":
            return {"line_spacing"}
        if category == "other" and "manual tab" in (getattr(issue, "description", "") or "").lower():
            return {"manual_tabs"}
        return set()

    def _should_fix_category(self, index: int, categories: List[str]) -> bool:
        """Return True when a paragraph has a detected issue for the target categories."""
        if not self._has_explicit_issues:
            return True
        found_categories = self._issue_map.get(index, set())
        return any(
            category in found_categories or category in self._global_issue_categories
            for category in categories
        )

    def _allowed_properties_for(
        self,
        index: int,
        categories: Optional[List[str]] = None,
        fallback: Optional[List[str]] = None,
    ) -> Optional[set]:
        """Return properties that may be fixed for a paragraph."""
        if not self._has_explicit_issues:
            return None
        properties = set(self._issue_property_map.get(index, set()))
        for category in categories or []:
            properties.update(self._global_issue_property_map.get(category, set()))
        if properties:
            return properties
        return set()

    def _property_allowed(self, property_name: str, allowed_properties: Optional[set]) -> bool:
        """Return True when a property can be changed under the current issue filter."""
        return allowed_properties is None or property_name in allowed_properties

    def _has_category_issue(self, category: str) -> bool:
        """Return True when the checker reported at least one issue in a category."""
        if not self._has_explicit_issues:
            return True
        return bool(self.issues_by_category.get(category))

    def _category_allows_global_property(self, category: str, property_name: str) -> bool:
        """Return True when a category-level issue permits a document-level fix."""
        if not self._has_explicit_issues:
            return True
        return property_name in self._global_issue_property_map.get(category, set())
    
    def fix_all(self) -> Tuple[Document, List[ChangeRecord]]:
        """
        Apply all formatting fixes to the document
        
        Returns:
            Tuple of (fixed document, list of changes made)
        """
        if not self.document:
            raise ValueError("No document loaded. Call load_manuscript() first.")
        
        self.changes = []

        if self._has_category_issue("layout"):
            self._fix_layout()
        
        # Fix margins first only when margin issues were detected.
        if self._has_category_issue("margins"):
            self._fix_margins()
        if self._category_allows_global_property("other", "manual_tabs"):
            self._fix_page_header_layout()
        
        # Fix paragraphs based on their classification
        for i, para in enumerate(self.document.paragraphs):
            classification = self._classification_map.get(i)
            
            if not classification:
                continue
            if (
                not classification.should_fix
                and classification.paragraph_type != ParagraphType.AUTHOR_INFO
            ):
                continue
            
            para_type = classification.paragraph_type
            
            if para_type == ParagraphType.JOURNAL_HEADER and self._is_journal_title_header(classification.text):
                if not self._should_fix_category(i, ["journal_header"]):
                    continue
                self._fix_journal_header(para, i)
            elif para_type == ParagraphType.PAPER_TITLE:
                if self._should_fix_category(i, ["title"]):
                    self._fix_title(para, i)
            elif para_type == ParagraphType.AUTHOR_INFO:
                if self._should_fix_category(i, ["author_info"]):
                    self._fix_author_info(para, i)
            elif para_type == ParagraphType.SECTION_HEADING:
                if self._should_fix_category(i, ["headings"]):
                    self._fix_heading(para, i)
            elif para_type == ParagraphType.BODY:
                if self._should_fix_category(i, ["body_text", "line_spacing"]):
                    self._fix_body_text(para, i)
            elif para_type == ParagraphType.ABSTRACT_CONTENT:
                if self._should_fix_category(i, ["body_text", "line_spacing"]):
                    self._fix_abstract(para, i)
            elif para_type == ParagraphType.KEYWORDS_CONTENT:
                if self._should_fix_category(i, ["body_text"]):
                    self._fix_keywords(para, i)
            elif para_type == ParagraphType.CAPTION:
                if self._should_fix_category(i, ["figures", "tables"]):
                    self._fix_caption(para, i)
            elif para_type == ParagraphType.REFERENCE:
                if self._should_fix_category(i, ["references"]):
                    self._fix_reference(para, i)

        if self._has_category_issue("references"):
            self._fix_sdt_references()
        
        return self.document, self.changes

    def _is_journal_title_header(self, text: str) -> bool:
        """Return True for journal title lines, excluding volume and ISSN metadata."""
        text_lower = text.lower()
        return "journal of informatics" in text_lower or text_lower.strip() == "web engineering"
    
    def _fix_margins(self):
        """Fix page margins with structured change records."""
        expected_margins = self.rules.get("margins", {})

        for section in self.document.sections:
            margin_specs = [
                ("left_margin", "left", "Left Margin"),
                ("right_margin", "right", "Right Margin"),
                ("top_margin", "top", "Top Margin"),
                ("bottom_margin", "bottom", "Bottom Margin"),
            ]

            for attr_name, rule_name, label in margin_specs:
                current_margin = getattr(section, attr_name)
                if not current_margin:
                    continue

                current_value = current_margin.inches
                expected_value = expected_margins.get(rule_name, 1.0)
                if abs(current_value - expected_value) <= 0.05:
                    continue

                setattr(section, attr_name, Inches(expected_value))
                self._add_change_record(
                    paragraph_index=-1,
                    location=label,
                    change_type="margins",
                    property_name=rule_name,
                    current_value=f"{current_value:.2f} in",
                    target_value=f"{expected_value:.2f} in",
                    text_preview="Document margins adjusted",
                    paragraph_type="document",
                    evidence="Detected margin issue"
                )

    def _fix_layout(self):
        """Fix page size, orientation, and column count."""
        rules = self.rules.get("layout", {})
        page_sizes = {
            "Letter": (8.5, 11.0),
            "A4": (8.27, 11.69),
        }
        expected_page_size = rules.get("page_size")
        expected_orientation = rules.get("orientation")
        expected_columns = rules.get("columns")

        for section_index, section in enumerate(self.document.sections):
            if expected_page_size in page_sizes:
                width, height = page_sizes[expected_page_size]
                if expected_orientation == "LANDSCAPE":
                    width, height = height, width
                section.page_width = Inches(width)
                section.page_height = Inches(height)

            if expected_orientation:
                section.orientation = (
                    WD_ORIENT.LANDSCAPE
                    if expected_orientation == "LANDSCAPE"
                    else WD_ORIENT.PORTRAIT
                )

            if expected_columns is not None:
                cols = section._sectPr.find(qn("w:cols"))
                if cols is None:
                    cols = OxmlElement("w:cols")
                    section._sectPr.append(cols)
                cols.set(qn("w:num"), str(int(expected_columns)))

        self._add_change_record(
            paragraph_index=-1,
            location="Page Layout",
            change_type="layout",
            text_preview="",
            property_name="layout",
            current_value="Non-matching layout",
            target_value=f"{expected_page_size}, {expected_orientation}, {expected_columns} column(s)",
            paragraph_type="document",
            evidence="Detected page layout did not match target rules",
        )

    def _fix_page_header_layout(self):
        """Normalize page header tab spacing to reduce Word wrapping."""
        for section_index, section in enumerate(self.document.sections):
            headers = [
                ("Page Header", section.header),
                ("First Page Header", section.first_page_header),
                ("Even Page Header", section.even_page_header),
            ]
            usable_width = section.page_width - section.left_margin - section.right_margin

            for label, header in headers:
                for paragraph in header.paragraphs:
                    original_text = paragraph.text
                    if not self._has_unstable_header_tabs(original_text):
                        continue
                    normalized_text = self._normalize_tabbed_header_text(original_text)
                    if normalized_text == original_text:
                        continue

                    self._replace_paragraph_text_preserving_first_run(paragraph, normalized_text)
                    try:
                        tab_stops = paragraph.paragraph_format.tab_stops
                        tab_stops.clear_all()
                        tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)
                    except Exception:
                        pass

                    self._add_change_record(
                        paragraph_index=-1,
                        location=f"{label} (Section {section_index + 1})",
                        change_type="page_header",
                        property_name="manual_tabs",
                        current_value=truncate_text(original_text, 80),
                        target_value=truncate_text(normalized_text, 80),
                        text_preview=truncate_text(normalized_text, 80),
                        paragraph_type="document_header",
                        evidence="Detected page header manual tabs/spaces that may wrap in Word",
                    )

    def _has_unstable_header_tabs(self, text: str) -> bool:
        """Return True for tab spacing patterns that are likely to wrap in Word."""
        return "\t\t" in text or bool(re.search(r"\t\s{2,}", text))

    def _normalize_tabbed_header_text(self, text: str) -> str:
        """Collapse unstable tab runs into one left/right tab-separated header line."""
        if "\t" not in text:
            return text
        parts = [re.sub(r"\s+", " ", part).strip() for part in text.split("\t") if part.strip()]
        if len(parts) < 2:
            return text
        return f"{parts[0]}\t{parts[-1]}"

    def _trim_paragraph_edge_whitespace(self, paragraph):
        """Trim paragraph edge spaces without replacing runs and losing formatting."""
        text = paragraph.text
        leading_count = len(text) - len(text.lstrip(" \t"))
        trailing_count = len(text) - len(text.rstrip(" \t"))

        remaining = leading_count
        for run in paragraph.runs:
            if remaining <= 0:
                break
            run_text = run.text
            if len(run_text) <= remaining:
                run.text = ""
                remaining -= len(run_text)
            else:
                run.text = run_text[remaining:]
                remaining = 0

        remaining = trailing_count
        for run in reversed(paragraph.runs):
            if remaining <= 0:
                break
            run_text = run.text
            if len(run_text) <= remaining:
                run.text = ""
                remaining -= len(run_text)
            else:
                run.text = run_text[:-remaining]
                remaining = 0

    def _replace_paragraph_text_preserving_first_run(self, paragraph, text: str):
        """Replace paragraph text while preserving the first visible run formatting."""
        first_index = next(
            (index for index, run in enumerate(paragraph.runs) if run.text),
            None,
        )
        if first_index is None:
            paragraph.add_run(text)
            return

        first_run = paragraph.runs[first_index]
        self._set_run_text_preserving_drawings(first_run, text)
        for index, run in enumerate(paragraph.runs):
            if index != first_index:
                self._clear_run_text_preserving_drawings(run)

    def _run_has_drawing_or_pict(self, run) -> bool:
        """Return True when a run contains a Word drawing/picture object."""
        xml = run._r.xml
        return "<w:drawing" in xml or "<w:pict" in xml or "AlternateContent" in xml

    def _clear_run_text_preserving_drawings(self, run):
        """Clear text, tabs, and line breaks while preserving drawing objects."""
        if not self._run_has_drawing_or_pict(run):
            run.text = ""
            return

        removable_tags = {
            qn("w:t"),
            qn("w:tab"),
            qn("w:br"),
            qn("w:cr"),
            qn("w:noBreakHyphen"),
            qn("w:softHyphen"),
        }
        for child in list(run._r):
            if child.tag in removable_tags:
                run._r.remove(child)

    def _set_run_text_preserving_drawings(self, run, text: str):
        """Set run text without deleting embedded header lines or images."""
        if not self._run_has_drawing_or_pict(run):
            run.text = text
            return

        self._clear_run_text_preserving_drawings(run)
        text_node = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = text
        insert_at = 1 if run._r.rPr is not None else 0
        run._r.insert(insert_at, text_node)

    def _add_change_record(
        self,
        paragraph_index: int,
        location: str,
        change_type: str,
        property_name: str,
        current_value: str,
        target_value: str,
        text_preview: str,
        paragraph_type: str,
        evidence: str,
    ):
        """Append one structured change record."""
        self.changes.append(ChangeRecord(
            paragraph_index=paragraph_index,
            location=location,
            change_type=change_type,
            before=f"{property_name}: {current_value}",
            after=f"{property_name}: {target_value}",
            text_preview=text_preview,
            property_name=property_name,
            current_value=current_value,
            target_value=target_value,
            paragraph_type=paragraph_type,
            evidence=evidence,
        ))

    def _add_property_changes(
        self,
        paragraph_index: int,
        location: str,
        change_type: str,
        details: List[Dict[str, str]],
        text_preview: str,
        paragraph_type: str,
    ):
        """Add de-duplicated property changes for one paragraph."""
        seen = set()
        for detail in details:
            key = (
                detail.get("property_name", "formatting"),
                detail.get("current_value", ""),
                detail.get("target_value", ""),
            )
            if key in seen:
                continue
            seen.add(key)
            self._add_change_record(
                paragraph_index=paragraph_index,
                location=location,
                change_type=change_type,
                property_name=detail.get("property_name", "formatting"),
                current_value=detail.get("current_value", ""),
                target_value=detail.get("target_value", ""),
                text_preview=text_preview,
                paragraph_type=paragraph_type,
                evidence=detail.get("evidence", "Detected formatting difference"),
            )

    def _fix_journal_header(self, paragraph, index: int):
        """Fix journal title/header formatting without using paper title rules."""
        header_rules = self.rules.get("journal_header", {})
        changes = []
        current_alignment = get_paragraph_alignment(paragraph)
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["journal_header"],
            fallback=["font_name", "font_size", "bold", "alignment", "manual_tabs"],
        )

        expected_font = header_rules.get("font_name", "Palatino Linotype")
        expected_size = header_rules.get("font_size", 24)
        expected_bold = header_rules.get("bold", True)
        expected_alignment = header_rules.get("alignment", "CENTER")

        original_text = paragraph.text
        cleaned_text = original_text.strip(" \t")
        if (
            self._property_allowed("manual_tabs", allowed_properties)
            and cleaned_text
            and cleaned_text != original_text
        ):
            self._trim_paragraph_edge_whitespace(paragraph)
            changes.append({
                "property_name": "manual_tabs",
                "current_value": "Manual tabs/spaces",
                "target_value": "No leading/trailing manual tabs/spaces",
                "evidence": "Journal header contained manual indentation that can shift layout",
            })

        if self._property_allowed("alignment", allowed_properties) and current_alignment != expected_alignment:
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 1)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Journal header alignment did not match target rule",
            })

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    expected_strike=False,
                    allowed_properties=allowed_properties,
                ))

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Journal Header",
                change_type="journal_header",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 50),
                paragraph_type=ParagraphType.JOURNAL_HEADER.value,
            )

    def _fix_run_formatting(self, run, expected_font: str, expected_size: float,
                           expected_bold: Optional[bool] = None,
                           expected_italic: Optional[bool] = None,
                           expected_strike: Optional[bool] = None,
                           allowed_properties: Optional[set] = None) -> List[Dict[str, str]]:
        """Fix a run and return structured property changes."""
        changes = []
        font = run.font

        preserve_italic = font.italic
        preserve_underline = font.underline
        preserve_subscript = font.subscript
        preserve_superscript = font.superscript
        preserve_strike = font.strike

        current_font = font.name
        if expected_font is not None and self._property_allowed("font_name", allowed_properties):
            if current_font is None or not is_font_equivalent(current_font, expected_font):
                font.name = expected_font
                changes.append({
                    "property_name": "font_name",
                    "current_value": current_font or "(inherited)",
                    "target_value": expected_font,
                    "evidence": "Run font did not match target rule",
                })

        current_size = font.size.pt if font.size else None
        if expected_size is not None and self._property_allowed("font_size", allowed_properties):
            if current_size is None or abs(current_size - expected_size) > 0.5:
                font.size = Pt(expected_size)
                changes.append({
                    "property_name": "font_size",
                    "current_value": f"{current_size} pt" if current_size is not None else "(inherited)",
                    "target_value": f"{expected_size} pt",
                    "evidence": "Run font size did not match target rule",
                })

        if expected_bold is not None and self._property_allowed("bold", allowed_properties):
            current_bold = font.bold
            current_bold_value = bool(current_bold) if current_bold is not None else False
            if current_bold_value != expected_bold:
                font.bold = expected_bold
                changes.append({
                    "property_name": "bold",
                    "current_value": "Bold" if current_bold_value else "Not Bold",
                    "target_value": "Bold" if expected_bold else "Not Bold",
                    "evidence": "Run bold setting did not match target rule",
                })

        if expected_italic is not None and self._property_allowed("italic", allowed_properties):
            current_italic = font.italic
            current_italic_value = bool(current_italic) if current_italic is not None else False
            if current_italic_value != expected_italic:
                font.italic = expected_italic
                changes.append({
                    "property_name": "italic",
                    "current_value": "Italic" if current_italic_value else "Not Italic",
                    "target_value": "Italic" if expected_italic else "Not Italic",
                    "evidence": "Run italic setting did not match target rule",
                })
        else:
            font.italic = preserve_italic

        if expected_strike is not None and self._property_allowed("strike", allowed_properties):
            if font.strike != expected_strike:
                font.strike = expected_strike
        else:
            font.strike = preserve_strike

        font.underline = preserve_underline
        font.subscript = preserve_subscript
        font.superscript = preserve_superscript

        return changes

    def _fix_title(self, paragraph, index: int):
        """Fix paper title formatting with structured change records."""
        title_rules = self.rules.get("title", {})
        changes = []
        current_alignment = get_paragraph_alignment(paragraph)
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["title"],
            fallback=["font_name", "font_size", "bold", "italic", "alignment"],
        )

        expected_font = title_rules.get("font_name", "Times New Roman")
        expected_size = title_rules.get("font_size", 24)
        expected_bold = title_rules.get("bold", None)
        expected_italic = title_rules.get("italic", None)
        expected_alignment = title_rules.get("alignment", "CENTER")

        if self._property_allowed("alignment", allowed_properties) and current_alignment != expected_alignment:
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 1)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Paragraph alignment did not match target rule",
            })

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    expected_italic,
                    expected_strike=False,
                    allowed_properties=allowed_properties,
                ))

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Paper Title",
                change_type="title",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 50),
                paragraph_type=ParagraphType.PAPER_TITLE.value,
            )

    def _fix_author_info(self, paragraph, index: int):
        """Fix front-matter author information without changing its text."""
        role = classify_author_info_role(get_paragraph_text(paragraph), self.rules)
        rules = self.rules.get(role, self.rules.get("author", {}))
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["author_info"],
            fallback=["font_name", "font_size", "bold", "italic", "alignment"],
        )
        changes = []
        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    rules.get("font_name"),
                    rules.get("font_size"),
                    rules.get("bold"),
                    rules.get("italic"),
                    allowed_properties=allowed_properties,
                ))

        expected_alignment = rules.get("alignment")
        current_alignment = get_paragraph_alignment(paragraph)
        if (
            expected_alignment
            and self._property_allowed("alignment", allowed_properties)
            and current_alignment != expected_alignment
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 1)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Author-information alignment did not match target rule",
            })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location=role.replace("_", " ").title(),
                change_type="author_info",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 50),
                paragraph_type=ParagraphType.AUTHOR_INFO.value,
            )

    def _fix_heading(self, paragraph, index: int):
        """Fix section heading formatting with structured change records."""
        heading_rules = self._heading_rules_for_text(get_paragraph_text(paragraph))
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["headings"],
            fallback=[
                "font_name",
                "font_size",
                "bold",
                "italic",
                "capitalization",
                "number_font_name",
                "number_font_size",
                "number_bold",
            ],
        )

        expected_font = heading_rules.get("font_name", "Times New Roman")
        expected_size = heading_rules.get("font_size", 10)
        expected_bold = heading_rules.get("bold", None)
        expected_italic = heading_rules.get("italic", None)
        expected_all_caps = heading_rules.get("all_caps")

        changes.extend(self._fix_numbering_formatting(
            paragraph,
            expected_font,
            expected_size,
            expected_bold,
            allowed_properties,
        ))

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    expected_italic,
                    expected_strike=False,
                    allowed_properties=allowed_properties,
                ))

        current_text = get_paragraph_text(paragraph)
        if (
            expected_all_caps
            and self._property_allowed("capitalization", allowed_properties)
            and not self._heading_text_is_all_caps(current_text)
        ):
            target_text = current_text.upper()
            self._replace_paragraph_text_preserving_first_run(paragraph, target_text)
            changes.append({
                "property_name": "capitalization",
                "current_value": current_text,
                "target_value": target_text,
                "evidence": "Heading capitalization did not match target rule",
            })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Section Heading",
                change_type="heading",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.SECTION_HEADING.value,
            )

    def _heading_rules_for_text(self, text: str) -> Dict[str, Any]:
        """Return the correct heading rule for main headings or subheadings."""
        stripped = re.sub(r"\s+", " ", text.strip())
        if re.match(r"^\d+\.\d+", stripped):
            return self.rules.get("subheading", self.rules.get("heading", {}))
        return self.rules.get("heading", {})

    @staticmethod
    def _heading_text_is_all_caps(text: str) -> bool:
        """Return True when heading words are uppercase after removing numbering."""
        stripped = re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", text or "").strip()
        letters = [char for char in stripped if char.isalpha()]
        return bool(letters) and all(char.isupper() for char in letters)

    def _get_numbering_level(self, paragraph):
        """Return the numbering level XML element for a numbered paragraph."""
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.numPr is None:
            return None

        num_id_element = p_pr.numPr.numId
        ilvl_element = p_pr.numPr.ilvl
        if num_id_element is None:
            return None

        num_id = num_id_element.val
        ilvl = str(ilvl_element.val if ilvl_element is not None else 0)
        numbering = self.document.part.numbering_part.element

        abstract_num_id = None
        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) == str(num_id):
                abstract_node = num.find(qn("w:abstractNumId"))
                if abstract_node is not None:
                    abstract_num_id = abstract_node.get(qn("w:val"))
                break

        if abstract_num_id is None:
            return None

        for abstract_num in numbering.findall(qn("w:abstractNum")):
            if abstract_num.get(qn("w:abstractNumId")) != str(abstract_num_id):
                continue
            for level in abstract_num.findall(qn("w:lvl")):
                if level.get(qn("w:ilvl")) == ilvl:
                    return level

        return None

    def _get_or_create_child(self, parent, tag_name: str):
        """Return a child XML node, creating it when missing."""
        child = parent.find(qn(tag_name))
        if child is None:
            child = OxmlElement(tag_name)
            parent.append(child)
        return child

    def _numbering_bold_value(self, r_pr) -> Optional[bool]:
        """Read numbering bold value from a numbering rPr node."""
        bold = r_pr.find(qn("w:b")) if r_pr is not None else None
        if bold is None:
            return None
        value = bold.get(qn("w:val"))
        return value not in {"0", "false", "False", "off"}

    def _fix_numbering_formatting(
        self,
        paragraph,
        expected_font: str,
        expected_size: float,
        expected_bold: Optional[bool],
        allowed_properties: Optional[set] = None,
    ) -> List[Dict[str, str]]:
        """Fix Word list numbering formatting for numbered headings."""
        level = self._get_numbering_level(paragraph)
        if level is None:
            return []

        changes = []
        r_pr = self._get_or_create_child(level, "w:rPr")

        if expected_font and self._property_allowed("number_font_name", allowed_properties):
            r_fonts = self._get_or_create_child(r_pr, "w:rFonts")
            current_font = r_fonts.get(qn("w:ascii")) or r_fonts.get(qn("w:hAnsi"))
            if current_font and not is_font_equivalent(current_font, expected_font):
                changes.append({
                    "property_name": "number_font_name",
                    "current_value": current_font,
                    "target_value": expected_font,
                    "evidence": "Heading number font did not match target rule",
                })
            r_fonts.set(qn("w:ascii"), expected_font)
            r_fonts.set(qn("w:hAnsi"), expected_font)

        if expected_size and self._property_allowed("number_font_size", allowed_properties):
            size_node = self._get_or_create_child(r_pr, "w:sz")
            current_size = None
            if size_node.get(qn("w:val")):
                current_size = int(size_node.get(qn("w:val"))) / 2
            if current_size is not None and abs(current_size - expected_size) > 0.5:
                changes.append({
                    "property_name": "number_font_size",
                    "current_value": f"{current_size:.1f} pt",
                    "target_value": f"{expected_size} pt",
                    "evidence": "Heading number font size did not match target rule",
                })
            size_node.set(qn("w:val"), str(int(float(expected_size) * 2)))

        if expected_bold is not None and self._property_allowed("number_bold", allowed_properties):
            current_bold = self._numbering_bold_value(r_pr)
            if current_bold != expected_bold:
                changes.append({
                    "property_name": "number_bold",
                    "current_value": "Bold" if current_bold else "Not Bold",
                    "target_value": "Bold" if expected_bold else "Not Bold",
                    "evidence": "Heading number bold setting did not match target rule",
                })
            bold_node = self._get_or_create_child(r_pr, "w:b")
            bold_node.set(qn("w:val"), "1" if expected_bold else "0")

        return changes

    def _fix_body_text(self, paragraph, index: int):
        """Fix body text formatting with structured change records."""
        body_rules = self.rules.get("body", {})
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["body_text", "line_spacing"],
            fallback=[
                "font_name",
                "font_size",
                "bold",
                "line_spacing",
                "alignment",
                "space_after",
                "manual_line_breaks",
            ],
        )

        expected_font = body_rules.get("font_name", "Times New Roman")
        expected_size = body_rules.get("font_size", 12)
        expected_bold = body_rules.get("bold", None)
        expected_alignment = body_rules.get("alignment")
        expected_space_after = body_rules.get("space_after")

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    allowed_properties=allowed_properties,
                ))

        if (
            self._property_allowed("manual_line_breaks", allowed_properties)
            and paragraph_has_manual_line_breaks(paragraph)
        ):
            replacements = replace_manual_line_breaks_with_spaces(paragraph)
            if replacements:
                changes.append({
                    "property_name": "manual_line_breaks",
                    "current_value": f"{replacements} manual line break(s)",
                    "target_value": "Normal paragraph wrapping",
                    "evidence": "Body paragraph contained manual line breaks that can stretch justified text",
                })

        expected_spacing = body_rules.get("line_spacing")
        if expected_spacing and self._property_allowed("line_spacing", allowed_properties):
            paragraph_format = paragraph.paragraph_format
            if paragraph_format.line_spacing != expected_spacing:
                current_spacing = paragraph_format.line_spacing
                paragraph_format.line_spacing = expected_spacing
                changes.append({
                    "property_name": "line_spacing",
                    "current_value": str(current_spacing) if current_spacing is not None else "(inherited)",
                    "target_value": str(expected_spacing),
                    "evidence": "Paragraph line spacing did not match target rule",
                })

        if (
            expected_space_after is not None
            and self._property_allowed("space_after", allowed_properties)
        ):
            current_space_after = get_space_after_pt(paragraph)
            if (
                current_space_after is None
                or abs(float(current_space_after) - float(expected_space_after)) > 0.5
            ):
                paragraph.paragraph_format.space_after = Pt(float(expected_space_after))
                changes.append({
                    "property_name": "space_after",
                    "current_value": (
                        f"{current_space_after}pt"
                        if current_space_after is not None
                        else "(inherited)"
                    ),
                    "target_value": f"{expected_space_after}pt",
                    "evidence": "Body paragraph spacing after did not match target rule",
                })

        current_alignment = get_paragraph_alignment(paragraph)
        if (
            expected_alignment
            and self._property_allowed("alignment", allowed_properties)
            and current_alignment != expected_alignment
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 3)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Body text alignment did not match target rule",
            })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location=f"Body Text (Para {index + 1})",
                change_type="body",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.BODY.value,
            )

    def _fix_abstract(self, paragraph, index: int):
        """Fix abstract formatting with structured change records."""
        abstract_rules = self.rules.get("abstract", {})
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["body_text", "line_spacing"],
            fallback=["font_name", "font_size", "bold", "line_spacing", "alignment", "space_after"],
        )

        expected_font = abstract_rules.get("font_name", "Times New Roman")
        expected_size = abstract_rules.get("font_size", 9)
        expected_bold = abstract_rules.get("bold", None)
        expected_alignment = abstract_rules.get("alignment")
        expected_space_after = abstract_rules.get(
            "space_after",
            self.rules.get("body", {}).get("space_after"),
        )

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    allowed_properties=allowed_properties,
                ))

        expected_spacing = abstract_rules.get("line_spacing", self.rules.get("body", {}).get("line_spacing"))
        if expected_spacing and self._property_allowed("line_spacing", allowed_properties):
            paragraph_format = paragraph.paragraph_format
            if paragraph_format.line_spacing != expected_spacing:
                current_spacing = paragraph_format.line_spacing
                paragraph_format.line_spacing = expected_spacing
                changes.append({
                    "property_name": "line_spacing",
                    "current_value": str(current_spacing) if current_spacing is not None else "(inherited)",
                    "target_value": str(expected_spacing),
                    "evidence": "Abstract line spacing did not match target rule",
                })

        if (
            expected_space_after is not None
            and self._property_allowed("space_after", allowed_properties)
        ):
            current_space_after = get_space_after_pt(paragraph)
            if (
                current_space_after is None
                or abs(float(current_space_after) - float(expected_space_after)) > 0.5
            ):
                paragraph.paragraph_format.space_after = Pt(float(expected_space_after))
                changes.append({
                    "property_name": "space_after",
                    "current_value": (
                        f"{current_space_after}pt"
                        if current_space_after is not None
                        else "(inherited)"
                    ),
                    "target_value": f"{expected_space_after}pt",
                    "evidence": "Abstract paragraph spacing after did not match target rule",
                })

        current_alignment = get_paragraph_alignment(paragraph)
        if (
            expected_alignment
            and self._property_allowed("alignment", allowed_properties)
            and current_alignment != expected_alignment
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 3)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Abstract alignment did not match target rule",
            })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Abstract",
                change_type="abstract",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.ABSTRACT_CONTENT.value,
            )

    def _fix_keywords(self, paragraph, index: int):
        """Fix keyword formatting with structured change records."""
        keywords_rules = self.rules.get("keywords", self.rules.get("abstract", {}))
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["body_text"],
            fallback=["font_name", "font_size", "bold", "italic", "capitalization"],
        )

        expected_font = keywords_rules.get("font_name", "Times New Roman")
        expected_size = keywords_rules.get("font_size", 9)
        expected_bold = keywords_rules.get("bold", None)
        expected_italic = keywords_rules.get("italic", None)
        capitalize_first_letter = keywords_rules.get("capitalize_first_letter")

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    expected_italic,
                    allowed_properties=allowed_properties,
                ))

        current_text = get_paragraph_text(paragraph)
        if (
            capitalize_first_letter
            and self._property_allowed("capitalization", allowed_properties)
        ):
            target_text = self._capitalize_keyword_text(current_text)
            if target_text != current_text:
                self._replace_paragraph_text_preserving_first_run(paragraph, target_text)
                changes.append({
                    "property_name": "capitalization",
                    "current_value": current_text,
                    "target_value": target_text,
                    "evidence": "Keyword capitalization did not match target rule",
                })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Keywords",
                change_type="keywords",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.KEYWORDS_CONTENT.value,
            )

    @staticmethod
    def _capitalize_keyword_text(text: str) -> str:
        """Capitalize the first letter of each existing keyword item."""
        match = re.match(
            r"^(\s*(?:keywords?|key\s+words?)\s*(?:-|\u2013|\u2014|:)?\s*)(.*)$",
            text or "",
            flags=re.IGNORECASE,
        )
        if not match:
            return text

        prefix, keyword_text = match.groups()
        parts = re.split(r"([,;])", keyword_text)
        for index in range(0, len(parts), 2):
            parts[index] = re.sub(
                r"^(\s*)([A-Za-z])",
                lambda item: item.group(1) + item.group(2).upper(),
                parts[index],
                count=1,
            )
        return prefix + "".join(parts)

    def _fix_caption(self, paragraph, index: int):
        """Fix caption formatting with structured change records."""
        caption_rules = self.rules.get("caption", {})
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["figures", "tables"],
            fallback=["font_name", "font_size", "italic", "space_after", "capitalization"],
        )

        expected_font = caption_rules.get("font_name", "Times New Roman")
        expected_size = caption_rules.get("font_size", 10)
        expected_italic = caption_rules.get("italic", None)
        expected_space_after = caption_rules.get("space_after")
        title_case_required = bool(caption_rules.get("title_case"))

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    None,
                    expected_italic,
                    allowed_properties=allowed_properties,
                ))

        if (
            expected_space_after is not None
            and self._property_allowed("space_after", allowed_properties)
        ):
            current_space_after = get_space_after_pt(paragraph)
            if (
                current_space_after is None
                or abs(float(current_space_after) - float(expected_space_after)) > 0.5
            ):
                paragraph.paragraph_format.space_after = Pt(float(expected_space_after))
                changes.append({
                    "property_name": "space_after",
                    "current_value": (
                        f"{current_space_after}pt"
                        if current_space_after is not None
                        else "(inherited)"
                    ),
                    "target_value": f"{expected_space_after}pt",
                    "evidence": "Caption paragraph spacing after did not match target rule",
                })

        current_text = get_paragraph_text(paragraph)
        if (
            title_case_required
            and self._property_allowed("capitalization", allowed_properties)
        ):
            target_text = to_journal_caption_title_case(current_text)
            if target_text != current_text:
                self._replace_paragraph_text_preserving_first_run(paragraph, target_text)
                changes.append({
                    "property_name": "capitalization",
                    "current_value": current_text,
                    "target_value": target_text,
                    "evidence": "Caption capitalization did not match target rule",
                })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location="Caption",
                change_type="caption",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.CAPTION.value,
            )

    def _fix_reference(self, paragraph, index: int, location: str = "Reference"):
        """Fix reference formatting with structured change records."""
        reference_rules = self.rules.get("reference", {})
        changes = []
        allowed_properties = self._allowed_properties_for(
            index,
            categories=["references"],
            fallback=[
                "font_name",
                "font_size",
                "bold",
                "alignment",
                "line_spacing",
                "space_after",
                "left_indent",
                "hanging_indent",
            ],
        )

        expected_font = reference_rules.get("font_name", "Times New Roman")
        expected_size = reference_rules.get("font_size", 9)
        expected_bold = reference_rules.get("bold", None)
        expected_alignment = reference_rules.get("alignment")
        expected_line_spacing = reference_rules.get("line_spacing")
        expected_space_after = reference_rules.get("space_after")
        has_left_indent_rule = "left_indent" in reference_rules
        expected_left_indent = reference_rules.get("left_indent")
        expected_hanging_indent = reference_rules.get("hanging_indent")

        for run in paragraph.runs:
            if run.text.strip():
                changes.extend(self._fix_run_formatting(
                    run,
                    expected_font,
                    expected_size,
                    expected_bold,
                    allowed_properties=allowed_properties,
                ))

        current_alignment = get_paragraph_alignment(paragraph)
        if (
            expected_alignment
            and self._property_allowed("alignment", allowed_properties)
            and current_alignment != expected_alignment
        ):
            paragraph.alignment = WD_ALIGN_PARAGRAPH(
                ALIGNMENT_REVERSE_MAP.get(expected_alignment, 3)
            )
            changes.append({
                "property_name": "alignment",
                "current_value": current_alignment,
                "target_value": expected_alignment,
                "evidence": "Reference alignment did not match target rule",
            })

        if (
            expected_line_spacing is not None
            and self._property_allowed("line_spacing", allowed_properties)
        ):
            current_line_spacing = paragraph.paragraph_format.line_spacing
            effective_line_spacing = 1.0 if current_line_spacing is None else float(current_line_spacing)
            if abs(effective_line_spacing - float(expected_line_spacing)) > 0.05:
                paragraph.paragraph_format.line_spacing = expected_line_spacing
                changes.append({
                    "property_name": "line_spacing",
                    "current_value": str(current_line_spacing) if current_line_spacing is not None else "(inherited)",
                    "target_value": str(expected_line_spacing),
                    "evidence": "Reference line spacing did not match target rule",
                })

        if (
            expected_space_after is not None
            and self._property_allowed("space_after", allowed_properties)
        ):
            current_space_after = get_space_after_pt(paragraph)
            if (
                current_space_after is None
                or abs(float(current_space_after) - float(expected_space_after)) > 0.5
            ):
                paragraph.paragraph_format.space_after = Pt(float(expected_space_after))
                changes.append({
                    "property_name": "space_after",
                    "current_value": (
                        f"{current_space_after}pt"
                        if current_space_after is not None
                        else "(inherited)"
                    ),
                    "target_value": f"{expected_space_after}pt",
                    "evidence": "Reference paragraph spacing after did not match target rule",
                })

        if has_left_indent_rule and self._property_allowed("left_indent", allowed_properties):
            current_left_indent = get_direct_left_indent_inches(paragraph)
            if expected_left_indent is None:
                if current_left_indent is not None:
                    paragraph.paragraph_format.left_indent = None
                    changes.append({
                        "property_name": "left_indent",
                        "current_value": f"{float(current_left_indent):.2f}in",
                        "target_value": "No explicit left indent",
                        "evidence": "Reference left indent did not match target rule",
                    })
            else:
                effective_left_indent = 0.0 if current_left_indent is None else float(current_left_indent)
                if abs(effective_left_indent - float(expected_left_indent)) > 0.03:
                    paragraph.paragraph_format.left_indent = Inches(float(expected_left_indent))
                    changes.append({
                        "property_name": "left_indent",
                        "current_value": f"{effective_left_indent:.2f}in",
                        "target_value": f"{expected_left_indent}in",
                        "evidence": "Reference left indent did not match target rule",
                    })

        if (
            expected_hanging_indent is not None
            and self._property_allowed("hanging_indent", allowed_properties)
        ):
            current_hanging_indent = get_hanging_indent_inches(paragraph)
            if (
                current_hanging_indent is None
                or abs(float(current_hanging_indent) - float(expected_hanging_indent)) > 0.03
            ):
                paragraph.paragraph_format.first_line_indent = Inches(-float(expected_hanging_indent))
                changes.append({
                    "property_name": "hanging_indent",
                    "current_value": (
                        f"{current_hanging_indent:.2f}in"
                        if current_hanging_indent is not None
                        else "No hanging indent"
                    ),
                    "target_value": f"{expected_hanging_indent}in",
                    "evidence": "Reference hanging indent did not match target rule",
                })

        if changes:
            self._add_property_changes(
                paragraph_index=index,
                location=location,
                change_type="reference",
                details=changes,
                text_preview=truncate_text(get_paragraph_text(paragraph), 40),
                paragraph_type=ParagraphType.REFERENCE.value,
            )

    def _fix_sdt_references(self):
        """Fix reference paragraphs stored inside Word content controls."""
        for index, paragraph in enumerate(get_sdt_reference_paragraphs(self.document), start=1):
            self._fix_reference(
                paragraph,
                -1,
                location=f"Reference Content Control {index}",
            )

    def get_fixed_document_bytes(self) -> bytes:
        """Get the fixed document as bytes for download"""
        if not self.document:
            raise ValueError("No document to export")
        
        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def get_changes_summary(self) -> Dict[str, Any]:
        """Get summary of all changes made"""
        summary = {
            "total_changes": len(self.changes),
            "changes_by_type": {},
            "changes_list": []
        }
        
        for change in self.changes:
            change_type = change.change_type
            summary["changes_by_type"][change_type] = \
                summary["changes_by_type"].get(change_type, 0) + 1
            
            summary["changes_list"].append({
                "location": change.location,
                "type": change.change_type,
                "before": change.before,
                "after": change.after,
                "property_name": change.property_name,
                "current_value": change.current_value,
                "target_value": change.target_value,
                "paragraph_type": change.paragraph_type,
                "evidence": change.evidence,
                "preview": change.text_preview
            })
        
        return summary
    
    def get_change_records(self) -> List[ChangeRecord]:
        """Get list of all change records"""
        return self.changes

    def _run_matches_change(self, run, change: ChangeRecord) -> bool:
        """Return True when a run likely contains the changed property."""
        property_name = change.property_name
        info = get_run_font_info(run)

        if property_name == "font_name":
            current_font = info.get("font_name")
            target_font = change.target_value
            return current_font is None or not is_font_equivalent(current_font, target_font)

        if property_name == "font_size":
            current_size = info.get("font_size")
            target_match = re.search(r"([\d.]+)", change.target_value)
            if not target_match:
                return True
            target_size = float(target_match.group(1))
            return current_size is None or abs(current_size - target_size) > 0.5

        if property_name == "bold":
            target_bold = change.target_value.lower() == "bold"
            return bool(info.get("bold")) != target_bold

        if property_name == "italic":
            target_italic = change.target_value.lower() == "italic"
            return bool(info.get("italic")) != target_italic

        return True

    def _highlight_page_header_change(self, document: Document, change: ChangeRecord) -> bool:
        """Highlight page header text for document-level header changes."""
        if change.change_type != "page_header":
            return False

        highlighted = False
        target_text = (change.target_value or "").strip()
        for section in document.sections:
            headers = [
                section.header,
                section.first_page_header,
                section.even_page_header,
            ]
            for header in headers:
                for paragraph in header.paragraphs:
                    if not paragraph.text.strip():
                        continue
                    paragraph_text = paragraph.text.strip()
                    matches_original = self._has_unstable_header_tabs(paragraph.text)
                    matches_fixed = target_text and paragraph_text == target_text
                    if not (matches_original or matches_fixed):
                        continue
                    for run in paragraph.runs:
                        if run.text.strip():
                            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                            highlighted = True
        return highlighted

    def get_highlighted_document_bytes(self) -> bytes:
        """
        Get the corrected document with highlighted changed locations.
        Changed runs are highlighted in yellow without adding summary content.
        
        Returns:
            Document bytes with yellow highlighting on corrected changed locations
        """
        source_document = self.document or self.original_document
        if not source_document:
            raise ValueError("No document to export")

        highlighted_document = deepcopy(source_document)

        for change in self.changes:
            index = change.paragraph_index
            if index < 0 or index >= len(highlighted_document.paragraphs):
                self._highlight_page_header_change(highlighted_document, change)
                continue

            paragraph = highlighted_document.paragraphs[index]
            matched_any_run = False
            for run in paragraph.runs:
                if not run.text.strip():
                    continue
                if self._run_matches_change(run, change):
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    matched_any_run = True

            if not matched_any_run:
                for run in paragraph.runs:
                    if run.text.strip():
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        buffer = BytesIO()
        highlighted_document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
