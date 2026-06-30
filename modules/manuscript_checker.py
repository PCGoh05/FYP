"""
Manuscript Checker Module
Checks manuscript formatting against template rules
"""

from docx.oxml.ns import qn
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
import re
from difflib import SequenceMatcher

from .utils import (
    load_document, get_paragraph_text, get_paragraph_alignment, get_paragraph_font_info,
    get_sdt_reference_paragraphs, get_margins, get_line_spacing,
    detect_reference_style, calculate_compliance_score, count_columns,
    truncate_text, is_font_equivalent, classify_author_info_role
)
from .paragraph_classifier import ParagraphClassifier, ParagraphType, ClassifiedParagraph
from config import REQUIRED_SECTIONS


@dataclass
class FormatIssue:
    """Represents a formatting issue found in the manuscript"""
    category: str
    location: str
    paragraph_index: int
    description: str
    current_value: str
    expected_value: str
    severity: str = "warning"  # "error", "warning", "info"
    text_preview: str = ""


@dataclass
class CheckResult:
    """Result of manuscript check"""
    is_compliant: bool
    compliance_score: float
    total_issues: int
    issues_by_category: Dict[str, List[FormatIssue]]
    classifications: List[ClassifiedParagraph]
    document_structure: Dict[str, Any]
    statistics: Dict[str, Any]


class ManuscriptChecker:
    """
    Checks manuscript formatting against extracted template rules
    Performs 10 category checks and calculates compliance score
    """
    
    CATEGORIES = [
        "margins",
        "layout",
        "journal_header",
        "title",
        "author_info",
        "body_text",
        "headings",
        "structure",
        "tables",
        "figures",
        "references",
        "line_spacing",
        "other"
    ]
    
    def __init__(self, rules: Dict[str, Any], llm_integration=None):
        """Initialize checker with template rules"""
        self.rules = rules
        self.llm = llm_integration
        self.document = None
        self.profile = rules.get("_profile", {})
        self.classifier = ParagraphClassifier(llm_integration, self.profile)
        self.issues: Dict[str, List[FormatIssue]] = {cat: [] for cat in self.CATEGORIES}
        self.classifications: List[ClassifiedParagraph] = []
        self.required_sections = self.profile.get("required_sections", REQUIRED_SECTIONS)
        self.required_declarations = self.profile.get("required_declarations", [])
    
    def load_manuscript(self, file_path_or_bytes):
        """Load the manuscript to check"""
        self.document = load_document(file_path_or_bytes)
        return self
    
    def check_all(self) -> CheckResult:
        """Perform all format checks"""
        if not self.document:
            raise ValueError("No manuscript loaded. Call load_manuscript() first.")
        
        # Reset issues
        self.issues = {cat: [] for cat in self.CATEGORIES}
        
        # Classify paragraphs first
        self.classifications = self.classifier.classify_document(self.document)
        
        # Run all checks
        self._check_margins()
        self._check_layout()
        self._check_journal_header()
        self._check_layout_stability()
        self._check_title()
        self._check_author_info()
        self._check_body_text()
        self._check_headings()
        structure = self._check_document_structure()
        self._check_tables()
        self._check_figures()
        self._check_equations()
        self._check_references()
        self._check_line_spacing()
        
        # Calculate compliance score
        score = calculate_compliance_score(self.issues)
        total_issues = sum(len(issues) for issues in self.issues.values())
        
        # Gather statistics
        stats = self._gather_statistics()
        
        return CheckResult(
            is_compliant=total_issues == 0,
            compliance_score=score,
            total_issues=total_issues,
            issues_by_category=self.issues,
            classifications=self.classifications,
            document_structure=structure,
            statistics=stats
        )
    
    def _add_issue(self, category: str, location: str, para_index: int,
                   description: str, current: str, expected: str,
                   severity: str = "warning", text_preview: str = ""):
        """Add a formatting issue"""
        issue = FormatIssue(
            category=category,
            location=location,
            paragraph_index=para_index,
            description=description,
            current_value=current,
            expected_value=expected,
            severity=severity,
            text_preview=text_preview
        )
        self.issues[category].append(issue)
    
    def _check_margins(self):
        """Check page margins against template"""
        current_margins = get_margins(self.document)
        expected_margins = self.rules.get("margins", {})
        
        margin_names = ["left", "right", "top", "bottom"]
        
        for margin in margin_names:
            current = current_margins.get(margin, 1.0)
            expected = expected_margins.get(margin, 1.0)
            
            # Allow small tolerance (0.05 inches)
            if abs(current - expected) > 0.05:
                self._add_issue(
                    category="margins",
                    location=f"{margin.capitalize()} Margin",
                    para_index=-1,
                    description=f"{margin.capitalize()} margin does not match template",
                    current=f"{current:.2f} inches",
                    expected=f"{expected:.2f} inches",
                    severity="error"
                )
    
    def _check_title(self):
        """Check paper title formatting"""
        title_rules = self.rules.get("title", {})
        found_title = False
        
        # Find the paper title
        for cp in self.classifications:
            if cp.paragraph_type == ParagraphType.PAPER_TITLE:
                found_title = True
                font_info = cp.font_info
                alignment = cp.alignment
                
                # Check font name (using font equivalence)
                expected_font = title_rules.get("font_name", "Times New Roman")
                current_font = font_info.get("font_name")
                if current_font and not is_font_equivalent(current_font, expected_font):
                    self._add_issue(
                        category="title",
                        location="Paper Title",
                        para_index=cp.index,
                        description="Title font does not match template",
                        current=current_font,
                        expected=expected_font,
                        severity="error",
                        text_preview=truncate_text(cp.text, 60)
                    )
                
                # Check font size (with 0.5pt tolerance)
                expected_size = title_rules.get("font_size", 24)
                current_size = font_info.get("font_size")
                if current_size and abs(current_size - expected_size) > 0.5:
                    self._add_issue(
                        category="title",
                        location="Paper Title",
                        para_index=cp.index,
                        description="Title font size does not match template",
                        current=f"{current_size}pt",
                        expected=f"{expected_size}pt",
                        severity="error",
                        text_preview=truncate_text(cp.text, 60)
                    )
                
                # Check bold
                expected_bold = title_rules.get("bold")
                current_bold = self._is_paragraph_mostly_bold(cp.index)
                if expected_bold is not None and current_bold != expected_bold:
                    self._add_issue(
                        category="title",
                        location="Paper Title",
                        para_index=cp.index,
                        description="Title bold formatting does not match template",
                        current="Bold" if current_bold else "Not Bold",
                        expected="Bold" if expected_bold else "Not Bold",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 60)
                    )

                expected_italic = title_rules.get("italic")
                current_italic = bool(font_info.get("italic"))
                if expected_italic is not None and current_italic != expected_italic:
                    self._add_issue(
                        category="title",
                        location="Paper Title",
                        para_index=cp.index,
                        description="Title italic formatting does not match template",
                        current="Italic" if current_italic else "Not Italic",
                        expected="Italic" if expected_italic else "Not Italic",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 60)
                    )

                # Check alignment
                expected_align = title_rules.get("alignment", "CENTER")
                if alignment != expected_align:
                    self._add_issue(
                        category="title",
                        location="Paper Title",
                        para_index=cp.index,
                        description="Title alignment does not match template",
                        current=alignment,
                        expected=expected_align,
                        severity="warning",
                        text_preview=truncate_text(cp.text, 60)
                    )
                
                break  # Only check the first title

        if not found_title:
            self._add_issue(
                category="title",
                location="Paper Title",
                para_index=-1,
                description="Paper title could not be identified",
                current="Not found",
                expected="One clear paper title near the beginning of the manuscript",
                severity="error"
            )

    def _check_journal_header(self):
        """Check journal header formatting separately from the paper title."""
        header_rules = self.rules.get("journal_header", {})
        expected_font = header_rules.get("font_name", "Palatino Linotype")
        expected_size = header_rules.get("font_size", 24)
        expected_bold = header_rules.get("bold", True)
        expected_align = header_rules.get("alignment", "CENTER")

        header_paragraphs = [
            cp for cp in self.classifications
            if (
                cp.paragraph_type == ParagraphType.JOURNAL_HEADER
                and cp.index <= 1
                and ("journal of" in cp.text.lower() or "web engineering" in cp.text.lower())
            )
        ]

        if not header_paragraphs:
            self._add_issue(
                category="journal_header",
                location="Journal Header",
                para_index=-1,
                description="Journal header could not be identified",
                current="Not found",
                expected="Journal title/header near the top of the manuscript",
                severity="warning"
            )
            return

        for cp in header_paragraphs:
            font_info = cp.font_info
            location = f"Journal Header (Para {cp.index + 1})"
            preview = truncate_text(cp.text, 60)
            raw_text = self.document.paragraphs[cp.index].text

            current_font = font_info.get("font_name")
            if current_font and not is_font_equivalent(current_font, expected_font):
                self._add_issue(
                    category="journal_header",
                    location=location,
                    para_index=cp.index,
                    description="Journal header font does not match template",
                    current=current_font,
                    expected=expected_font,
                    severity="warning",
                    text_preview=preview
                )

            current_size = font_info.get("font_size")
            if current_size and abs(current_size - expected_size) > 0.5:
                self._add_issue(
                    category="journal_header",
                    location=location,
                    para_index=cp.index,
                    description="Journal header font size does not match template",
                    current=f"{current_size}pt",
                    expected=f"{expected_size}pt",
                    severity="warning",
                    text_preview=preview
                )

            current_bold = bool(font_info.get("bold"))
            if expected_bold is not None and current_bold != expected_bold:
                self._add_issue(
                    category="journal_header",
                    location=location,
                    para_index=cp.index,
                    description="Journal header bold formatting does not match template",
                    current="Bold" if current_bold else "Not Bold",
                    expected="Bold" if expected_bold else "Not Bold",
                    severity="warning",
                    text_preview=preview
                )

            if cp.alignment != expected_align:
                self._add_issue(
                    category="journal_header",
                    location=location,
                    para_index=cp.index,
                    description="Journal header alignment does not match template",
                    current=cp.alignment,
                    expected=expected_align,
                    severity="warning",
                    text_preview=preview
                )

            has_manual_indentation = raw_text.startswith((" ", "\t")) or "\t" in raw_text.strip()
            if has_manual_indentation:
                self._add_issue(
                    category="journal_header",
                    location=location,
                    para_index=cp.index,
                    description="Journal header contains manual tab indentation that can shift layout",
                    current="Manual tabs/spaces",
                    expected="No manual indentation before centered journal header text",
                    severity="warning",
                    text_preview=preview
                )

    def _check_layout_stability(self):
        """Check manual page header spacing patterns that commonly wrap in Word."""
        for section_index, section in enumerate(self.document.sections):
            header_parts = [
                ("Page Header", section.header),
                ("First Page Header", section.first_page_header),
                ("Even Page Header", section.even_page_header),
            ]
            for label, header in header_parts:
                for paragraph in header.paragraphs:
                    text = get_paragraph_text(paragraph)
                    if not text.strip():
                        continue
                    has_unstable_tabs = "\t\t" in text or re.search(r"\t\s{2,}", text)
                    if not has_unstable_tabs:
                        continue
                    self._add_issue(
                        category="other",
                        location=f"{label} (Section {section_index + 1})",
                        para_index=-1,
                        description="Page header uses multiple manual tabs/spaces that can wrap in Word",
                        current="Multiple manual tabs/spaces",
                        expected="Single right-aligned tab stop between left and right header text",
                        severity="warning",
                        text_preview=truncate_text(text, 80)
                    )
                    break

    def _check_author_info(self):
        """Check author names, affiliations, and corresponding-author details."""
        abstract_indices = [
            cp.index for cp in self.classifications
            if cp.paragraph_type in {ParagraphType.ABSTRACT_LABEL, ParagraphType.ABSTRACT_CONTENT}
        ]
        front_matter_end = min(abstract_indices) if abstract_indices else 25
        author_paragraphs = [
            cp for cp in self.classifications
            if cp.paragraph_type == ParagraphType.AUTHOR_INFO and cp.index < front_matter_end
        ]

        for cp in author_paragraphs:
            role = classify_author_info_role(cp.text, self.rules)
            rules = self.rules.get(role, self.rules.get("author", {}))
            expected_font = rules.get("font_name")
            expected_size = rules.get("font_size")
            expected_bold = rules.get("bold")
            expected_italic = rules.get("italic")
            expected_alignment = rules.get("alignment")
            mismatches = []

            current_font = cp.font_info.get("font_name")
            if expected_font and current_font and not is_font_equivalent(current_font, expected_font):
                mismatches.append(f"font {current_font}")
            current_size = cp.font_info.get("font_size")
            if expected_size is not None and current_size and abs(current_size - expected_size) > 0.5:
                mismatches.append(f"size {current_size}pt")
            current_bold = self._is_paragraph_mostly_bold(cp.index)
            if expected_bold is not None and current_bold != bool(expected_bold):
                mismatches.append("bold" if current_bold else "not bold")
            current_italic = bool(cp.font_info.get("italic"))
            if expected_italic is not None and current_italic != bool(expected_italic):
                mismatches.append("italic" if current_italic else "not italic")
            if expected_alignment and cp.alignment != expected_alignment:
                mismatches.append(f"alignment {cp.alignment}")

            if mismatches:
                role_label = {
                    "author": "Author name",
                    "affiliation": "Affiliation",
                    "corresponding_author": "Corresponding author",
                }[role]
                self._add_issue(
                    category="author_info",
                    location=role_label,
                    para_index=cp.index,
                    description=f"{role_label} formatting does not match template",
                    current=", ".join(mismatches),
                    expected=f"{expected_font or 'Template font'}, {expected_size or 'template'}pt",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 60),
                )

            if role == "corresponding_author":
                valid_email = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", cp.text)
                valid_orcid = re.search(
                    r"\b\d{4}-\d{4}-\d{4}-\d{3}[\dXx]\b",
                    cp.text,
                )
                if rules.get("email_required") and not valid_email:
                    self._add_issue(
                        category="author_info",
                        location="Corresponding Author",
                        para_index=cp.index,
                        description="Corresponding author email is missing or invalid",
                        current="No valid email detected",
                        expected="A valid corresponding-author email address",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 60),
                    )
                if rules.get("orcid_required") and not valid_orcid:
                    self._add_issue(
                        category="author_info",
                        location="Corresponding Author",
                        para_index=cp.index,
                        description="Corresponding author ORCID is missing or invalid",
                        current="No valid ORCID detected",
                        expected="ORCID format: 0000-0000-0000-0000",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 60),
                    )

    def _check_layout(self):
        """Check page size, orientation, and column count."""
        layout_rules = self.rules.get("layout", {})
        expected_page_size = layout_rules.get("page_size")
        expected_orientation = layout_rules.get("orientation")
        expected_columns = layout_rules.get("columns")
        section = self.document.sections[0]
        width = section.page_width.inches
        height = section.page_height.inches
        short_side, long_side = sorted((width, height))

        if abs(short_side - 8.27) < 0.1 and abs(long_side - 11.69) < 0.1:
            current_page_size = "A4"
        elif abs(short_side - 8.5) < 0.1 and abs(long_side - 11.0) < 0.1:
            current_page_size = "Letter"
        else:
            current_page_size = f"{width:.2f}x{height:.2f}"

        current_orientation = "LANDSCAPE" if width > height else "PORTRAIT"
        current_columns = count_columns(self.document)

        if expected_page_size and current_page_size != expected_page_size:
            self._add_issue(
                "layout", "Page Layout", -1,
                "Page size does not match template",
                current_page_size, expected_page_size, "error",
            )
        if expected_orientation and current_orientation != expected_orientation:
            self._add_issue(
                "layout", "Page Layout", -1,
                "Page orientation does not match template",
                current_orientation, expected_orientation, "warning",
            )
        if expected_columns is not None and current_columns != int(expected_columns):
            self._add_issue(
                "layout", "Page Layout", -1,
                "Column count does not match template",
                str(current_columns), str(expected_columns), "warning",
            )
    
    def _check_body_text(self):
        """Check body text formatting - ENHANCED with tolerance"""
        body_rules = self.rules.get("body", {})
        expected_font = body_rules.get("font_name", "Times New Roman")
        expected_size = body_rules.get("font_size", 12)
        expected_bold = body_rules.get("bold")
        expected_alignment = body_rules.get("alignment")
        
        # Font size tolerance (allow +/-1.0pt difference for body text)
        size_tolerance = 1.0
        
        # Common font name variations that are equivalent
        font_equivalents = {
            "times new roman": ["times", "times-roman", "timesnewroman", "times new roman"],
            "arial": ["arial", "arial mt", "arialmt"],
            "calibri": ["calibri", "calibri light"],
        }
        
        def fonts_match(current, expected):
            if not current or not expected:
                return True  # Skip if font not detected
            current_lower = current.lower().strip()
            expected_lower = expected.lower().strip()
            if current_lower == expected_lower:
                return True
            # Check equivalents
            for base, variants in font_equivalents.items():
                if expected_lower in variants and current_lower in variants:
                    return True
            return False
        
        issues_found = 0
        max_issues_to_report = 10  # Limit to avoid too many issues
        alignment_mismatches = []
        
        for cp in self.classifications:
            if cp.paragraph_type == ParagraphType.BODY:
                font_info = cp.font_info
                
                # Check font name
                current_font = font_info.get("font_name")
                if current_font and not fonts_match(current_font, expected_font):
                    if issues_found < max_issues_to_report:
                        self._add_issue(
                            category="body_text",
                            location=f"Paragraph {cp.index + 1}",
                            para_index=cp.index,
                            description="Body text font does not match template",
                            current=current_font,
                            expected=expected_font,
                            severity="warning",
                            text_preview=truncate_text(cp.text, 50)
                        )
                    issues_found += 1
                
                # Check font size with tolerance
                current_size = font_info.get("font_size")
                if current_size and abs(current_size - expected_size) > size_tolerance:
                    if issues_found < max_issues_to_report:
                        self._add_issue(
                            category="body_text",
                            location=f"Paragraph {cp.index + 1}",
                            para_index=cp.index,
                            description="Body text font size does not match template",
                            current=f"{current_size}pt",
                            expected=f"{expected_size}pt",
                            severity="warning",
                            text_preview=truncate_text(cp.text, 50)
                        )
                    issues_found += 1

                current_bold = self._is_paragraph_mostly_bold(cp.index)
                if expected_bold is not None and current_bold != bool(expected_bold):
                    if issues_found < max_issues_to_report:
                        self._add_issue(
                            category="body_text",
                            location=f"Paragraph {cp.index + 1}",
                            para_index=cp.index,
                            description="Body text bold formatting does not match template",
                            current="Bold" if current_bold else "Not Bold",
                            expected="Bold" if expected_bold else "Not Bold",
                            severity="warning",
                            text_preview=truncate_text(cp.text, 50)
                        )
                    issues_found += 1

                if expected_alignment and cp.alignment != expected_alignment:
                    alignment_mismatches.append(cp)

        if alignment_mismatches:
            self._add_issue(
                category="body_text",
                location="Body Paragraphs",
                para_index=-1,
                description="Body text alignment does not match template",
                current=f"{len(alignment_mismatches)} non-{expected_alignment.lower()} paragraphs",
                expected=expected_alignment,
                severity="warning",
                text_preview=truncate_text(alignment_mismatches[0].text, 50)
            )
        
        # Add summary if many issues
        if issues_found > max_issues_to_report:
            self._add_issue(
                category="body_text",
                location="Multiple Paragraphs",
                para_index=-1,
                description=f"Additional {issues_found - max_issues_to_report} body text formatting issues found",
                current="Various",
                expected=f"{expected_font} {expected_size}pt",
                severity="info"
            )
        
        # Check abstract content formatting (usually smaller font like 9pt)
        self._check_abstract_content()
        
        # Check keywords content formatting (usually same as abstract)
        self._check_keywords_content()
    
    def _check_abstract_content(self):
        """Check abstract content formatting"""
        abstract_rules = self.rules.get("abstract", {})
        expected_font = abstract_rules.get("font_name", "Times New Roman")
        expected_size = abstract_rules.get("font_size", 9)
        expected_bold = abstract_rules.get("bold")
        expected_alignment = abstract_rules.get("alignment")
        minimum_words = abstract_rules.get("min_words")
        maximum_words = abstract_rules.get("max_words")
        one_paragraph = abstract_rules.get("one_paragraph")
        prohibit_equations = abstract_rules.get("prohibit_equations")
        prohibit_tables = abstract_rules.get("prohibit_tables")
        prohibit_citations = abstract_rules.get("prohibit_citations")
        size_tolerance = 0.5

        abstract_paragraphs = [
            cp for cp in self.classifications
            if cp.paragraph_type == ParagraphType.ABSTRACT_CONTENT
        ]
        for cp in abstract_paragraphs:
            font_info = cp.font_info

            current_font = font_info.get("font_name")
            if current_font and not is_font_equivalent(current_font, expected_font):
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=cp.index,
                    description="Abstract font does not match template",
                    current=current_font,
                    expected=expected_font,
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            if expected_alignment and cp.alignment != expected_alignment:
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=cp.index,
                    description="Abstract alignment does not match template",
                    current=cp.alignment,
                    expected=expected_alignment,
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            current_size = font_info.get("font_size")
            if current_size and abs(current_size - expected_size) > size_tolerance:
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=cp.index,
                    description="Abstract font size does not match template",
                    current=f"{current_size}pt",
                    expected=f"{expected_size}pt",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            current_bold = self._is_paragraph_mostly_bold(cp.index)
            if expected_bold is not None and current_bold != bool(expected_bold):
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=cp.index,
                    description="Abstract bold formatting does not match template",
                    current="Bold" if current_bold else "Not Bold",
                    expected="Bold" if expected_bold else "Not Bold",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

        if abstract_paragraphs and (minimum_words is not None or maximum_words is not None):
            abstract_text = " ".join(cp.text for cp in abstract_paragraphs)
            abstract_text = re.sub(
                r"^\s*abstract\s*(?:-|\u2013|\u2014|:)?\s*",
                "",
                abstract_text,
                flags=re.IGNORECASE,
            )
            word_count = len(re.findall(r"\b[\w]+(?:[-'][\w]+)*\b", abstract_text))
            below_minimum = minimum_words is not None and word_count < int(minimum_words)
            above_maximum = maximum_words is not None and word_count > int(maximum_words)
            if below_minimum or above_maximum:
                if minimum_words is not None and maximum_words is not None:
                    expected_range = f"{minimum_words}-{maximum_words} words"
                elif minimum_words is not None:
                    expected_range = f"At least {minimum_words} words"
                else:
                    expected_range = f"No more than {maximum_words} words"
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=abstract_paragraphs[0].index,
                    description="Abstract word count is outside the template limit",
                    current=f"{word_count} words",
                    expected=expected_range,
                    severity="warning",
                    text_preview=truncate_text(abstract_text, 50)
                )

        if one_paragraph and len(abstract_paragraphs) > 1:
            self._add_issue(
                category="body_text",
                location="Abstract",
                para_index=abstract_paragraphs[0].index,
                description="Abstract must be written as one paragraph",
                current=f"{len(abstract_paragraphs)} paragraphs",
                expected="One paragraph",
                severity="warning",
                text_preview=truncate_text(" ".join(cp.text for cp in abstract_paragraphs), 50),
            )

        if abstract_paragraphs:
            abstract_text = " ".join(cp.text for cp in abstract_paragraphs)
            first_index = abstract_paragraphs[0].index
            if prohibit_citations and (
                re.search(r"\[\d+(?:\s*[-,]\s*\d+)*\]", abstract_text)
                or re.search(r"\([A-Z][A-Za-z-]+(?:\s+et\s+al\.)?,?\s+(?:19|20)\d{2}\)", abstract_text)
            ):
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=first_index,
                    description="Abstract contains citation or reference markers",
                    current="Citation marker detected",
                    expected="No references or citations in the abstract",
                    severity="warning",
                    text_preview=truncate_text(abstract_text, 50),
                )

            has_math_xml = any(
                self.document.paragraphs[cp.index]._p.find(".//" + qn("m:oMath")) is not None
                or self.document.paragraphs[cp.index]._p.find(".//" + qn("m:oMathPara")) is not None
                for cp in abstract_paragraphs
            )
            has_equation_text = bool(re.search(r"\b[\w]+\s*[=<>]\s*[\w]+", abstract_text))
            if prohibit_equations and (has_math_xml or has_equation_text):
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=first_index,
                    description="Abstract contains equation-like content",
                    current="Equation or formula detected",
                    expected="No mathematical equations in the abstract",
                    severity="warning",
                    text_preview=truncate_text(abstract_text, 50),
                )

            if prohibit_tables and re.search(r"\b(?:table|figure|fig\.)\s*\d+", abstract_text, re.IGNORECASE):
                self._add_issue(
                    category="body_text",
                    location="Abstract",
                    para_index=first_index,
                    description="Abstract contains table or figure material",
                    current="Table or figure reference detected",
                    expected="No tabular or figure material in the abstract",
                    severity="warning",
                    text_preview=truncate_text(abstract_text, 50),
                )
    
    def _check_keywords_content(self):
        """Check keywords content formatting"""
        # Keywords typically use same format as abstract
        keywords_rules = self.rules.get("keywords", self.rules.get("abstract", {}))
        expected_font = keywords_rules.get("font_name", "Times New Roman")
        expected_size = keywords_rules.get("font_size", 9)
        expected_bold = keywords_rules.get("bold")
        expected_italic = keywords_rules.get("italic")
        minimum_count = keywords_rules.get("min_count")
        capitalize_first_letter = keywords_rules.get("capitalize_first_letter")
        size_tolerance = 0.5
        
        for cp in self.classifications:
            if cp.paragraph_type == ParagraphType.KEYWORDS_CONTENT:
                font_info = cp.font_info

                current_font = font_info.get("font_name")
                if current_font and not is_font_equivalent(current_font, expected_font):
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keywords font does not match template",
                        current=current_font,
                        expected=expected_font,
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                # Check font size
                current_size = font_info.get("font_size")
                if current_size and abs(current_size - expected_size) > size_tolerance:
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keywords font size does not match template",
                        current=f"{current_size}pt",
                        expected=f"{expected_size}pt",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                current_bold = self._is_paragraph_mostly_bold(cp.index)
                if expected_bold is not None and current_bold != bool(expected_bold):
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keywords bold formatting does not match template",
                        current="Bold" if current_bold else "Not Bold",
                        expected="Bold" if expected_bold else "Not Bold",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                current_italic = bool(font_info.get("italic"))
                if expected_italic is not None and current_italic != bool(expected_italic):
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keywords italic formatting does not match template",
                        current="Italic" if current_italic else "Not Italic",
                        expected="Italic" if expected_italic else "Not Italic",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                keyword_text = re.sub(
                    r"^\s*(?:keywords?|key\s+words?)\s*(?:-|\u2013|\u2014|:)?\s*",
                    "",
                    cp.text,
                    flags=re.IGNORECASE,
                )
                keyword_text = re.sub(r"\([^)]*\)\s*$", "", keyword_text).strip()
                keywords = [
                    item.strip()
                    for item in re.split(r"[,;]", keyword_text)
                    if item.strip()
                ]
                if minimum_count is not None and len(keywords) < int(minimum_count):
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keyword count is below the template minimum",
                        current=f"{len(keywords)} keywords",
                        expected=f"At least {minimum_count} keywords",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                invalid_capitalization = [
                    keyword for keyword in keywords
                    if keyword and keyword[0].isalpha() and not keyword[0].isupper()
                ]
                if capitalize_first_letter and invalid_capitalization:
                    self._add_issue(
                        category="body_text",
                        location="Keywords",
                        para_index=cp.index,
                        description="Keyword capitalization does not match template",
                        current=", ".join(invalid_capitalization[:5]),
                        expected="Capitalize the first letter of each keyword",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )
    
    def _check_headings(self):
        """Check section heading formatting"""
        for cp in self.classifications:
            if cp.paragraph_type == ParagraphType.SECTION_HEADING:
                heading_rules = self._heading_rules_for_text(cp.text)
                expected_font = heading_rules.get("font_name", "Times New Roman")
                expected_size = heading_rules.get("font_size", 10)
                expected_bold = heading_rules.get("bold")
                expected_italic = heading_rules.get("italic")
                font_info = cp.font_info
                
                # Check font
                current_font = font_info.get("font_name")
                if current_font and not is_font_equivalent(current_font, expected_font):
                    self._add_issue(
                        category="headings",
                        location=f"Heading: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading font does not match template",
                        current=current_font,
                        expected=expected_font,
                        severity="warning",
                        text_preview=cp.text
                    )
                
                # Check size with 1.0pt tolerance
                current_size = font_info.get("font_size")
                if current_size and abs(current_size - expected_size) > 1.0:
                    self._add_issue(
                        category="headings",
                        location=f"Heading: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading font size does not match template",
                        current=f"{current_size}pt",
                        expected=f"{expected_size}pt",
                        severity="warning",
                        text_preview=cp.text
                    )
                
                # Check bold - handle None values (None means not bold)
                current_bold = bool(font_info.get("bold"))
                if expected_bold is not None and current_bold != bool(expected_bold):
                    self._add_issue(
                        category="headings",
                        location=f"Heading: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading bold formatting does not match template",
                        current="Bold" if current_bold else "Not Bold",
                        expected="Bold" if expected_bold else "Not Bold",
                        severity="warning",
                        text_preview=cp.text
                    )

                current_italic = bool(font_info.get("italic"))
                if expected_italic is not None and current_italic != bool(expected_italic):
                    self._add_issue(
                        category="headings",
                        location=f"Heading: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading italic formatting does not match template",
                        current="Italic" if current_italic else "Not Italic",
                        expected="Italic" if expected_italic else "Not Italic",
                        severity="warning",
                        text_preview=cp.text
                    )

                expected_all_caps = heading_rules.get("all_caps")
                if expected_all_caps and not self._heading_text_is_all_caps(cp.text):
                    self._add_issue(
                        category="headings",
                        location=f"Heading: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading capitalization does not match template",
                        current=cp.text,
                        expected="All capital letters",
                        severity="warning",
                        text_preview=cp.text
                    )

                paragraph = self.document.paragraphs[cp.index]
                numbering_bold = self._get_numbering_bold(paragraph)
                if expected_bold is not None and numbering_bold is not None and numbering_bold != bool(expected_bold):
                    self._add_issue(
                        category="headings",
                        location=f"Heading Number: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Heading number bold formatting does not match template",
                        current="Bold" if numbering_bold else "Not Bold",
                        expected="Bold" if expected_bold else "Not Bold",
                        severity="warning",
                        text_preview=cp.text
                    )

    def _heading_rules_for_text(self, text: str) -> Dict[str, Any]:
        """Return the correct heading rule for main headings or subheadings."""
        stripped = re.sub(r"\s+", " ", text.strip())
        if re.match(r"^\d+\.\d+", stripped):
            return self.rules.get("subheading", self.rules.get("heading", {}))
        return self.rules.get("heading", {})

    @staticmethod
    def _heading_text_is_all_caps(text: str) -> bool:
        """Return True when heading words are uppercase after removing numbering."""
        stripped = re.sub(r"^\s*\d+(?:\.\d+)*\.?\s+", "", text or "").strip()
        letters = [char for char in stripped if char.isalpha()]
        return bool(letters) and all(char.isupper() for char in letters)
    
    def _check_document_structure(self) -> Dict[str, Any]:
        """Check required sections, order, and heading-role confidence."""
        sections = {
            section: {"found": False, "index": None, "format_status": "not_checked"}
            for section in self.required_sections
        }

        def mark_section(section_name: str, cp: ClassifiedParagraph, format_status: str):
            if section_name not in sections:
                return
            existing = sections[section_name]
            if existing["found"] and not (
                existing["format_status"] != "valid" and format_status == "valid"
            ):
                return
            sections[section_name] = {
                "found": True,
                "index": cp.index,
                "format_status": format_status,
            }

        for cp in self.classifications:
            text_lower = cp.text.lower().strip()

            if "abstract" in text_lower or cp.paragraph_type in {
                ParagraphType.ABSTRACT_LABEL,
                ParagraphType.ABSTRACT_CONTENT,
            }:
                status = "valid" if cp.paragraph_type in {
                    ParagraphType.ABSTRACT_LABEL,
                    ParagraphType.ABSTRACT_CONTENT,
                } else "weak"
                mark_section("abstract", cp, status)

            if any(keyword in text_lower for keyword in ["keywords", "key words"]) or \
               cp.paragraph_type == ParagraphType.KEYWORDS_CONTENT:
                status = "valid" if cp.paragraph_type in {
                    ParagraphType.KEYWORDS_LABEL,
                    ParagraphType.KEYWORDS_CONTENT,
                } else "weak"
                mark_section("keywords", cp, status)

            if "introduction" in text_lower and cp.paragraph_type != ParagraphType.BODY:
                status = "valid" if cp.paragraph_type == ParagraphType.SECTION_HEADING else "weak"
                mark_section("introduction", cp, status)

            if "conclusion" in text_lower and cp.paragraph_type != ParagraphType.BODY:
                status = "valid" if cp.paragraph_type == ParagraphType.SECTION_HEADING else "weak"
                mark_section("conclusion", cp, status)

            if text_lower.startswith(("references", "bibliography", "works cited")):
                status = "valid" if cp.paragraph_type == ParagraphType.SECTION_HEADING else "weak"
                mark_section("references", cp, status)
            elif cp.paragraph_type == ParagraphType.REFERENCE:
                mark_section("references", cp, "weak")

        self._check_required_heading_typos(sections)
        declarations = self._check_required_declarations()

        for section, details in sections.items():
            if not details["found"]:
                self._add_issue(
                    category="structure",
                    location="Document Structure",
                    para_index=-1,
                    description=f"Missing required section: {section.capitalize()}",
                    current="Not Found",
                    expected=f"{section.capitalize()} section",
                    severity="error"
                )
            elif details["format_status"] == "weak":
                self._add_issue(
                    category="structure",
                    location=f"{section.capitalize()} Section",
                    para_index=details["index"],
                    description="Section was found but its heading role is not confidently detected",
                    current="Weak section evidence",
                    expected="Recognized section heading or section label",
                    severity="warning"
                )

        found_positions = [
            sections[section]["index"]
            for section in self.required_sections
            if sections[section]["found"] and sections[section]["index"] is not None
        ]
        order_correct = found_positions == sorted(found_positions)
        if len(found_positions) > 1 and not order_correct:
            self._add_issue(
                category="structure",
                location="Document Section Order",
                para_index=-1,
                description="Required sections are not in the expected order",
                current="Out of order",
                expected=", ".join(section.title() for section in self.required_sections),
                severity="warning"
            )

        return {
            "sections": sections,
            "declarations": declarations,
            "order_correct": order_correct,
            "expected_order": self.required_sections,
        }

    def _check_required_declarations(self) -> Dict[str, Dict[str, Any]]:
        """Check mandatory JIWE declaration headings without generating content."""
        aliases = {
            "funding_statement": ("funding statement", "funding"),
            "author_contributions": ("author contributions", "authors contributions"),
            "conflict_of_interests": (
                "conflict of interests",
                "conflict of interest",
                "competing interests",
            ),
            "data_availability": ("data availability", "data availability statement"),
        }
        results = {}
        normalized_paragraphs = []
        for cp in self.classifications:
            text = re.sub(r"\([^)]*\)", "", cp.text)
            text = re.sub(r"^\d+(?:\.\d+)*\.?\s+", "", text)
            normalized = re.sub(r"[^a-z ]", " ", text.lower())
            normalized = re.sub(r"\s+", " ", normalized).strip()
            normalized_paragraphs.append((cp, normalized))

        for declaration in self.required_declarations:
            found = None
            for cp, normalized in normalized_paragraphs:
                if normalized in aliases.get(declaration, (declaration.replace("_", " "),)):
                    found = cp
                    break
            results[declaration] = {
                "found": found is not None,
                "index": found.index if found else None,
            }
            if found is None:
                label = declaration.replace("_", " ").title()
                self._add_issue(
                    category="structure",
                    location="Declaration Sections",
                    para_index=-1,
                    description=f"Missing required declaration section: {label}",
                    current="Not Found",
                    expected=label,
                    severity="warning",
                )
        return results

    def _check_required_heading_typos(self, sections: Dict[str, Dict[str, Any]]) -> None:
        """Report likely misspellings of required section headings."""
        reported_sections = set()
        for cp in self.classifications:
            candidate = self._normalize_heading_candidate(cp.text)
            if not candidate:
                continue
            for section in self.required_sections:
                if sections.get(section, {}).get("found") or section in reported_sections:
                    continue
                ratio = SequenceMatcher(None, candidate, section).ratio()
                if ratio >= 0.82:
                    reported_sections.add(section)
                    self._add_issue(
                        category="structure",
                        location=f"Possible {section.capitalize()} Heading",
                        para_index=cp.index,
                        description="Possible misspelled required section heading",
                        current=cp.text,
                        expected=section.upper(),
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

    def _normalize_heading_candidate(self, text: str) -> str:
        """Normalize short heading-like text for typo detection."""
        stripped = re.sub(r"\s+", " ", text.strip())
        if not stripped or len(stripped) > 80:
            return ""
        if not (stripped.isupper() or re.match(r"^\d+(\.\d+)*\.?\s+\S+", stripped)):
            return ""
        stripped = re.sub(r"^\d+(\.\d+)*\.?\s+", "", stripped)
        stripped = re.sub(r"\([^)]*\)", "", stripped)
        normalized = re.sub(r"[^a-z ]", "", stripped.lower()).strip()
        if len(normalized.split()) > 4:
            return ""
        return normalized
    
    def _check_tables(self):
        """Check tables in the document"""
        tables = self.document.tables
        
        if not tables:
            # Check if tables are mentioned in text but not found
            for cp in self.classifications:
                if re.search(r'table\s*\d+', cp.text.lower()):
                    self._add_issue(
                        category="tables",
                        location="Tables",
                        para_index=-1,
                        description="Table reference found but no tables in document",
                        current="No tables",
                        expected="Tables matching references",
                        severity="warning"
                    )
                    break
        else:
            # Check table captions
            table_captions = [
                cp for cp in self.classifications
                if cp.paragraph_type == ParagraphType.CAPTION
                and re.match(r"^table\s*\d+\s*[\.:]", cp.text.lower())
            ]
            caption_count = len(table_captions)
            expected_caption_count = self._count_captionable_tables()
            
            if caption_count < expected_caption_count:
                self._add_issue(
                    category="tables",
                    location="Table Captions",
                    para_index=-1,
                    description="Some tables may be missing captions",
                    current=f"{caption_count} captions",
                    expected=f"{expected_caption_count} captions",
                    severity="warning"
                )

            caption_rules = self.rules.get("caption", {})
            expected_font = caption_rules.get("font_name", "Times New Roman")
            expected_size = caption_rules.get("font_size", 10)
            expected_italic = caption_rules.get("italic")

            for cp in table_captions:
                current_font = cp.font_info.get("font_name")
                current_size = cp.font_info.get("font_size")
                current_italic = bool(cp.font_info.get("italic"))

                if current_font and not is_font_equivalent(current_font, expected_font):
                    self._add_issue(
                        category="tables",
                        location=f"Caption: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Table caption font does not match template",
                        current=current_font,
                        expected=expected_font,
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                if current_size and abs(current_size - expected_size) > 0.5:
                    self._add_issue(
                        category="tables",
                        location=f"Caption: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Table caption size does not match template",
                        current=f"{current_size}pt",
                        expected=f"{expected_size}pt",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

                if expected_italic is not None and current_italic != bool(expected_italic):
                    self._add_issue(
                        category="tables",
                        location=f"Caption: {truncate_text(cp.text, 30)}",
                        para_index=cp.index,
                        description="Table caption italic formatting does not match template",
                        current="Italic" if current_italic else "Not Italic",
                        expected="Italic" if expected_italic else "Not Italic",
                        severity="warning",
                        text_preview=truncate_text(cp.text, 50)
                    )

        self._check_caption_order_and_numbering("table")
    
    def _check_figures(self):
        """Check figures in the document"""
        # Count figure captions
        figure_captions = [
            cp for cp in self.classifications
            if cp.paragraph_type == ParagraphType.CAPTION and
            re.match(r'^(figure|fig\.?)\s*\d+\s*[\.:]', cp.text.lower())
        ]
        
        image_count = self._count_body_images_before_back_matter()
        
        if image_count > 0 and len(figure_captions) < image_count:
            self._add_issue(
                category="figures",
                location="Figure Captions",
                para_index=-1,
                description="Some figures may be missing captions",
                current=f"{len(figure_captions)} captions",
                expected=f"{image_count} captions (for {image_count} images)",
                severity="warning"
            )
        
        # Check caption formatting
        caption_rules = self.rules.get("caption", {})
        expected_font = caption_rules.get("font_name", "Times New Roman")
        expected_size = caption_rules.get("font_size", 10)
        expected_italic = caption_rules.get("italic")
        
        for cp in figure_captions:
            current_font = cp.font_info.get("font_name")
            current_size = cp.font_info.get("font_size")
            current_italic = bool(cp.font_info.get("italic"))
            
            if current_font and not is_font_equivalent(current_font, expected_font):
                self._add_issue(
                    category="figures",
                    location=f"Caption: {truncate_text(cp.text, 30)}",
                    para_index=cp.index,
                    description="Figure caption font does not match template",
                    current=current_font,
                    expected=expected_font,
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )
            
            if current_size and current_size != expected_size:
                self._add_issue(
                    category="figures",
                    location=f"Caption: {truncate_text(cp.text, 30)}",
                    para_index=cp.index,
                    description="Figure caption size does not match template",
                    current=f"{current_size}pt",
                    expected=f"{expected_size}pt",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            if expected_italic is not None and current_italic != bool(expected_italic):
                self._add_issue(
                    category="figures",
                    location=f"Caption: {truncate_text(cp.text, 30)}",
                    para_index=cp.index,
                    description="Figure caption italic formatting does not match template",
                    current="Italic" if current_italic else "Not Italic",
                    expected="Italic" if expected_italic else "Not Italic",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

        self._check_caption_order_and_numbering("figure")
    
    def _check_references(self):
        """Check references section formatting"""
        reference_rules = self.rules.get("reference", {})
        expected_font = reference_rules.get("font_name", "Times New Roman")
        expected_size = reference_rules.get("font_size", 9)
        expected_bold = reference_rules.get("bold")
        expected_alignment = reference_rules.get("alignment")
        expected_line_spacing = reference_rules.get("line_spacing")
        publication_italic_required = reference_rules.get("publication_italic_required")
        if publication_italic_required is None:
            publication_italic_required = self.profile.get("name", "").lower() == "jiwe"
        
        references = [
            cp for cp in self.classifications
            if cp.paragraph_type == ParagraphType.REFERENCE
        ]
        
        if not references:
            sdt_references = get_sdt_reference_paragraphs(self.document)
            if sdt_references:
                self._check_ieee_reference_consistency()
                self._check_sdt_references(
                    sdt_references,
                    expected_font,
                    expected_size,
                    expected_bold,
                    expected_alignment,
                    expected_line_spacing,
                    publication_italic_required,
                )
                return
            self._add_issue(
                category="references",
                location="References Section",
                para_index=-1,
                description="No references found in document",
                current="0 references",
                expected="At least 1 reference",
                severity="error"
            )
            return
        
        # Detect reference style
        ref_texts = [cp.text for cp in references]
        detected_style = detect_reference_style(ref_texts)
        self._check_ieee_reference_consistency()
        
        # Check for style consistency
        ieee_pattern = r'^\[\d+\]'
        apa_pattern = r'\(\d{4}\)'
        
        ieee_count = sum(1 for ref in ref_texts if re.match(ieee_pattern, ref.strip()))
        apa_count = sum(1 for ref in ref_texts if re.search(apa_pattern, ref))
        
        if ieee_count > 0 and apa_count > 0:
            self._add_issue(
                category="references",
                location="Reference Style",
                para_index=-1,
                description="Mixed reference styles detected (IEEE and APA)",
                current=f"IEEE: {ieee_count}, APA: {apa_count}",
                expected="Consistent style throughout",
                severity="warning"
            )
        
        # Check font formatting for every reference entry.
        alignment_mismatches = []
        line_spacing_mismatches = []
        for i, cp in enumerate(references):
            current_font = cp.font_info.get("font_name")
            current_size = cp.font_info.get("font_size")
            
            if current_font and not is_font_equivalent(current_font, expected_font):
                self._add_issue(
                    category="references",
                    location=f"Reference {i + 1}",
                    para_index=cp.index,
                    description="Reference font does not match template",
                    current=current_font,
                    expected=expected_font,
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            if current_size and current_size != expected_size:
                self._add_issue(
                    category="references",
                    location=f"Reference {i + 1}",
                    para_index=cp.index,
                    description="Reference font size does not match template",
                    current=f"{current_size}pt",
                    expected=f"{expected_size}pt",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            current_bold = self._is_paragraph_mostly_bold(cp.index)
            if expected_bold is not None and current_bold != bool(expected_bold):
                self._add_issue(
                    category="references",
                    location=f"Reference {i + 1}",
                    para_index=cp.index,
                    description="Reference bold formatting does not match template",
                    current="Bold" if current_bold else "Not Bold",
                    expected="Bold" if expected_bold else "Not Bold",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

            if expected_alignment and cp.alignment != expected_alignment:
                alignment_mismatches.append(cp)

            if expected_line_spacing is not None:
                current_line_spacing = get_line_spacing(self.document.paragraphs[cp.index])
                effective_line_spacing = 1.0 if current_line_spacing is None else float(current_line_spacing)
                if abs(effective_line_spacing - float(expected_line_spacing)) > 0.05:
                    line_spacing_mismatches.append(cp)

            if (
                publication_italic_required
                and self._reference_likely_has_publication_source(cp.text)
                and not self._paragraph_has_italic_text(cp.index)
            ):
                self._add_issue(
                    category="references",
                    location=f"Reference {i + 1}",
                    para_index=cp.index,
                    description="Reference publication source may need italic formatting",
                    current="No italic publication source segment detected",
                    expected="Italic journal, conference, book, or proceedings source segment",
                    severity="warning",
                    text_preview=truncate_text(cp.text, 50)
                )

        if alignment_mismatches:
            self._add_issue(
                category="references",
                location="Reference Entries",
                para_index=-1,
                description="Reference alignment does not match template",
                current=f"{len(alignment_mismatches)} non-{expected_alignment.lower()} references",
                expected=expected_alignment,
                severity="warning",
                text_preview=truncate_text(alignment_mismatches[0].text, 50)
            )

        if line_spacing_mismatches:
            self._add_issue(
                category="references",
                location="Reference Entries",
                para_index=-1,
                description="Reference line spacing does not match template",
                current=f"{len(line_spacing_mismatches)} references use another spacing",
                expected=str(expected_line_spacing),
                severity="warning",
                text_preview=truncate_text(line_spacing_mismatches[0].text, 50)
            )

    def _check_equations(self) -> None:
        """Check display-equation numbering and explicit equation citations."""
        equation_numbers = []
        unnumbered_count = 0
        citation_numbers = set()
        in_references = False
        paragraphs = [
            paragraph
            for paragraph in self.document.element.body.iter()
            if paragraph.tag == qn("w:p")
        ]

        for index, paragraph in enumerate(paragraphs):
            text = self._get_equation_paragraph_text(paragraph)
            normalized = re.sub(r"\s+", " ", text).strip()
            lower_text = normalized.lower()
            if lower_text.startswith(("references", "bibliography", "works cited")):
                in_references = True

            has_display_math = any(
                node.tag == qn("m:oMathPara")
                for node in paragraph.iter()
            )
            has_math = any(
                node.tag == qn("m:oMath")
                for node in paragraph.iter()
            )
            if has_math:
                math_text = "".join(
                    text_node.text or ""
                    for text_node in paragraph.iter()
                    if text_node.tag == qn("m:t")
                )
                number = self._extract_equation_label(math_text)
                if number is None:
                    number = self._extract_equation_label(normalized)
                if number is None and has_display_math:
                    row = paragraph.getparent()
                    while row is not None and row.tag not in {qn("w:tr"), qn("w:body")}:
                        row = row.getparent()
                    if row is not None and row.tag == qn("w:tr"):
                        row_text = "".join(
                            text_node.text or ""
                            for text_node in row.iter()
                            if text_node.tag in {qn("w:t"), qn("m:t")}
                        )
                        number = self._extract_equation_label(row_text)
                if number is None and has_display_math:
                    for adjacent_index in (index - 1, index + 1):
                        if not 0 <= adjacent_index < len(paragraphs):
                            continue
                        if any(
                            node.tag == qn("m:oMath")
                            for node in paragraphs[adjacent_index].iter()
                        ):
                            continue
                        adjacent_text = self._get_equation_paragraph_text(
                            paragraphs[adjacent_index]
                        )
                        adjacent_match = re.fullmatch(
                            r"\s*\((\d+)\)\s*",
                            adjacent_text,
                        )
                        if adjacent_match:
                            number = int(adjacent_match.group(1))
                            break
                if number is not None:
                    equation_numbers.append(number)
                elif has_display_math:
                    unnumbered_count += 1

            if not in_references:
                citation_numbers.update(self._extract_equation_citation_numbers(normalized))

        if unnumbered_count and equation_numbers and citation_numbers:
            self._add_issue(
                category="other",
                location="Equation Numbering",
                para_index=-1,
                description="Display equation may be missing a number",
                current=f"{unnumbered_count} unnumbered display equations",
                expected="Numbered display equations such as (1), (2), and (3)",
                severity="warning",
            )

        equation_numbers = list(dict.fromkeys(equation_numbers))
        expected_numbers = list(range(1, len(equation_numbers) + 1))
        if equation_numbers and equation_numbers != expected_numbers:
            self._add_issue(
                category="other",
                location="Equation Numbering",
                para_index=-1,
                description="Equation numbering is not continuous",
                current=", ".join(f"({number})" for number in equation_numbers),
                expected=", ".join(f"({number})" for number in expected_numbers),
                severity="warning",
            )

        missing_equations = sorted(citation_numbers - set(equation_numbers))
        if equation_numbers and missing_equations:
            self._add_issue(
                category="other",
                location="Equation Citations",
                para_index=-1,
                description="Equation citation has no matching equation",
                current=", ".join(f"({number})" for number in missing_equations),
                expected="Every explicit equation citation matches a numbered equation",
                severity="error",
            )

    @staticmethod
    def _get_equation_paragraph_text(paragraph) -> str:
        """Return visible Word and OMML text from a paragraph."""
        return "".join(
            text_node.text or ""
            for text_node in paragraph.iter()
            if text_node.tag in {qn("w:t"), qn("m:t")}
        )

    @staticmethod
    def _extract_equation_label(text: str) -> Optional[int]:
        """Extract a trailing equation label such as (3) or Word's #3 form."""
        match = re.search(r"#\s*(\d+)", text)
        if match:
            return int(match.group(1))
        match = re.search(r"^\s*\((\d+)\)|\((\d+)\)\s*$", text)
        if not match:
            return None
        return int(match.group(1) or match.group(2))

    @staticmethod
    def _extract_equation_citation_numbers(text: str) -> set:
        """Extract numbers from explicit Equation or Eq. citations."""
        numbers = set()
        range_pattern = (
            r"\b(?:equations?|eqs?\.?)\s*\(?(\d+)\)?"
            r"\s*[-\u2013\u2014]\s*\(?(\d+)\)?"
        )
        for match in re.finditer(range_pattern, text, re.IGNORECASE):
            start, end = int(match.group(1)), int(match.group(2))
            if start <= end:
                numbers.update(range(start, end + 1))

        single_pattern = r"\b(?:equations?|eqs?\.?)\s*\(?(\d+)\)?"
        numbers.update(
            int(match.group(1))
            for match in re.finditer(single_pattern, text, re.IGNORECASE)
        )
        return numbers

    def _check_sdt_references(
        self,
        references,
        expected_font,
        expected_size,
        expected_bold,
        expected_alignment,
        expected_line_spacing,
        publication_italic_required,
    ):
        """Check references stored inside Word content controls."""
        alignment_mismatches = []
        line_spacing_mismatches = []

        for index, paragraph in enumerate(references):
            text = get_paragraph_text(paragraph)
            font_info = get_paragraph_font_info(paragraph)
            current_font = font_info.get("font_name")
            current_size = font_info.get("font_size")
            current_bold = bool(font_info.get("bold"))

            if current_font and not is_font_equivalent(current_font, expected_font):
                self._add_issue(
                    category="references",
                    location=f"Reference Content Control {index + 1}",
                    para_index=-1,
                    description="Reference font does not match template",
                    current=current_font,
                    expected=expected_font,
                    severity="warning",
                    text_preview=truncate_text(text, 50),
                )

            if current_size and current_size != expected_size:
                self._add_issue(
                    category="references",
                    location=f"Reference Content Control {index + 1}",
                    para_index=-1,
                    description="Reference font size does not match template",
                    current=f"{current_size}pt",
                    expected=f"{expected_size}pt",
                    severity="warning",
                    text_preview=truncate_text(text, 50),
                )

            if expected_bold is not None and current_bold != bool(expected_bold):
                self._add_issue(
                    category="references",
                    location=f"Reference Content Control {index + 1}",
                    para_index=-1,
                    description="Reference bold formatting does not match template",
                    current="Bold" if current_bold else "Not Bold",
                    expected="Bold" if expected_bold else "Not Bold",
                    severity="warning",
                    text_preview=truncate_text(text, 50),
                )

            current_alignment = get_paragraph_alignment(paragraph)
            if expected_alignment and current_alignment != expected_alignment:
                alignment_mismatches.append(paragraph)

            if expected_line_spacing is not None:
                current_spacing = get_line_spacing(paragraph)
                effective_spacing = 1.0 if current_spacing is None else float(current_spacing)
                if abs(effective_spacing - float(expected_line_spacing)) > 0.05:
                    line_spacing_mismatches.append(paragraph)

            has_italic = any(
                bool(run.font.italic)
                for run in paragraph.runs
                if run.text.strip()
            )
            if (
                publication_italic_required
                and self._reference_likely_has_publication_source(text)
                and not has_italic
            ):
                self._add_issue(
                    category="references",
                    location=f"Reference Content Control {index + 1}",
                    para_index=-1,
                    description="Reference publication source may need italic formatting",
                    current="No italic publication source segment detected",
                    expected="Italic journal, conference, book, or proceedings source segment",
                    severity="warning",
                    text_preview=truncate_text(text, 50),
                )

        if alignment_mismatches:
            self._add_issue(
                category="references",
                location="Reference Content Controls",
                para_index=-1,
                description="Reference alignment does not match template",
                current=f"{len(alignment_mismatches)} non-{expected_alignment.lower()} reference paragraphs",
                expected=expected_alignment,
                severity="warning",
                text_preview=truncate_text(get_paragraph_text(alignment_mismatches[0]), 50),
            )

        if line_spacing_mismatches:
            self._add_issue(
                category="references",
                location="Reference Content Controls",
                para_index=-1,
                description="Reference line spacing does not match template",
                current=f"{len(line_spacing_mismatches)} references use another spacing",
                expected=str(expected_line_spacing),
                severity="warning",
                text_preview=truncate_text(get_paragraph_text(line_spacing_mismatches[0]), 50),
            )

    def _is_paragraph_mostly_bold(self, paragraph_index: int) -> bool:
        """Return True when most visible paragraph text is explicitly bold."""
        if paragraph_index < 0 or paragraph_index >= len(self.document.paragraphs):
            return False
        bold_chars = 0
        total_chars = 0
        for run in self.document.paragraphs[paragraph_index].runs:
            text_length = len(run.text.strip())
            if not text_length:
                continue
            total_chars += text_length
            if bool(run.font.bold):
                bold_chars += text_length
        return total_chars > 0 and (bold_chars / total_chars) >= 0.8

    def _paragraph_has_italic_text(self, paragraph_index: int) -> bool:
        """Return True when a paragraph contains any explicitly italic visible text."""
        if paragraph_index < 0 or paragraph_index >= len(self.document.paragraphs):
            return False
        for run in self.document.paragraphs[paragraph_index].runs:
            if run.text.strip() and bool(run.font.italic):
                return True
        return False

    def _reference_likely_has_publication_source(self, text: str) -> bool:
        """Return True for references likely to contain an italic publication source."""
        normalized = re.sub(r"\s+", " ", text or "").strip()
        if len(normalized) < 80:
            return False

        lower_text = normalized.lower()
        source_keywords = (
            "journal",
            "proceedings",
            "conference",
            "transactions",
            "communications",
            "informatics",
            "engineering",
            "springer",
            "elsevier",
            "ieee",
            "acm",
        )
        detail_patterns = (
            r"\bvol\.",
            r"\bvolume\b",
            r"\bno\.",
            r"\bpp\.",
            r"\bdoi\b",
            r"\(\d{4}\)",
            r"\".+?\"",
            r"\u201c.+?\u201d",
        )

        has_source_keyword = any(keyword in lower_text for keyword in source_keywords)
        has_bibliographic_detail = any(re.search(pattern, normalized, re.IGNORECASE) for pattern in detail_patterns)
        return has_source_keyword and has_bibliographic_detail

    def _count_captionable_tables(self) -> int:
        """Count tables that are likely manuscript tables needing captions."""
        count = 0
        for table in self.document.tables:
            text = "\n".join(cell.text.strip() for row in table.rows for cell in row.cells)
            normalized = re.sub(r"\s+", " ", text.lower()).strip()
            first_cell = ""
            if table.rows and table.rows[0].cells:
                first_cell = table.rows[0].cells[0].text.strip().lower()

            if not normalized:
                continue
            if normalized.startswith("received:") and "published:" in normalized:
                continue
            if first_cell.startswith("algorithm"):
                continue
            if " is a " in normalized and "research" in normalized and len(table.columns) <= 2:
                continue

            count += 1
        return count

    def _check_caption_order_and_numbering(self, caption_type: str) -> None:
        """Check JIWE caption placement and numbering in body XML order."""
        blocks = self._get_body_blocks()
        if caption_type == "table":
            caption_pattern = re.compile(r"^table\s*(\d+)\b", re.IGNORECASE)
            expected_side = "above"
            object_kind = "table"
            category = "tables"
            label = "Table"
        else:
            caption_pattern = re.compile(r"^(?:figure|fig\.?)\s*(\d+)\b", re.IGNORECASE)
            expected_side = "below"
            object_kind = "drawing"
            category = "figures"
            label = "Figure"

        captions = []
        for index, block in enumerate(blocks):
            match = caption_pattern.match(block["text"])
            if match:
                captions.append((index, block, int(match.group(1))))

        numbers = [number for _, _, number in captions]
        expected_numbers = list(range(1, len(numbers) + 1))
        if numbers and numbers != expected_numbers:
            self._add_issue(
                category=category,
                location=f"{label} Captions",
                para_index=-1,
                description=f"{label} numbering is not continuous",
                current=", ".join(str(number) for number in numbers),
                expected=", ".join(str(number) for number in expected_numbers),
                severity="warning",
            )

        for index, block, _ in captions:
            previous_block = self._nearest_content_block(blocks, index, -1)
            next_block = self._nearest_content_block(blocks, index, 1)
            reversed_position = (
                expected_side == "above"
                and previous_block
                and previous_block["kind"] == object_kind
            ) or (
                expected_side == "below"
                and next_block
                and next_block["kind"] == object_kind
            )
            if reversed_position:
                self._add_issue(
                    category=category,
                    location=f"Caption: {truncate_text(block['text'], 30)}",
                    para_index=block["paragraph_index"],
                    description=f"{label} caption should appear {expected_side} the {caption_type}",
                    current=f"Caption appears on the wrong side of the {caption_type}",
                    expected=f"Caption immediately {expected_side} the {caption_type}",
                    severity="warning",
                    text_preview=truncate_text(block["text"], 50),
                )

    def _get_body_blocks(self) -> List[Dict[str, Any]]:
        """Return body paragraphs, tables, and drawings in document order."""
        blocks = []
        paragraph_index = -1
        for child in self.document.element.body.iterchildren():
            if child.tag == qn("w:p"):
                paragraph_index += 1
                text = "".join(
                    text_node.text or ""
                    for text_node in child.iter()
                    if text_node.tag == qn("w:t")
                )
                has_drawing = any(
                    node.tag == qn("w:drawing")
                    for node in child.iter()
                )
                blocks.append({
                    "kind": "drawing" if has_drawing else "paragraph",
                    "text": re.sub(r"\s+", " ", text).strip(),
                    "paragraph_index": paragraph_index,
                })
            elif child.tag == qn("w:tbl"):
                blocks.append({
                    "kind": "table",
                    "text": "",
                    "paragraph_index": -1,
                })
        return blocks

    @staticmethod
    def _nearest_content_block(
        blocks: List[Dict[str, Any]],
        start_index: int,
        direction: int,
    ) -> Optional[Dict[str, Any]]:
        """Find the nearest non-empty paragraph, table, or drawing."""
        index = start_index + direction
        while 0 <= index < len(blocks):
            block = blocks[index]
            if block["kind"] != "paragraph" or block["text"]:
                return block
            index += direction
        return None

    def _count_body_images_before_back_matter(self) -> int:
        """Count body images while ignoring author photos in biography sections."""
        count = 0
        seen_figure_caption = False
        for child in self.document.element.body.iterchildren():
            block_text = " ".join(
                text_node.text or ""
                for text_node in child.iter()
                if text_node.tag == qn("w:t")
            )
            normalized = re.sub(r"\s+", " ", block_text.lower()).strip()
            if normalized.startswith(("biographies of authors", "appendix")):
                break
            if re.match(r"^(figure|fig\.?)\s*\d+\s*[\.:]", normalized):
                seen_figure_caption = True
            if seen_figure_caption:
                count += sum(1 for node in child.iter() if node.tag == qn("w:drawing"))
        return count

    def _extract_reference_texts_from_body_xml(self) -> List[str]:
        """Extract references from body XML, including Word content controls."""
        chunks = []
        in_references = False

        for child in self.document.element.body.iterchildren():
            block_text = "\n".join(
                text_node.text or ""
                for text_node in child.iter()
                if text_node.tag == qn("w:t")
            )
            normalized = re.sub(r"\s+", " ", block_text.lower()).strip()
            if not normalized:
                continue
            if normalized.startswith(("biographies of authors", "appendix")):
                break
            if normalized.startswith(("references", "bibliography", "works cited")):
                in_references = True
                block_text = re.sub(r"^\s*(references|bibliography|works cited)\s*", "", block_text, flags=re.IGNORECASE)
            if in_references and block_text.strip():
                chunks.append(block_text)

        reference_text = "\n".join(chunks)
        markers = list(re.finditer(r"(?m)^\s*\[\d+\]\s*$", reference_text))
        if markers:
            references = []
            for index, marker in enumerate(markers):
                start = marker.end()
                end = markers[index + 1].start() if index + 1 < len(markers) else len(reference_text)
                entry = reference_text[start:end].strip()
                if entry:
                    references.append(entry)
            return references

        return [
            line.strip()
            for line in reference_text.splitlines()
            if "doi" in line.lower() or re.search(r"\b(19|20)\d{2}\b", line)
        ]

    def _check_ieee_reference_consistency(self) -> None:
        """Check reference numbering and links between citations and entries."""
        reference_numbers = []
        citation_numbers = set()
        in_references = False

        for child in self.document.element.body.iterchildren():
            paragraph_texts = [
                "".join(
                    text_node.text or ""
                    for text_node in paragraph.iter()
                    if text_node.tag == qn("w:t")
                )
                for paragraph in child.iter()
                if paragraph.tag == qn("w:p")
            ]
            for block_text in paragraph_texts:
                normalized = re.sub(r"\s+", " ", block_text).strip()
                lower_text = normalized.lower()
                if not normalized:
                    continue

                if lower_text.startswith(("references", "bibliography", "works cited")):
                    in_references = True
                    block_text = re.sub(
                        r"^\s*(references|bibliography|works cited)\s*",
                        "",
                        block_text,
                        flags=re.IGNORECASE,
                    )
                elif in_references and lower_text.startswith(("biographies of authors", "appendix")):
                    break

                if in_references:
                    match = re.match(r"^\s*\[(\d+)\]", block_text)
                    if match:
                        reference_numbers.append(int(match.group(1)))
                else:
                    citation_numbers.update(self._extract_citation_numbers(block_text))

        if not reference_numbers:
            return

        citation_numbers.discard(0)
        highest_plausible_citation = max(reference_numbers) + 1
        citation_numbers = {
            number
            for number in citation_numbers
            if number <= highest_plausible_citation
        }

        expected_numbers = list(range(1, len(reference_numbers) + 1))
        if reference_numbers != expected_numbers:
            self._add_issue(
                category="references",
                location="Reference Numbering",
                para_index=-1,
                description="Reference numbering is not continuous",
                current=", ".join(f"[{number}]" for number in reference_numbers),
                expected=", ".join(f"[{number}]" for number in expected_numbers),
                severity="warning",
            )

        reference_number_set = set(reference_numbers)
        missing_references = sorted(citation_numbers - reference_number_set)
        if missing_references:
            self._add_issue(
                category="references",
                location="In-text Citations",
                para_index=-1,
                description="In-text citation has no matching reference",
                current=", ".join(f"[{number}]" for number in missing_references),
                expected="Every in-text citation matches a numbered reference entry",
                severity="error",
            )

        uncited_references = sorted(reference_number_set - citation_numbers)
        if uncited_references:
            self._add_issue(
                category="references",
                location="Reference Entries",
                para_index=-1,
                description="Reference is not cited in the manuscript",
                current=", ".join(f"[{number}]" for number in uncited_references),
                expected="Every reference entry is cited in the manuscript",
                severity="warning",
            )

    @staticmethod
    def _extract_citation_numbers(text: str) -> set:
        """Extract IEEE citation numbers, including comma lists and ranges."""
        numbers = set()
        for match in re.finditer(r"\[(\d+)\]\s*[-\u2013\u2014]\s*\[(\d+)\]", text):
            start, end = int(match.group(1)), int(match.group(2))
            if start <= end:
                numbers.update(range(start, end + 1))

        for match in re.finditer(
            r"\[(\d+(?:\s*(?:,|[-\u2013\u2014])\s*\d+)*)\]",
            text,
        ):
            content = match.group(1)
            for part in re.split(r"\s*,\s*", content):
                range_match = re.fullmatch(r"(\d+)\s*[-\u2013\u2014]\s*(\d+)", part)
                if range_match:
                    start, end = int(range_match.group(1)), int(range_match.group(2))
                    if start <= end:
                        numbers.update(range(start, end + 1))
                else:
                    numbers.add(int(part))
        return numbers

    def _get_numbering_level(self, paragraph):
        """Return the numbering level XML element for a numbered paragraph."""
        p_pr = paragraph._p.pPr
        if p_pr is None or p_pr.numPr is None or p_pr.numPr.numId is None:
            return None

        num_id = p_pr.numPr.numId.val
        ilvl = str(p_pr.numPr.ilvl.val if p_pr.numPr.ilvl is not None else 0)
        numbering = self.document.part.numbering_part.element

        abstract_num_id = None
        for num in numbering.findall(qn("w:num")):
            if num.get(qn("w:numId")) == str(num_id):
                abstract_node = num.find(qn("w:abstractNumId"))
                if abstract_node is not None:
                    abstract_num_id = abstract_node.get(qn("w:val"))
                break

        if abstract_num_id is None:
            return None

        for abstract_num in numbering.findall(qn("w:abstractNum")):
            if abstract_num.get(qn("w:abstractNumId")) != str(abstract_num_id):
                continue
            for level in abstract_num.findall(qn("w:lvl")):
                if level.get(qn("w:ilvl")) == ilvl:
                    return level
        return None

    def _get_numbering_bold(self, paragraph) -> Optional[bool]:
        """Read bold formatting from a paragraph numbering level."""
        level = self._get_numbering_level(paragraph)
        if level is None:
            return None
        r_pr = level.find(qn("w:rPr"))
        if r_pr is None:
            return None
        bold = r_pr.find(qn("w:b"))
        if bold is None:
            return None
        value = bold.get(qn("w:val"))
        return value not in {"0", "false", "False", "off"}
    def _check_line_spacing(self):
        """Check line spacing throughout document"""
        body_rules = self.rules.get("body", {})
        expected_spacing = body_rules.get("line_spacing", 1.5)
        
        spacing_issues = 0
        
        for cp in self.classifications:
            if cp.paragraph_type in [ParagraphType.BODY, ParagraphType.ABSTRACT_CONTENT]:
                para = self.document.paragraphs[cp.index]
                current_spacing = get_line_spacing(para)
                
                if current_spacing and abs(current_spacing - expected_spacing) > 0.1:
                    spacing_issues += 1
        
        if spacing_issues > 0:
            self._add_issue(
                category="line_spacing",
                location="Line Spacing",
                para_index=-1,
                description=f"{spacing_issues} paragraphs have incorrect line spacing",
                current="Varies",
                expected=f"{expected_spacing} line spacing",
                severity="warning"
            )
    
    def _gather_statistics(self) -> Dict[str, Any]:
        """Gather document statistics"""
        # Count paragraphs by type
        type_counts = {}
        for cp in self.classifications:
            type_name = cp.paragraph_type.value
            type_counts[type_name] = type_counts.get(type_name, 0) + 1
        
        # Count words
        total_words = sum(
            len(cp.text.split()) for cp in self.classifications
            if cp.paragraph_type not in [ParagraphType.EMPTY]
        )
        
        # Count tables and figures
        table_count = len(self.document.tables)
        figure_count = sum(
            1 for cp in self.classifications
            if cp.paragraph_type == ParagraphType.CAPTION and
            re.match(r'^(figure|fig\.?)\s*\d+', cp.text.lower())
        )
        
        # Count references
        reference_count = sum(
            1 for cp in self.classifications
            if cp.paragraph_type == ParagraphType.REFERENCE
        )
        
        return {
            "total_paragraphs": len(self.document.paragraphs),
            "classified_paragraphs": len(self.classifications),
            "paragraph_types": type_counts,
            "total_words": total_words,
            "table_count": table_count,
            "figure_count": figure_count,
            "reference_count": reference_count,
            "pages": len(self.document.sections)
        }
    
    def get_comparison_data(self) -> List[Dict[str, Any]]:
        """Get data for Turnitin-style comparison view"""
        comparison_data = []
        
        for category, issues in self.issues.items():
            for issue in issues:
                comparison_data.append({
                    "category": category,
                    "location": issue.location,
                    "paragraph_index": issue.paragraph_index,
                    "text_preview": issue.text_preview,
                    "current": issue.current_value,
                    "expected": issue.expected_value,
                    "description": issue.description,
                    "severity": issue.severity
                })
        
        return comparison_data
