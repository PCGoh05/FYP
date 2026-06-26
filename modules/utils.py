"""
Utility functions for the Academic Manuscript Format Checker
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from collections import Counter
import re
from typing import Dict, List, Tuple, Any, Optional
import os
import platform
import shutil
import subprocess
import tempfile


# Font equivalence groups - fonts that are considered equivalent
FONT_EQUIVALENTS = {
    "times": ["times new roman", "times", "times-roman", "timesnewroman"],
    "arial": ["arial", "arial unicode ms", "arial narrow"],
    "calibri": ["calibri", "calibri light"],
    "cambria": ["cambria", "cambria math"],
    "palatino": ["palatino linotype", "palatino", "book antiqua"],
}


def contains_east_asian_text(text: str) -> bool:
    """Return True when text contains CJK characters."""
    return bool(re.search(r"[\u3040-\u30ff\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]", text or ""))


def is_font_equivalent(font1: str, font2: str) -> bool:
    """
    Check if two fonts are equivalent (e.g., 'Times New Roman' and 'Times')
    
    Args:
        font1: First font name
        font2: Second font name
        
    Returns:
        True if fonts are equivalent
    """
    if font1 is None or font2 is None:
        return False
    
    f1 = font1.lower().strip()
    f2 = font2.lower().strip()
    
    # Exact match
    if f1 == f2:
        return True
    
    # Check equivalence groups
    for group in FONT_EQUIVALENTS.values():
        if f1 in group and f2 in group:
            return True
    
    return False


# PDF conversion imports
try:
    from pdf2docx import Converter as PDF2DOCXConverter
    PDF2DOCX_AVAILABLE = True
except ImportError:
    PDF2DOCX_AVAILABLE = False

try:
    from docx2pdf import convert as docx2pdf_convert
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False


def resolve_docx_to_pdf_backend(
    system_name: Optional[str] = None,
    docx2pdf_available: Optional[bool] = None,
    executable_finder=shutil.which,
) -> Optional[str]:
    """Return the supported DOCX-to-PDF backend for the current environment."""
    system = system_name or platform.system()
    has_docx2pdf = DOCX2PDF_AVAILABLE if docx2pdf_available is None else docx2pdf_available

    if executable_finder("soffice") or executable_finder("libreoffice"):
        return "libreoffice"

    if has_docx2pdf and system in {"Windows", "Darwin"}:
        return "docx2pdf"

    return None


def is_docx_to_pdf_supported() -> bool:
    """Return whether DOCX-to-PDF conversion can run on this machine."""
    return resolve_docx_to_pdf_backend() is not None


def get_docx_to_pdf_status() -> str:
    """Return a user-facing DOCX-to-PDF conversion status."""
    backend = resolve_docx_to_pdf_backend()
    if backend == "libreoffice":
        return "PDF Download Supported through LibreOffice"
    if backend == "docx2pdf":
        return "PDF Download Supported through Microsoft Word"
    if platform.system() == "Linux":
        return "PDF Download not available on this server. Install LibreOffice or download DOCX instead."
    return "PDF Download not available. Install Microsoft Word or LibreOffice."


def pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """Convert PDF bytes to DOCX bytes"""
    if not PDF2DOCX_AVAILABLE:
        raise ImportError("pdf2docx library not installed. Run: pip install pdf2docx")
    
    # Create temp files
    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as pdf_temp:
        pdf_temp.write(pdf_bytes)
        pdf_path = pdf_temp.name
    
    docx_path = pdf_path.replace('.pdf', '.docx')
    
    try:
        # Convert PDF to DOCX
        cv = PDF2DOCXConverter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        # Read the DOCX file
        with open(docx_path, 'rb') as f:
            docx_bytes = f.read()
        
        return docx_bytes
    finally:
        # Clean up temp files
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
        if os.path.exists(docx_path):
            os.remove(docx_path)


def docx_to_pdf(docx_bytes: bytes) -> bytes:
    """Convert DOCX bytes to PDF bytes"""
    backend = resolve_docx_to_pdf_backend()
    if not backend:
        raise RuntimeError(get_docx_to_pdf_status())

    with tempfile.TemporaryDirectory() as temp_dir:
        docx_path = os.path.join(temp_dir, "input.docx")
        pdf_path = os.path.join(temp_dir, "input.pdf")
        with open(docx_path, "wb") as docx_file:
            docx_file.write(docx_bytes)

        if backend == "libreoffice":
            executable = shutil.which("soffice") or shutil.which("libreoffice")
            command = [
                executable,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                temp_dir,
                docx_path,
            ]
            result = subprocess.run(
                command,
                capture_output=True,
                text=True,
                timeout=90,
                check=False,
            )
            if result.returncode != 0 or not os.path.exists(pdf_path):
                details = (result.stderr or result.stdout or "Unknown LibreOffice conversion error").strip()
                raise RuntimeError(f"PDF conversion failed through LibreOffice: {details}")
        else:
            docx2pdf_convert(docx_path, pdf_path)

        with open(pdf_path, "rb") as pdf_file:
            return pdf_file.read()


def load_document(file_path_or_bytes) -> Document:
    """Load a Word document from file path or bytes"""
    try:
        if isinstance(file_path_or_bytes, str):
            return Document(file_path_or_bytes)
        else:
            return Document(file_path_or_bytes)
    except Exception as e:
        raise ValueError(f"Failed to load document: {str(e)}")


def get_paragraph_text(paragraph) -> str:
    """Get the full text of a paragraph"""
    return paragraph.text.strip()


def get_run_font_info(run) -> Dict[str, Any]:
    """Extract font information from a run - ENHANCED VERSION"""
    font = run.font
    
    font_name = None
    font_size = None
    bold = None
    italic = None
    
    # Method 1: Direct font properties
    if font.name:
        font_name = font.name
    
    if font.size is not None:
        font_size = font.size.pt
    
    bold = font.bold
    italic = font.italic
    
    # Method 2: Check XML directly for more accurate extraction
    rPr = run._element.rPr
    if rPr is not None:
        # Font name from XML
        if font_name is None:
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is not None:
                latin_font = (
                    rFonts.get(qn('w:ascii')) or
                    rFonts.get(qn('w:hAnsi')) or
                    rFonts.get(qn('w:cs'))
                )
                east_asia_font = rFonts.get(qn('w:eastAsia'))
                font_name = latin_font
                if font_name is None and contains_east_asian_text(run.text):
                    font_name = east_asia_font
        
        # Font size from XML (w:sz is in half-points)
        if font_size is None:
            sz = rPr.find(qn('w:sz'))
            if sz is not None:
                sz_val = sz.get(qn('w:val'))
                if sz_val:
                    try:
                        font_size = int(sz_val) / 2  # Convert half-points to points
                    except:
                        pass
            
            # Also check szCs for complex script size
            if font_size is None:
                szCs = rPr.find(qn('w:szCs'))
                if szCs is not None:
                    sz_val = szCs.get(qn('w:val'))
                    if sz_val:
                        try:
                            font_size = int(sz_val) / 2
                        except:
                            pass
        
        # Bold from XML
        if bold is None:
            b = rPr.find(qn('w:b'))
            if b is not None:
                b_val = b.get(qn('w:val'))
                bold = b_val != '0' if b_val else True
        
        # Italic from XML
        if italic is None:
            i = rPr.find(qn('w:i'))
            if i is not None:
                i_val = i.get(qn('w:val'))
                italic = i_val != '0' if i_val else True
    
    return {
        "font_name": font_name,
        "font_size": font_size,
        "bold": bold,
        "italic": italic,
        "underline": font.underline,
        "strike": font.strike,
        "subscript": font.subscript,
        "superscript": font.superscript
    }


def get_paragraph_font_info(paragraph) -> Dict[str, Any]:
    """Get the most common font info from a paragraph's runs - ENHANCED VERSION"""
    font_names = []
    font_sizes = []
    bold_values = []
    italic_values = []
    
    # Method 1: Extract from runs
    if paragraph.runs:
        for run in paragraph.runs:
            if run.text.strip():
                info = get_run_font_info(run)
                if info["font_name"]:
                    font_names.append(info["font_name"])
                if info["font_size"]:
                    font_sizes.append(info["font_size"])
                if info["bold"] is not None:
                    bold_values.append(info["bold"])
                if info["italic"] is not None:
                    italic_values.append(info["italic"])
    
    # Method 2: Check paragraph style if we didn't get values from runs
    if not font_names or not font_sizes:
        style = paragraph.style
        if style:
            # Check direct style font
            if style.font:
                if not font_names and style.font.name:
                    font_names.append(style.font.name)
                if not font_sizes and style.font.size:
                    font_sizes.append(style.font.size.pt)
                if not bold_values and style.font.bold is not None:
                    bold_values.append(style.font.bold)
            
            # Check base style
            base_style = style.base_style
            if base_style and base_style.font:
                if not font_names and base_style.font.name:
                    font_names.append(base_style.font.name)
                if not font_sizes and base_style.font.size:
                    font_sizes.append(base_style.font.size.pt)
    
    # Method 3: Check paragraph XML directly
    if not font_sizes:
        pPr = paragraph._element.pPr
        if pPr is not None:
            rPr = pPr.find(qn('w:rPr'))
            if rPr is not None:
                sz = rPr.find(qn('w:sz'))
                if sz is not None:
                    sz_val = sz.get(qn('w:val'))
                    if sz_val:
                        try:
                            font_sizes.append(int(sz_val) / 2)
                        except:
                            pass
    
    # Return most common values
    return {
        "font_name": Counter(font_names).most_common(1)[0][0] if font_names else None,
        "font_size": Counter(font_sizes).most_common(1)[0][0] if font_sizes else None,
        "bold": Counter(bold_values).most_common(1)[0][0] if bold_values else None,
        "italic": Counter(italic_values).most_common(1)[0][0] if italic_values else None
    }


def get_paragraph_alignment(paragraph) -> str:
    """Get paragraph alignment as string"""
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "LEFT",
        WD_ALIGN_PARAGRAPH.CENTER: "CENTER",
        WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY"
    }

    alignment = paragraph.alignment
    if alignment is None and paragraph.style is not None:
        alignment = paragraph.style.paragraph_format.alignment

    base_style = paragraph.style.base_style if paragraph.style is not None else None
    if alignment is None and base_style is not None:
        alignment = base_style.paragraph_format.alignment

    return alignment_map.get(alignment, "LEFT")


def get_sdt_reference_paragraphs(document) -> List[Paragraph]:
    """Return reference paragraphs stored inside Word content controls."""
    references = []
    in_references = False
    body = document.element.body

    for child in body.iterchildren():
        is_content_control = child.tag == qn("w:sdt")
        if child.tag == qn("w:p"):
            paragraphs = [Paragraph(child, document._body)]
        elif is_content_control:
            paragraphs = [
                Paragraph(node, document._body)
                for node in child.iter(qn("w:p"))
            ]
        else:
            paragraphs = []

        for paragraph in paragraphs:
            text = get_paragraph_text(paragraph).strip()
            normalized = re.sub(r"\s+", " ", text.lower())
            if normalized.startswith(("references", "bibliography", "works cited")):
                in_references = True
                continue
            if in_references and normalized.startswith(
                ("biographies of authors", "author biography", "appendix")
            ):
                return references
            if in_references and is_content_control and text:
                references.append(paragraph)

    return references


def get_line_spacing(paragraph) -> Optional[float]:
    """Get effective line spacing from direct, style, or document defaults."""
    line_spacing = paragraph.paragraph_format.line_spacing
    if line_spacing is None and paragraph.style is not None:
        line_spacing = paragraph.style.paragraph_format.line_spacing

    base_style = paragraph.style.base_style if paragraph.style is not None else None
    if line_spacing is None and base_style is not None:
        line_spacing = base_style.paragraph_format.line_spacing

    if line_spacing is None:
        document = getattr(paragraph.part, "document", None)
        styles = document.styles.element if document is not None else None
        doc_defaults = styles.find(qn("w:docDefaults")) if styles is not None else None
        spacing = (
            doc_defaults.find(
                f"{qn('w:pPrDefault')}/{qn('w:pPr')}/{qn('w:spacing')}"
            )
            if doc_defaults is not None
            else None
        )
        if spacing is not None:
            line_value = spacing.get(qn("w:line"))
            line_rule = spacing.get(qn("w:lineRule"), "auto")
            if line_value and line_rule == "auto":
                line_spacing = int(line_value) / 240

    return line_spacing


def get_margins(document) -> Dict[str, float]:
    """Extract page margins from document in inches"""
    try:
        section = document.sections[0]
        return {
            "left": section.left_margin.inches if section.left_margin else 1.0,
            "right": section.right_margin.inches if section.right_margin else 1.0,
            "top": section.top_margin.inches if section.top_margin else 1.0,
            "bottom": section.bottom_margin.inches if section.bottom_margin else 1.0
        }
    except Exception:
        return {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0}


def set_margins(document, margins: Dict[str, float]):
    """Set page margins for all sections"""
    for section in document.sections:
        section.left_margin = Inches(margins.get("left", 1.0))
        section.right_margin = Inches(margins.get("right", 1.0))
        section.top_margin = Inches(margins.get("top", 1.0))
        section.bottom_margin = Inches(margins.get("bottom", 1.0))


def format_change_description(before: Dict, after: Dict) -> str:
    """Create a description of formatting changes"""
    changes = []
    
    if before.get("font_name") != after.get("font_name"):
        changes.append(f"Font: {before.get('font_name', 'Unknown')} -> {after.get('font_name')}")
    
    if before.get("font_size") != after.get("font_size"):
        changes.append(f"Size: {before.get('font_size', 'Unknown')}pt -> {after.get('font_size')}pt")
    
    if before.get("bold") != after.get("bold"):
        changes.append(f"Bold: {before.get('bold')} -> {after.get('bold')}")
    
    if before.get("alignment") != after.get("alignment"):
        changes.append(f"Alignment: {before.get('alignment', 'Unknown')} -> {after.get('alignment')}")
    
    return "; ".join(changes) if changes else "No changes"


def truncate_text(text: str, max_length: int = 50) -> str:
    """Truncate text with ellipsis"""
    if len(text) <= max_length:
        return text
    return text[:max_length-3] + "..."


def is_empty_paragraph(paragraph) -> bool:
    """Check if paragraph is empty or contains only whitespace"""
    return not paragraph.text.strip()


def count_columns(document) -> int:
    """Detect number of columns in document"""
    try:
        section = document.sections[0]
        cols = section._sectPr.xpath('./w:cols/@w:num')
        if cols:
            return int(cols[0])
        return 1
    except Exception:
        return 1


def detect_reference_style(references: List[str]) -> str:
    """Detect reference citation style (IEEE, APA, etc.)"""
    ieee_pattern = r'^\[\d+\]'
    apa_pattern = r'\(\d{4}\)'
    
    ieee_count = sum(1 for ref in references if re.match(ieee_pattern, ref.strip()))
    apa_count = sum(1 for ref in references if re.search(apa_pattern, ref))
    
    if ieee_count > apa_count:
        return "IEEE"
    elif apa_count > ieee_count:
        return "APA"
    else:
        return "Unknown"


def clean_text_for_comparison(text: str) -> str:
    """Clean text for comparison purposes"""
    # Remove extra whitespace
    text = re.sub(r'\s+', ' ', text)
    return text.strip()


def classify_author_info_role(text: str) -> str:
    """Classify a front-matter author-information paragraph by its role."""
    normalized = re.sub(r"\s+", " ", text.lower()).strip()
    if (
        "corresponding author" in normalized
        or "orcid" in normalized
        or re.search(r"[\w.+-]+@[\w.-]+\.[a-z]{2,}", normalized)
    ):
        return "corresponding_author"
    if re.search(
        r"\b(university|universiti|faculty|department|school|college|institute|"
        r"centre|center|laboratory|address|jalan|street|campus|"
        r"sdn\.?\s*bhd\.?|bhd\.?|ltd\.?|limited|inc\.?|corp\.?|corporation|"
        r"company|systems|technologies)\b",
        normalized,
    ):
        return "affiliation"
    return "author"


def calculate_compliance_score(issues: Dict[str, List]) -> float:
    """Calculate overall compliance score based on issues found"""
    weights = {
        "margins": 10,
        "layout": 10,
        "journal_header": 5,
        "title": 15,
        "author_info": 10,
        "body_text": 20,
        "headings": 10,
        "structure": 15,
        "tables": 5,
        "figures": 5,
        "references": 10,
        "line_spacing": 5,
        "other": 5
    }
    
    total_weight = sum(weights.values())
    deductions = 0
    
    for category, issue_list in issues.items():
        category_weight = weights.get(category, 5)
        if issue_list:
            # More issues = more deduction (up to full category weight)
            deduction = min(len(issue_list) * 2, category_weight)
            deductions += deduction
    
    score = max(0, 100 - (deductions / total_weight * 100))
    return round(score, 1)


def get_document_sections(document) -> Dict[str, Tuple[int, int]]:
    """Identify sections in the document and their paragraph ranges"""
    sections = {}
    current_section = "preamble"
    section_start = 0
    
    section_keywords = {
        "abstract": ["abstract"],
        "keywords": ["keywords", "key words"],
        "introduction": ["introduction", "1. introduction", "1 introduction"],
        "literature_review": ["literature review", "related work", "2. literature"],
        "methodology": ["methodology", "method", "materials and methods", "3. methodology"],
        "results": ["results", "findings", "4. results"],
        "discussion": ["discussion", "5. discussion"],
        "conclusion": ["conclusion", "conclusions", "6. conclusion"],
        "references": ["references", "bibliography", "works cited"]
    }
    
    for i, para in enumerate(document.paragraphs):
        text = para.text.strip().lower()
        
        for section_name, keywords in section_keywords.items():
            if any(text == kw or text.startswith(kw + ":") for kw in keywords):
                if current_section != "preamble":
                    sections[current_section] = (section_start, i - 1)
                current_section = section_name
                section_start = i
                break
    
    # Add the last section
    if current_section:
        sections[current_section] = (section_start, len(document.paragraphs) - 1)
    
    return sections


def rgba_to_hex(r: int, g: int, b: int, a: float = 1.0) -> str:
    """Convert RGBA to hex color"""
    return f"#{r:02x}{g:02x}{b:02x}"


def hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    """Convert hex color to RGB tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
