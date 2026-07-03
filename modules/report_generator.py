"""
Report Generator Module
Generates comparison reports showing formatting changes
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Any
from io import BytesIO
from datetime import datetime

from .auto_fixer import ChangeRecord, summarize_remaining_issues
from .display_values import format_user_value


class ReportGenerator:
    """
    Generates formatted comparison reports showing all changes made
    """

    # Colors
    RED = RGBColor(204, 0, 0)  # Error/Before
    GREEN = RGBColor(0, 128, 0)  # Correct/After
    BLUE = RGBColor(0, 0, 139)  # Info
    BLACK = RGBColor(0, 0, 0)
    LIGHT_RED = "FFCCCC"
    LIGHT_GREEN = "CCFFCC"
    LIGHT_BLUE = "CCE5FF"
    LIGHT_GRAY = "F0F0F0"

    def __init__(
        self,
        rules: Dict[str, Any],
        changes: List[ChangeRecord],
        check_result: Any = None,
        post_fix_validation: Any = None,
        post_fix_result: Any = None,
    ):
        """
        Initialize report generator

        Args:
            rules: Template formatting rules used
            changes: List of changes made during auto-fix
            check_result: CheckResult object from ManuscriptChecker
            post_fix_validation: comparison of pre-fix and post-fix checker results
            post_fix_result: CheckResult object from the corrected document
        """
        self.rules = rules
        self.changes = changes
        self.check_result = check_result
        self.post_fix_validation = post_fix_validation
        self.post_fix_result = post_fix_result
        self.document = None

    @staticmethod
    def _format_change_value(value: str) -> str:
        """Return report-friendly text for raw change values."""
        return format_user_value(value)

    def generate_comparison_report(self) -> Document:
        """Generate a detailed comparison report document"""
        self.document = Document()

        # Set up document margins
        for section in self.document.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Add title
        self._add_title()

        # Add generation timestamp
        self._add_timestamp()

        # Add usage guidance
        self._add_usage_section()

        # Add scope boundaries for automatic fixes
        self._add_auto_fix_scope_section()

        # Add summary section
        self._add_summary_section()

        # Add target format rules
        self._add_rules_section()

        # Add compliance score (if available)
        if self.check_result:
            self._add_compliance_section()

        # Add post-fix validation (if available)
        if self.post_fix_validation:
            self._add_post_fix_validation_section()

        # Add detailed changes
        self._add_changes_section()

        # Add legend
        self._add_legend()

        return self.document

    def _add_title(self):
        """Add report title"""
        title = self.document.add_heading("Format Correction Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Style the title
        for run in title.runs:
            run.font.size = Pt(24)
            run.font.bold = True
            run.font.color.rgb = self.BLUE

    def _add_timestamp(self):
        """Add generation timestamp"""
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"Generated on: {timestamp}")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

        self.document.add_paragraph()  # Spacing

    def _add_usage_section(self):
        """Add short guidance that explains the report and marked original outputs."""
        self.document.add_heading("How to Use This Report", level=1)

        guidance_items = [
            (
                "Corrected Manuscript",
                "contains supported formatting fixes that were applied automatically.",
            ),
            (
                "Highlighted Corrected Manuscript",
                "uses the corrected manuscript and marks applied-change locations in yellow for quick comparison.",
            ),
            (
                "Fix Summary Report",
                "lists what changed, the target rules, post-fix validation, and issues that still need manual review.",
            ),
        ]

        for label, description in guidance_items:
            paragraph = self.document.add_paragraph(style=None)
            paragraph.style = self.document.styles["Normal"]
            paragraph.paragraph_format.left_indent = Inches(0.15)
            label_run = paragraph.add_run(f"{label}: ")
            label_run.font.bold = True
            paragraph.add_run(description)

        note = self.document.add_paragraph()
        note.add_run(
            "Hidden page-header spacing cleanup is listed in the summary report but is not highlighted "
            "in the manuscript because the visible header text is unchanged."
        )
        note.runs[0].font.italic = True
        note.runs[0].font.color.rgb = RGBColor(96, 96, 96)

        self.document.add_paragraph()

    def _add_auto_fix_scope_section(self):
        """Add the safety boundaries for automatic correction."""
        self.document.add_heading("Auto-Fix Scope", level=1)

        scope_items = [
            (
                "Automatic fixes are limited to detected formatting properties",
                "such as fonts, font sizes, alignment, margins, line spacing, capitalization, and stable header spacing.",
            ),
            (
                "Manual review remains required",
                "for missing content, citation meaning, reference source selection, figure/table movement, equations, and any issue that could change academic meaning.",
            ),
            (
                "Post-fix validation is a safety check",
                "that re-runs the rule checker on the corrected manuscript and reports remaining or increased issues.",
            ),
        ]

        for label, description in scope_items:
            paragraph = self.document.add_paragraph(style=None)
            label_run = paragraph.add_run(f"{label}: ")
            label_run.font.bold = True
            paragraph.add_run(description)

        self.document.add_paragraph()

    def _add_summary_section(self):
        """Add summary section with change statistics"""
        self.document.add_heading("Summary", level=1)

        # Count changes by category
        changes_by_type = {}
        for change in self.changes:
            change_type = change.change_type
            changes_by_type[change_type] = changes_by_type.get(change_type, 0) + 1

        # Create summary table
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Category"
        header_cells[1].text = "Changes Made"

        self._style_header_row(table.rows[0])

        # Data rows
        total = 0
        for change_type, count in changes_by_type.items():
            row = table.add_row()
            row.cells[0].text = change_type.replace("_", " ").title()
            row.cells[1].text = str(count)
            total += count

        # Total row
        total_row = table.add_row()
        total_row.cells[0].text = "TOTAL"
        total_row.cells[1].text = str(total)

        # Bold total row
        for cell in total_row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(2)

        self.document.add_paragraph()  # Spacing

    def _add_rules_section(self):
        """Add section showing target format rules"""
        self.document.add_heading("Target Format Rules", level=1)

        # Margins
        para = self.document.add_paragraph()
        run = para.add_run("Page Margins: ")
        run.font.bold = True

        margins = self.rules.get("margins", {})
        para.add_run(
            f"Left: {margins.get('left', 1.0):.2f}in, "
            f"Right: {margins.get('right', 1.0):.2f}in, "
            f"Top: {margins.get('top', 1.0):.2f}in, "
            f"Bottom: {margins.get('bottom', 1.0):.2f}in"
        )

        # Title style
        para = self.document.add_paragraph()
        run = para.add_run("Title Style: ")
        run.font.bold = True

        title_rules = self.rules.get("title", {})
        para.add_run(
            f"{title_rules.get('font_name', 'Times New Roman')} "
            f"{title_rules.get('font_size', 24)}pt, "
            f"{'Bold, ' if title_rules.get('bold') else ''}"
            f"{title_rules.get('alignment', 'CENTER')}"
        )

        # Body style
        para = self.document.add_paragraph()
        run = para.add_run("Body Text Style: ")
        run.font.bold = True

        body_rules = self.rules.get("body", {})
        para.add_run(
            f"{body_rules.get('font_name', 'Times New Roman')} "
            f"{body_rules.get('font_size', 12)}pt, "
            f"Line spacing: {body_rules.get('line_spacing', 1.5)}"
        )

        # Heading style
        para = self.document.add_paragraph()
        run = para.add_run("Heading Style: ")
        run.font.bold = True

        heading_rules = self.rules.get("heading", {})
        para.add_run(
            f"{heading_rules.get('font_name', 'Times New Roman')} "
            f"{heading_rules.get('font_size', 14)}pt"
            f"{', Bold' if heading_rules.get('bold') else ''}"
        )

        self.document.add_paragraph()  # Spacing

    def _add_compliance_section(self):
        """Add compliance score section"""
        self.document.add_heading("Compliance Index", level=1)

        score = self.check_result.compliance_score
        total_issues = self.check_result.total_issues

        # Score display
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Determine color based on score
        if score >= 80:
            color = self.GREEN
        elif score >= 60:
            color = RGBColor(255, 165, 0)  # Orange
        else:
            color = self.RED

        run = para.add_run(f"{score}%")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = color

        # Issues count
        para = self.document.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(f"Total issues found: {total_issues}")
        note = self.document.add_paragraph()
        note.add_run(
            "This index is a user-facing rule-weighted indicator. "
            "FYP accuracy should be reported with Precision, Recall, and F1."
        )

        # Issues breakdown
        if self.check_result.issues_by_category:
            para = self.document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            breakdown = []
            for category, issues in self.check_result.issues_by_category.items():
                if issues:
                    breakdown.append(f"{category.replace('_', ' ').title()}: {len(issues)}")

            if breakdown:
                para.add_run("Issues by category: " + ", ".join(breakdown))

        self.document.add_paragraph()  # Spacing

    def _add_post_fix_validation_section(self):
        """Add corrected-document validation details."""
        validation = self.post_fix_validation
        self.document.add_heading("Post-Fix Validation", level=1)

        status_para = self.document.add_paragraph()
        status_run = status_para.add_run("Status: ")
        status_run.font.bold = True
        value_run = status_para.add_run("Safe" if validation.is_safe else "Needs Manual Review")
        value_run.font.color.rgb = self.GREEN if validation.is_safe else self.RED
        value_run.font.bold = True

        self.document.add_paragraph(f"Issues before auto-fix: {validation.before_issues}")
        self.document.add_paragraph(f"Issues after auto-fix: {validation.after_issues}")
        self.document.add_paragraph(f"Compliance before auto-fix: {validation.before_score}%")
        self.document.add_paragraph(f"Compliance after auto-fix: {validation.after_score}%")
        self.document.add_paragraph(validation.message)

        remaining_rows = summarize_remaining_issues(self.post_fix_result) if self.post_fix_result else []
        if remaining_rows:
            self.document.add_heading("Remaining Issues After Auto-Fix", level=2)
            table = self.document.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            headers = ["Category", "Count", "First Location", "First Issue"]
            for index, header in enumerate(headers):
                table.rows[0].cells[index].text = header
            self._style_header_row(table.rows[0])

            for remaining in remaining_rows:
                row = table.add_row()
                row.cells[0].text = str(remaining["Category"])
                row.cells[1].text = str(remaining["Count"])
                row.cells[2].text = str(remaining["First Location"])
                row.cells[3].text = str(remaining["First Issue"])
        else:
            self.document.add_paragraph("No remaining issues were detected after auto-fix.")

        self.document.add_paragraph()  # Spacing

    def _add_changes_section(self):
        """Add detailed changes table"""
        self.document.add_heading("Detailed Changes", level=1)

        if not self.changes:
            para = self.document.add_paragraph()
            para.add_run("No formatting changes were made.")
            return

        # Create changes table
        table = self.document.add_table(rows=1, cols=5)
        table.style = 'Table Grid'

        # Header row
        headers = ["#", "Location", "Property", "Current", "Target"]
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header

        self._style_header_row(table.rows[0])

        # Data rows
        for i, change in enumerate(self.changes, 1):
            row = table.add_row()

            # Number
            row.cells[0].text = str(i)

            # Location
            location_cell = row.cells[1]
            location_para = location_cell.paragraphs[0]
            run = location_para.add_run(change.location)
            run.font.size = Pt(9)

            if change.text_preview:
                location_para.add_run("\n")
                preview_run = location_para.add_run(f'"{change.text_preview}"')
                preview_run.font.size = Pt(8)
                preview_run.font.italic = True
                preview_run.font.color.rgb = RGBColor(128, 128, 128)

            # Property
            property_cell = row.cells[2]
            property_para = property_cell.paragraphs[0]
            run = property_para.add_run(change.property_name or change.change_type)
            run.font.size = Pt(9)

            # Current value (red background)
            before_cell = row.cells[3]
            self._set_cell_background(before_cell, self.LIGHT_RED)
            before_para = before_cell.paragraphs[0]
            run = before_para.add_run(self._format_change_value(change.current_value or change.before))
            run.font.color.rgb = self.RED
            run.font.size = Pt(9)

            # Target value (green background)
            after_cell = row.cells[4]
            self._set_cell_background(after_cell, self.LIGHT_GREEN)
            after_para = after_cell.paragraphs[0]
            run = after_para.add_run(self._format_change_value(change.target_value or change.after))
            run.font.color.rgb = self.GREEN
            run.font.size = Pt(9)

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(0.4)
            row.cells[1].width = Inches(2.5)
            row.cells[2].width = Inches(1.2)
            row.cells[3].width = Inches(1.8)
            row.cells[4].width = Inches(1.8)

        self.document.add_paragraph()  # Spacing

    def _add_legend(self):
        """Add color legend"""
        self.document.add_heading("Legend", level=2)

        # Red = Before/Error
        para = self.document.add_paragraph()
        run = para.add_run("- ")
        run.font.color.rgb = self.RED
        run.font.size = Pt(14)
        para.add_run("Red: Original (incorrect) formatting")

        # Green = After/Correct
        para = self.document.add_paragraph()
        run = para.add_run("- ")
        run.font.color.rgb = self.GREEN
        run.font.size = Pt(14)
        para.add_run("Green: Corrected formatting")

    def _style_header_row(self, row):
        """Style a table header row"""
        for cell in row.cells:
            self._set_cell_background(cell, self.LIGHT_BLUE)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)

    def _set_cell_background(self, cell, color: str):
        """Set background color of a table cell"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def get_report_bytes(self) -> bytes:
        """Get the report document as bytes for download"""
        if not self.document:
            self.generate_comparison_report()

        buffer = BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
