"""
Template Extractor Module
Extracts formatting rules from journal templates automatically
"""

from docx.oxml.ns import qn
from collections import Counter
from typing import Dict, Any, Optional
import re

from .profile_loader import ProfileLoader
from .utils import (
    load_document, get_paragraph_text, get_paragraph_font_info,
    get_paragraph_alignment, get_margins, get_line_spacing,
    count_columns, get_run_font_info
)
from config import DEFAULT_RULES, SECTION_HEADING_PATTERNS


def _parse_instruction_format(text: str) -> Dict[str, Any]:
    """
    Parse formatting instructions from template instruction text.
    
    Examples:
        "(24- Font size, bold Palatino Linotype)" -> font_size=24, bold=True, font_name="Palatino Linotype"
        "(11-Font size, bold Times New Roman)" -> font_size=11, bold=True, font_name="Times New Roman"
        "(10-Font size, Times New Roman)" -> font_size=10, bold=False, font_name="Times New Roman"
    
    Returns:
        Dictionary with extracted format rules (font_name, font_size, bold)
        Empty dict if no format info found
    """
    result = {}
    
    # Extract font size: look for patterns like "(24-", "(11-Font", "(10 Font"
    size_match = re.search(r'\((\d+)\s*[-\u2013\u2014]?\s*[Ff]ont', text)
    if size_match:
        result["font_size"] = int(size_match.group(1))
    
    # Extract bold: CAREFUL distinction between format instruction vs font variant name
    # Pattern 1: ", bold," or ", bold)" - bold is a separate format instruction -> TRUE
    # Pattern 2: "bold Palatino Linotype" - bold is part of font variant name -> FALSE
    # 
    # Key rule: If "bold" is followed by a font name (without comma), it's a font variant
    # If "bold" is followed by comma, closing paren, or end of instruction, it's a format instruction
    bold_as_format = re.search(r',\s*bold\s*[,)\]]', text, re.IGNORECASE)  # ", bold," or ", bold)"
    bold_at_end = re.search(r',\s*bold\s*$', text, re.IGNORECASE)  # ends with ", bold"
    bold_standalone = re.search(r'\(\s*bold\s*[,)]', text, re.IGNORECASE)  # "(bold," or "(bold)"
    
    # Check if "bold" is followed by a font name (making it a font variant, not format instruction)
    bold_font_variant = re.search(
        r'\bbold\s+(palatino|times|arial|calibri|cambria|georgia|verdana|helvetica|garamond)',
        text, re.IGNORECASE
    )
    
    if bold_font_variant:
        # "bold [FontName]" = font variant, NOT bold formatting
        result["bold"] = False
    elif bold_as_format or bold_at_end or bold_standalone:
        # ", bold," or "(bold)" = actual bold format instruction
        result["bold"] = True
    else:
        result["bold"] = False
    
    # Extract font name: look for common font names
    font_patterns = [
        r'(Palatino\s*Linotype)',
        r'(Times\s*New\s*Roman)',
        r'(Arial)',
        r'(Calibri)',
        r'(Cambria)',
        r'(Georgia)',
        r'(Verdana)',
        r'(Helvetica)',
        r'(Garamond)',
    ]
    for pattern in font_patterns:
        font_match = re.search(pattern, text, re.IGNORECASE)
        if font_match:
            # Normalize font name
            font_name = font_match.group(1)
            # Fix spacing issues
            font_name = re.sub(r'\s+', ' ', font_name).strip()
            result["font_name"] = font_name
            break
    
    return result


def _parse_word_limits(text: str) -> Dict[str, int]:
    """Extract explicit minimum and maximum word limits from instruction text."""
    normalized = re.sub(r"\s+", " ", text.lower())
    range_match = re.search(
        r"(?:between\s+)?(\d+)\s*(?:-|\u2013|\u2014|to)\s*(\d+)\s+words?",
        normalized,
    )
    if range_match:
        return {
            "min_words": int(range_match.group(1)),
            "max_words": int(range_match.group(2)),
        }

    limits = {}
    minimum_match = re.search(r"(?:at least|minimum(?: of)?)\s+(\d+)\s+words?", normalized)
    maximum_match = re.search(
        r"(?:not exceed|no more than|maximum(?: of)?|up to)\s+(\d+)\s+words?",
        normalized,
    )
    if minimum_match:
        limits["min_words"] = int(minimum_match.group(1))
    if maximum_match:
        limits["max_words"] = int(maximum_match.group(1))
    return limits


class TemplateExtractor:
    """Extract formatting rules from a journal template document"""
    
    def __init__(self, document=None, llm_integration=None):
        """
        Initialize the template extractor.
        
        Args:
            document: Optional pre-loaded document
            llm_integration: Optional LLM integration for intelligent analysis
        """
        self.document = document
        self.profile_loader = ProfileLoader()
        self.profile = self.profile_loader.load("generic")
        self.rules = self.profile_loader.default_rules(self.profile)
        self.debug_info = {}  # Store debug information
        self.llm = llm_integration  # LLM for fallback classification
        
        self.template_name = "Unknown Template"
        
    def load(self, file_path_or_bytes, template_name=None):
        """Load the template document"""
        self.document = load_document(file_path_or_bytes)
        if template_name:
            self.template_name = template_name
        elif isinstance(file_path_or_bytes, str):
            import os
            self.template_name = os.path.basename(file_path_or_bytes)
        return self
    
    def scan_all_fonts(self) -> Dict[str, Any]:
        """Scan ALL fonts in the document for debugging purposes"""
        if not self.document:
            raise ValueError("No document loaded. Call load() first.")
        
        all_fonts = []
        all_sizes = []
        paragraph_details = []
        
        for i, para in enumerate(self.document.paragraphs):
            text = get_paragraph_text(para)
            preview = text[:50] + "..." if len(text) > 50 else text
            
            para_fonts = []
            para_sizes = []
            
            # Method 1: Scan each run individually
            for run in para.runs:
                if not run.text.strip():
                    continue
                    
                run_info = get_run_font_info(run)
                
                if run_info.get("font_name"):
                    para_fonts.append(run_info["font_name"])
                    all_fonts.append(run_info["font_name"])
                
                if run_info.get("font_size"):
                    para_sizes.append(run_info["font_size"])
                    all_sizes.append(run_info["font_size"])
            
            # Method 2: Check paragraph style directly
            style = para.style
            style_font = None
            style_size = None
            if style:
                if style.font:
                    style_font = style.font.name
                    if style.font.size:
                        style_size = style.font.size.pt
            
            # Method 3: Check XML directly for paragraph-level font
            pPr = para._element.pPr
            xml_size = None
            if pPr is not None:
                rPr = pPr.find(qn('w:rPr'))
                if rPr is not None:
                    sz = rPr.find(qn('w:sz'))
                    if sz is not None:
                        sz_val = sz.get(qn('w:val'))
                        if sz_val:
                            try:
                                xml_size = int(sz_val) / 2
                            except:
                                pass
            
            if para_fonts or para_sizes or style_font or style_size or xml_size:
                paragraph_details.append({
                    "index": i,
                    "preview": preview,
                    "fonts_from_runs": list(set(para_fonts)) if para_fonts else None,
                    "sizes_from_runs": list(set(para_sizes)) if para_sizes else None,
                    "style_name": style.name if style else None,
                    "style_font": style_font,
                    "style_size": style_size,
                    "xml_size": xml_size
                })
        
        # Also scan document styles
        style_details = []
        for style in self.document.styles:
            # Skip styles that don't have font attribute (like NumberingStyle)
            if not hasattr(style, 'font') or style.font is None:
                continue
            try:
                style_info = {
                    "name": style.name,
                    "type": str(style.type)
                }
                if style.font.name:
                    style_info["font_name"] = style.font.name
                if style.font.size:
                    style_info["font_size"] = style.font.size.pt
                if style.font.bold is not None:
                    style_info["bold"] = style.font.bold
                if any(k in style_info for k in ["font_name", "font_size", "bold"]):
                    style_details.append(style_info)
            except AttributeError:
                # Skip styles that cause attribute errors
                continue
        
        # Summary
        font_counter = Counter(all_fonts)
        size_counter = Counter(all_sizes)
        
        self.debug_info = {
            "total_paragraphs": len(self.document.paragraphs),
            "font_frequency": dict(font_counter.most_common(10)),
            "size_frequency": dict(size_counter.most_common(10)),
            "unique_fonts": list(set(all_fonts)),
            "unique_sizes": sorted(list(set(all_sizes))),
            "paragraph_samples": paragraph_details[:30],  # First 30 paragraphs
            "document_styles": style_details[:20]  # First 20 styles
        }
        
        return self.debug_info

    def _detect_template_profile(self) -> Dict[str, Any]:
        """Detect the closest known template profile."""
        self.profile = self.profile_loader.detect_from_document(self.document, self.template_name)
        return self.profile

    def _profile_default(self, category: str, fallback: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """Return profile default rules for one category."""
        defaults = self.profile_loader.default_rules(self.profile).get(category, fallback or {})
        return defaults.copy() if isinstance(defaults, dict) else {}

    def _has_caption_evidence(self) -> bool:
        """Return True when the template contains caption-like paragraphs."""
        for paragraph in self.document.paragraphs:
            text = get_paragraph_text(paragraph).lower()
            if re.match(r"^(figure|fig\.?|table)\s*\d+", text):
                return True
            if re.search(r"(figure|table)\s+caption", text):
                return True
        return False

    def _has_reference_evidence(self) -> bool:
        """Return True when the template contains a reference section or entries."""
        for paragraph in self.document.paragraphs:
            text = get_paragraph_text(paragraph).strip().lower()
            if text in {"references", "bibliography", "works cited"}:
                return True
            if re.match(r"^\[\d+\]", text):
                return True
        return False

    def _has_journal_header_evidence(self) -> bool:
        """Return True when journal header formatting evidence exists in the template."""
        for paragraph in self.document.paragraphs[:10]:
            text = get_paragraph_text(paragraph).lower()
            font_info = get_paragraph_font_info(paragraph)
            font_name = (font_info.get("font_name") or "").lower()
            if "journal title" in text or "palatino linotype" in text or "palatino linotype" in font_name:
                return True
        return False

    def _add_rule_provenance(self, rules: Dict[str, Any], abstract_size: float) -> None:
        """Attach source, confidence, and evidence metadata to extracted rules."""
        profile = self._detect_template_profile()
        has_journal_header = self._has_journal_header_evidence()
        has_caption = self._has_caption_evidence()
        has_reference = self._has_reference_evidence()
        provenance = {}

        category_sources = {
            "margins": ("extracted", 0.95, "Read directly from DOCX section margins"),
            "journal_header": (
                "extracted" if has_journal_header else "default",
                0.85 if has_journal_header else 0.55,
                "Detected from journal title instruction" if has_journal_header else f"{profile['name']} journal header fallback",
            ),
            "title": ("extracted", 0.85, "Detected from title instruction or title-like paragraph"),
            "author": ("default", 0.55, f"{profile['name']} fallback because author style is not reliably extractable"),
            "affiliation": ("inferred", 0.65, f"Inferred from smaller repeated text size {abstract_size} pt"),
            "body": ("extracted", 0.85, "Inferred from most frequent body paragraph run formatting"),
            "heading": ("extracted", 0.75, "Detected from section heading-like paragraphs"),
            "abstract": ("extracted", 0.75, "Detected from abstract section or template default when absent"),
            "keywords": ("inferred", 0.65, "Inferred from abstract style"),
            "caption": (
                "extracted" if has_caption else "default",
                0.75 if has_caption else 0.55,
                "Detected from caption-like paragraphs" if has_caption else f"{profile['name']} caption fallback",
            ),
            "reference": (
                "extracted" if has_reference else "default",
                0.75 if has_reference else 0.55,
                "Detected from reference section" if has_reference else f"{profile['name']} reference fallback",
            ),
            "layout": ("extracted", 0.90, "Read from DOCX section layout"),
        }

        counts = {"extracted": 0, "inferred": 0, "default": 0}
        for category, values in rules.items():
            if category.startswith("_") or not isinstance(values, dict):
                continue
            source, confidence, evidence = category_sources.get(
                category,
                ("default", 0.50, "No specific extraction evidence available"),
            )
            for field, value in values.items():
                key = f"{category}.{field}"
                provenance[key] = {
                    "value": value,
                    "source": source,
                    "confidence": confidence,
                    "evidence": evidence,
                }
                counts[source] = counts.get(source, 0) + 1

        rules["_profile"] = profile
        rules["_provenance"] = provenance
        rules["_extraction_summary"] = {
            "extracted": counts.get("extracted", 0),
            "inferred": counts.get("inferred", 0),
            "default": counts.get("default", 0),
            "total": sum(counts.values()),
        }
        
    def extract_all_rules(self) -> Dict[str, Any]:
        """Extract template rules with deterministic DOCX parsing first."""
        if not self.document:
            raise ValueError("No document loaded. Call load() first.")

        self.scan_all_fonts()
        profile = self._detect_template_profile()
        default_rules = self.profile_loader.default_rules(profile)

        size_freq = self.debug_info.get("size_frequency", {})
        sizes_sorted = sorted(size_freq.items(), key=lambda x: x[1], reverse=True)

        abstract_size = 9
        if sizes_sorted:
            for size, count in sizes_sorted:
                if size < sizes_sorted[0][0] and count > 5:
                    abstract_size = size
                    break

        rules = {
            "margins": self._extract_margins(),
            "journal_header": self._extract_journal_header_style(),
            "title": self._extract_title_style(),
            "author": {
                **default_rules.get("author", {}),
            },
            "affiliation": {
                **default_rules.get("affiliation", {}),
                "font_size": abstract_size,
            },
            "body": self._extract_body_style(),
            "heading": self._extract_heading_style(),
            "abstract": self._extract_abstract_style(),
            "keywords": {
                **default_rules.get("keywords", {}),
                "font_size": abstract_size
            },
            "caption": self._extract_caption_style(),
            "reference": self._extract_reference_style(),
            "layout": self._extract_layout(),
            "_ai_enhanced": False,
            "_ai_primary": False,
            "_debug": self.debug_info
        }

        if self.llm and self.llm.is_available():
            ai_rules = self._extract_rules_with_ai()
            if ai_rules and ai_rules.get("_ai_extracted"):
                rules = self._fill_missing_rules_with_ai(rules, ai_rules)
                rules["_ai_enhanced"] = True

        rules = self.profile_loader.apply_rule_defaults(rules, profile)
        self._add_rule_provenance(rules, abstract_size)
        self.rules = rules
        return rules

    def _extract_journal_header_style(self) -> Dict[str, Any]:
        """Extract journal title/header style separately from the paper title style."""
        default_rule = self._profile_default("journal_header", DEFAULT_RULES.get("journal_header", {}))

        for i, para in enumerate(self.document.paragraphs[:10]):
            text = get_paragraph_text(para)
            text_lower = text.lower()
            if not text:
                continue

            font_info = get_paragraph_font_info(para)
            instruction = _parse_instruction_format(text)
            font_name = font_info.get("font_name") or instruction.get("font_name")
            mentions_journal_header = (
                "journal title" in text_lower
                or "palatino linotype" in text_lower
                or (font_name and font_name.lower() == "palatino linotype")
            )
            if not mentions_journal_header:
                continue

            return {
                "font_name": font_name or default_rule.get("font_name", "Palatino Linotype"),
                "font_size": font_info.get("font_size") or instruction.get("font_size") or default_rule.get("font_size", 24),
                "bold": font_info.get("bold") if font_info.get("bold") is not None else default_rule.get("bold", True),
                "alignment": get_paragraph_alignment(para) or default_rule.get("alignment", "CENTER"),
            }

        return default_rule.copy()

    def _fill_missing_rules_with_ai(self, rules: Dict[str, Any], ai_rules: Dict[str, Any]) -> Dict[str, Any]:
        """Use AI only to fill missing values from deterministic extraction."""
        merged = rules.copy()
        rule_mapping = {
            "title": "title",
            "heading": "heading",
            "body": "body",
            "abstract": "abstract",
            "reference": "reference",
            "caption": "caption",
        }

        for ai_key, local_key in rule_mapping.items():
            if local_key not in merged:
                continue

            ai_rule = self._convert_ai_rule(ai_rules.get(ai_key, {}))
            if not ai_rule:
                continue

            for field in ["font_name", "font_size", "bold", "italic"]:
                if merged[local_key].get(field) is None and ai_rule.get(field) is not None:
                    merged[local_key][field] = ai_rule[field]

        return merged
    
    def _convert_ai_rule(self, ai_rule: Dict) -> Dict:
        """Convert AI rule format to internal format"""
        if not ai_rule:
            return {}
        return {
            "font_name": ai_rule.get("font", ai_rule.get("font_name", "Times New Roman")),
            "font_size": ai_rule.get("size", ai_rule.get("font_size")),
            "bold": ai_rule.get("bold"),
            "italic": ai_rule.get("italic"),
        }
    
    def _extract_rules_with_ai(self) -> Dict[str, Any]:
        """Use AI only as an optional fallback for missing template fields."""
        if not self.llm:
            return {}
        
        # Collect paragraph info for AI analysis - include MORE context
        paragraphs_info = []
        for i, para in enumerate(self.document.paragraphs[:50]):  # First 50 paragraphs
            text = get_paragraph_text(para)
            if not text.strip():
                continue
            
            # Get font info including italic
            font_info = get_paragraph_font_info(para)
            
            # Also check for italic in runs
            italic = False
            for run in para.runs:
                if run.font.italic:
                    italic = True
                    break
            
            paragraphs_info.append({
                'text': text[:120],  # More text for context
                'font': font_info.get('font_name', 'Unknown'),
                'size': font_info.get('font_size', '?'),
                'bold': font_info.get('bold', 'Unknown'),
                'italic': italic
            })
        
        if not paragraphs_info:
            return {}
        
        # Ask AI to analyze only as optional fallback support.
        return self.llm.analyze_template_rules(paragraphs_info)
    
    def _extract_margins(self) -> Dict[str, float]:
        """Extract page margins from the template"""
        return get_margins(self.document)
    
    def _extract_title_style(self) -> Dict[str, Any]:
        """
        Extract paper title formatting style using intelligent detection.
        
        This method distinguishes between:
        - Journal Title (e.g., "(Journal Title) Journal of Informatics...")
        - Paper Title (e.g., "(Title) Preparation template...")
        
        The paper title is identified by looking for format instructions in template text.
        """
        from docx.oxml.ns import qn
        import re
        
        # NOTE: Do NOT check Word's built-in Heading 1 style first!
        # Template instructions (like "24-Font size, Times New Roman") are more accurate.
        # Heading 1 style often has default bold=True which doesn't match template intent.
        
        # STEP 1: Look for paragraphs with format instructions
        # e.g., "(24-Font size, Times New Roman)" or "(Title) 24pt"
        # Parse ALL paragraphs in first 10 to find title format pattern
        title_candidates = []

        
        for i, para in enumerate(self.document.paragraphs[:10]):
            text = get_paragraph_text(para)
            if not text or len(text) < 10:
                continue
            
            # Try to parse instruction format for SIZE and BOLD only
            instruction = _parse_instruction_format(text)
            if instruction and instruction.get("font_size"):
                font_info = get_paragraph_font_info(para)
                actual_size = font_info.get("font_size", 0) or 0
                instruction_size = instruction.get("font_size", 0)
                
                # Title should have large font (>= 16pt)
                if instruction_size >= 16 or actual_size >= 16:
                    title_candidates.append({
                        "index": i,
                        "instruction": instruction,
                        "font_info": font_info,
                        "size": max(instruction_size, actual_size),
                        "text": text[:100]
                    })
        
        # Pick the candidate with largest font size
        if title_candidates:
            title_candidates.sort(key=lambda x: -x["size"])
            best = title_candidates[0]
            instruction = best["instruction"]
            # Check if instruction text explicitly mentions 'bold' as FORMAT INSTRUCTION
            # NOT as part of font name like "bold Palatino Linotype"
            instruction_text = best.get("text", "").lower()
            # Bold as format instruction patterns:
            # - "24pt, bold" or "bold, 24pt" (bold followed by comma or end of string)
            # - "(bold)" (bold in parentheses)
            # - "bold Times New Roman" would still match, so exclude font names
            import re
            # Only match if bold is:
            # 1. Followed by comma, closing paren, or end of text
            # 2. OR preceded by comma and not followed by a font name
            has_bold_instruction = bool(re.search(r'\bbold\s*[,)\]]', instruction_text) or 
                                       re.search(r',\s*bold\b(?!\s*(times|arial|palatino|calibri|helvetica|georgia))', instruction_text))
            # Use instruction-specified bold, or None if not mentioned (preserve original)
            return {
                "font_name": "Times New Roman",  # Default to Times New Roman for academic papers
                "font_size": instruction.get("font_size", 24),
                "bold": True if has_bold_instruction else None,  # Only bold if explicitly specified as format
                "alignment": get_paragraph_alignment(self.document.paragraphs[best["index"]]) or "CENTER"
            }
        
        # STEP 2: Fall back to original logic if "(Title)" not found
        # Keywords that indicate journal name (not paper title)
        journal_name_keywords = [
            'journal of', 'proceedings of', 'transactions on',
            'international journal', 'ieee', 'acm', 'springer',
            'elsevier', 'wiley', 'taylor & francis', 'mdpi',
            'frontiers in', 'annals of', 'advances in',
            'review of', 'letters in', 'communications in',
            '(journal title)'  # Added this!
        ]
        
        # Keywords that indicate journal info (volume, issue, etc.)
        journal_info_keywords = [
            'vol.', 'volume', 'issue', 'issn', 'eissn', 'e-issn',
            'p-issn', 'doi:', 'doi ', 'http', 'www.', '\u00a9',
            'copyright', 'open access', 'received', 'accepted',
            'published', 'article info', 'article history',
            # Note: Do NOT add font names here - they appear in title instructions too!
        ]
        
        # Track paragraphs to find paper title
        journal_header_end = -1
        candidates = []
        
        for i, para in enumerate(self.document.paragraphs[:15]):
            text = get_paragraph_text(para)
            if not text or len(text) < 10:
                continue
            
            text_lower = text.lower().strip()
            font_info = get_paragraph_font_info(para)
            alignment = get_paragraph_alignment(para)
            
            # Check if this is a journal name
            is_journal_name = any(kw in text_lower for kw in journal_name_keywords)
            
            # Check if this is journal info (volume, ISSN, etc.)
            is_journal_info = any(kw in text_lower for kw in journal_info_keywords)
            
            # Check if this is a section heading
            is_heading = any(
                re.match(pattern, text_lower) 
                for pattern in SECTION_HEADING_PATTERNS
            )
            
            if is_journal_name or is_journal_info:
                journal_header_end = i
                continue
            
            if is_heading:
                continue
            
            # This could be the paper title - add as candidate
            # Score based on multiple factors
            score = 0
            
            # Position: should be after journal header
            if i > journal_header_end:
                score += 2
            
            # Length: typical title is 50-300 characters
            if 30 <= len(text) <= 300:
                score += 2
            elif len(text) > 300:
                score -= 1  # Too long, probably body text
            
            # Font size: larger fonts are more likely titles
            font_size = font_info.get("font_size")
            if font_size:
                if font_size >= 16:
                    score += 3
                elif font_size >= 12:
                    score += 2
                elif font_size >= 10:
                    score += 1
            
            # Bold text is common for titles
            if font_info.get("bold"):
                score += 2
            
            # Center alignment is common for titles
            if alignment == "CENTER":
                score += 2
            
            # Title usually doesn't end with period
            if not text.endswith('.'):
                score += 1
            
            # Title has multiple words
            word_count = len(text.split())
            if 4 <= word_count <= 25:
                score += 1
            
            # Check if next paragraph looks like author info
            if i + 1 < len(self.document.paragraphs):
                next_text = get_paragraph_text(self.document.paragraphs[i + 1]).lower()
                author_indicators = ['@', 'university', 'faculty', 'department', 
                                    'institute', 'college', 'school of']
                if any(ind in next_text for ind in author_indicators):
                    score += 3
            
            if score >= 4:
                candidates.append({
                    'index': i,
                    'score': score,
                    'font_info': font_info,
                    'alignment': alignment,
                    'text': text[:100]
                })
        
        # Select best candidate
        if candidates:
            # Sort by score (descending), then by position (ascending)
            candidates.sort(key=lambda x: (-x['score'], x['index']))
            best = candidates[0]
            
            # Use LLM to verify if score is borderline (4-6)
            if self.llm and self.llm.is_available() and best['score'] < 7:
                verified = self._llm_verify_title(best['text'])
                if not verified and len(candidates) > 1:
                    # Try second best candidate
                    best = candidates[1]
            
            # FIRST: Try to parse instruction format from the text itself
            # This handles templates like "(24-Font size, bold Times New Roman)"
            instruction_rules = _parse_instruction_format(best['text'])
            
            if instruction_rules:
                # Use instruction rules if found
                return {
                    "font_name": instruction_rules.get("font_name", "Times New Roman"),
                    "font_size": instruction_rules.get("font_size", 24),
                    "bold": instruction_rules.get("bold", False),  # Default to False, parser sets it explicitly
                    "alignment": best['alignment'] or "CENTER"
                }
            else:
                # Fall back to actual paragraph formatting
                return {
                    "font_name": best['font_info'].get("font_name", "Times New Roman"),
                    "font_size": best['font_info'].get("font_size", 14) or 14,
                    "bold": best['font_info'].get("bold", True),
                    "alignment": best['alignment'] or "CENTER"
                }
        
        # Default title style if no candidates found
        return self._profile_default("title", DEFAULT_RULES["title"])
    
    def _llm_verify_title(self, text: str) -> bool:
        """
        Use LLM to verify if text is a paper title.
        
        Args:
            text: The text to verify
            
        Returns:
            True if LLM confirms this is a paper title
        """
        if not self.llm:
            return True  # Assume true if no LLM available
        
        prompt = f"""Is this text a paper title from an academic manuscript?

Text: "{text[:200]}"

A paper title is the main title of a research paper, NOT:
- Journal name (e.g., "Journal of Computer Science")
- Volume/Issue info (e.g., "Vol. 3 No. 2")
- Author information
- Section headings (e.g., "Introduction", "Methodology")

Answer with ONLY "yes" or "no"."""

        try:
            response = self.llm.generate(prompt)
            return 'yes' in response.lower()
        except Exception:
            return True  # Assume true on error
    
    def _extract_body_style(self) -> Dict[str, Any]:
        """Extract body text style using frequency analysis - ENHANCED"""
        font_names = []
        font_sizes = []
        line_spacings = []
        alignments = []
        
        # Analyze ALL runs in body paragraphs to get accurate font info
        for para in self.document.paragraphs[5:]:
            text = get_paragraph_text(para)
            
            # Skip empty paragraphs and short headings
            if not text or len(text) < 50:
                continue
            
            # Skip section headings (usually ALL CAPS or numbered)
            text_stripped = text.strip()
            if text_stripped.isupper() and len(text_stripped) < 50:
                continue
            
            is_heading = any(
                re.match(pattern, text.lower()) 
                for pattern in SECTION_HEADING_PATTERNS
            )
            if is_heading:
                continue
            
            # Get font info from EACH run individually
            for run in para.runs:
                if not run.text.strip():
                    continue
                
                run_info = get_run_font_info(run)
                
                if run_info.get("font_name"):
                    font_names.append(run_info["font_name"])
                if run_info.get("font_size"):
                    font_sizes.append(run_info["font_size"])
            
            spacing = get_line_spacing(para)
            if spacing:
                line_spacings.append(spacing)
            if para.alignment is not None:
                alignments.append(get_paragraph_alignment(para))
        
        # Get most common values
        body_style = {
            "font_name": Counter(font_names).most_common(1)[0][0] if font_names else self._profile_default("body", DEFAULT_RULES["body"]).get("font_name", "Times New Roman"),
            "font_size": Counter(font_sizes).most_common(1)[0][0] if font_sizes else self._profile_default("body", DEFAULT_RULES["body"]).get("font_size", 10),
            "line_spacing": Counter(line_spacings).most_common(1)[0][0] if line_spacings else self._profile_default("body", DEFAULT_RULES["body"]).get("line_spacing", 1.0),
            "alignment": Counter(alignments).most_common(1)[0][0] if alignments else self._profile_default("body", DEFAULT_RULES["body"]).get("alignment"),
        }
        
        return body_style
    
    def _extract_heading_style(self) -> Dict[str, Any]:
        """Extract section heading style - ENHANCED with better bold detection"""
        heading_fonts = []
        heading_sizes = []
        heading_bold = []
        all_caps_count = 0
        
        for para in self.document.paragraphs:
            text = get_paragraph_text(para)
            text_stripped = text.strip()
            
            # Check if this is a heading (ALL CAPS, numbered, or matches pattern)
            is_all_caps = text_stripped.isupper() and 5 < len(text_stripped) < 60
            is_numbered = re.match(r'^\d+\.?\s+[A-Z]', text_stripped)
            is_pattern_match = any(
                re.match(pattern, text.lower()) 
                for pattern in SECTION_HEADING_PATTERNS
            )
            
            if (is_all_caps or is_numbered or is_pattern_match) and len(text) < 100:
                if is_all_caps:
                    all_caps_count += 1
                
                # Get font info from the FIRST run only (heading text, not annotations)
                for run in para.runs:
                    run_text = run.text.strip()
                    if not run_text:
                        continue
                    
                    # Skip annotation runs (e.g., "(10-Font size...)")
                    if run_text.startswith('(') or 'font' in run_text.lower():
                        continue
                    
                    run_info = get_run_font_info(run)
                    
                    if run_info.get("font_name"):
                        heading_fonts.append(run_info["font_name"])
                    if run_info.get("font_size"):
                        heading_sizes.append(run_info["font_size"])
                    
                    # Check bold from multiple sources:
                    # 1. Run level bold
                    run_bold = run.font.bold
                    if run_bold is not None:
                        heading_bold.append(run_bold)
                    
                    # 2. Also check instruction text for bold/italic keywords
                    # e.g., "(10-Font size, bold Times New Roman)" or "(10-Font size, italic)"
                    if 'bold' in text.lower():
                        heading_bold.append(True)
                    elif 'italic' in text.lower():
                        heading_bold.append(False)  # italic usually means NOT bold
                    
                    # Only take first meaningful run
                    break
        
        # Determine heading style from extracted data
        if heading_fonts or heading_sizes:
            # Determine bold from actual template analysis
            detected_bold = None
            
            if heading_bold:
                bold_count = sum(1 for b in heading_bold if b)
                # If majority are bold, use bold
                if bold_count > len(heading_bold) / 2:
                    detected_bold = True
                elif bold_count == 0:
                    detected_bold = False
            # If no bold info detected, leave as None (preserve original formatting)
            # DO NOT assume bold=True for ALL CAPS headings - template must specify
            
            return {
                "font_name": Counter(heading_fonts).most_common(1)[0][0] if heading_fonts else self._profile_default("heading", DEFAULT_RULES["heading"]).get("font_name", "Times New Roman"),
                "font_size": Counter(heading_sizes).most_common(1)[0][0] if heading_sizes else self._profile_default("heading", DEFAULT_RULES["heading"]).get("font_size", 10),
                "bold": detected_bold,  # Based on template analysis
                "all_caps": all_caps_count > 0
            }
        
        return self._profile_default("heading", DEFAULT_RULES["heading"])
    
    def _extract_abstract_style(self) -> Dict[str, Any]:
        """Extract abstract text style"""
        in_abstract = False
        abstract_style = self._profile_default(
            "abstract",
            DEFAULT_RULES.get("abstract", {"font_name": "Times New Roman", "font_size": 9}),
        ).copy()
        
        for para in self.document.paragraphs:
            text = get_paragraph_text(para)
            
            if text.lower().startswith("abstract"):
                in_abstract = True
                abstract_style.update(_parse_word_limits(text))
                if para.alignment is not None:
                    abstract_style["alignment"] = get_paragraph_alignment(para)
                # Check if abstract text is on same line
                if len(text) > 20 and "-" in text:
                    # Format: "Abstract - content..."
                    for run in para.runs:
                        if run.text.strip() and len(run.text) > 10:
                            run_info = get_run_font_info(run)
                            if run_info.get("font_size"):
                                abstract_style["font_name"] = run_info.get(
                                    "font_name",
                                    abstract_style.get("font_name", "Times New Roman"),
                                )
                                abstract_style["font_size"] = run_info.get(
                                    "font_size",
                                    abstract_style.get("font_size", 9),
                                )
                                return abstract_style
                continue
            
            if in_abstract and text and len(text) > 50:
                abstract_style.update(_parse_word_limits(text))
                if para.alignment is not None:
                    abstract_style["alignment"] = get_paragraph_alignment(para)
                # Get font info from runs for accuracy
                for run in para.runs:
                    if run.text.strip() and len(run.text) > 10:
                        run_info = get_run_font_info(run)
                        if run_info.get("font_size"):
                            abstract_style["font_name"] = run_info.get(
                                "font_name",
                                abstract_style.get("font_name", "Times New Roman"),
                            )
                            abstract_style["font_size"] = run_info.get(
                                "font_size",
                                abstract_style.get("font_size", 9),
                            )
                            return abstract_style
            
            # Stop if we hit another section
            if in_abstract and any(
                re.match(pattern, text.lower()) 
                for pattern in SECTION_HEADING_PATTERNS
            ):
                break
        
        return abstract_style
    
    def _extract_caption_style(self) -> Dict[str, Any]:
        """Extract figure/table caption style"""
        caption_fonts = []
        caption_sizes = []
        caption_italic = []
        instruction_fonts = []
        instruction_sizes = []
        
        caption_patterns = [
            r'^figure\s*\d+',
            r'^fig\.\s*\d+',
            r'^table\s*\d+',
        ]
        
        for para in self.document.paragraphs:
            text = get_paragraph_text(para)
            text_lower = text.lower()
            
            is_caption = any(
                re.match(pattern, text_lower)
                for pattern in caption_patterns
            )
            
            if is_caption:
                font_info = get_paragraph_font_info(para)
                
                if font_info.get("font_name"):
                    caption_fonts.append(font_info["font_name"])
                if font_info.get("font_size"):
                    caption_sizes.append(font_info["font_size"])
                if font_info.get("italic") is not None:
                    caption_italic.append(font_info["italic"])

            instruction_match = re.search(
                r"(?:use\s+)?(\d+(?:\.\d+)?)\s*(?:point|pt)\s+"
                r"(times\s+new\s+roman|palatino\s+linotype|arial|calibri|cambria)"
                r".{0,80}?(?:figure|table)\s+caption",
                text_lower,
            )
            if instruction_match:
                instruction_sizes.append(float(instruction_match.group(1)))
                font_name = re.sub(r"\s+", " ", instruction_match.group(2)).title()
                if font_name.lower() == "times new roman":
                    font_name = "Times New Roman"
                elif font_name.lower() == "palatino linotype":
                    font_name = "Palatino Linotype"
                instruction_fonts.append(font_name)
        
        if caption_fonts or caption_sizes:
            return {
                "font_name": Counter(caption_fonts).most_common(1)[0][0] if caption_fonts else self._profile_default("caption", DEFAULT_RULES["caption"]).get("font_name", "Times New Roman"),
                "font_size": Counter(caption_sizes).most_common(1)[0][0] if caption_sizes else self._profile_default("caption", DEFAULT_RULES["caption"]).get("font_size", 10),
                "italic": Counter(caption_italic).most_common(1)[0][0] if caption_italic else True
            }

        if instruction_fonts or instruction_sizes:
            return {
                "font_name": Counter(instruction_fonts).most_common(1)[0][0] if instruction_fonts else self._profile_default("caption", DEFAULT_RULES["caption"]).get("font_name", "Times New Roman"),
                "font_size": Counter(instruction_sizes).most_common(1)[0][0] if instruction_sizes else self._profile_default("caption", DEFAULT_RULES["caption"]).get("font_size", 10),
                "italic": False
            }
        
        return self._profile_default("caption", DEFAULT_RULES.get("caption", {"font_name": "Times New Roman", "font_size": 10, "italic": True}))
    
    def _extract_reference_style(self) -> Dict[str, Any]:
        """Extract reference entry style - ENHANCED"""
        in_references = False
        ref_fonts = []
        ref_sizes = []
        ref_alignments = []
        ref_line_spacings = []
        
        for para in self.document.paragraphs:
            text = get_paragraph_text(para)
            text_upper = text.upper().strip()
            
            # Check for REFERENCES heading
            if text_upper in ["REFERENCES", "BIBLIOGRAPHY", "WORKS CITED"] or \
               text.lower() in ["references", "bibliography", "works cited"]:
                in_references = True
                continue
            
            # Stop at next major section
            if in_references and text_upper.startswith("BIOGRAPHIES"):
                break
            
            if in_references and text:
                # Get font info from runs for accuracy
                for run in para.runs:
                    if run.text.strip():
                        run_info = get_run_font_info(run)
                        
                        if run_info.get("font_name"):
                            ref_fonts.append(run_info["font_name"])
                        if run_info.get("font_size"):
                            ref_sizes.append(run_info["font_size"])

                if re.match(r"^\s*(?:\[\d+\]|\d+\.)", text):
                    if para.alignment is not None:
                        ref_alignments.append(get_paragraph_alignment(para))
                    spacing = get_line_spacing(para)
                    if spacing:
                        ref_line_spacings.append(spacing)
        
        if ref_fonts or ref_sizes or ref_alignments or ref_line_spacings:
            default_rule = self._profile_default("reference", DEFAULT_RULES["reference"])
            return {
                "font_name": Counter(ref_fonts).most_common(1)[0][0] if ref_fonts else default_rule.get("font_name", "Times New Roman"),
                "font_size": Counter(ref_sizes).most_common(1)[0][0] if ref_sizes else default_rule.get("font_size", 9),
                "alignment": Counter(ref_alignments).most_common(1)[0][0] if ref_alignments else default_rule.get("alignment"),
                "line_spacing": Counter(ref_line_spacings).most_common(1)[0][0] if ref_line_spacings else default_rule.get("line_spacing"),
            }
        
        return self._profile_default("reference", DEFAULT_RULES.get("reference", {"font_name": "Times New Roman", "font_size": 9}))
    
    def _extract_layout(self) -> Dict[str, Any]:
        """Extract document layout settings"""
        columns = count_columns(self.document)
        
        # Detect page size
        try:
            section = self.document.sections[0]
            width = section.page_width.inches if section.page_width else 8.5
            height = section.page_height.inches if section.page_height else 11
            
            # Determine page size
            if abs(width - 8.27) < 0.1 and abs(height - 11.69) < 0.1:
                page_size = "A4"
            elif abs(width - 8.5) < 0.1 and abs(height - 11) < 0.1:
                page_size = "Letter"
            else:
                page_size = f"{width:.2f}x{height:.2f}"
        except Exception:
            page_size = self._profile_default("layout", DEFAULT_RULES["layout"]).get("page_size", "A4")
        
        return {
            "columns": columns,
            "page_size": page_size
        }
    
    def get_rules(self) -> Dict[str, Any]:
        """Get the extracted rules"""
        return self.rules
    
    def get_rules_summary(self) -> str:
        """Generate an evidence-based summary of extracted formatting rules."""
        rules = self.rules
        extraction_summary = rules.get("_extraction_summary", {})
        profile = rules.get("_profile", {})
        provenance = rules.get("_provenance", {})

        lines = [f"=== Formatting Rules Summary ({self.template_name}) ==="]
        if profile:
            lines.append(f"Template Profile: {profile.get('name', 'Generic')}")
        if extraction_summary:
            lines.append(
                "Rule Sources: "
                f"{extraction_summary.get('extracted', 0)} extracted, "
                f"{extraction_summary.get('inferred', 0)} inferred, "
                f"{extraction_summary.get('default', 0)} defaulted"
            )
        lines.append("")

        display_order = [
            ("margins", "Page Margins"),
            ("journal_header", "Journal Header Style"),
            ("title", "Paper Title Style"),
            ("author", "Author Style"),
            ("affiliation", "Affiliation Style"),
            ("abstract", "Abstract Style"),
            ("keywords", "Keywords Style"),
            ("body", "Body Text Style"),
            ("heading", "Section Heading Style"),
            ("caption", "Caption Style"),
            ("reference", "Reference Style"),
            ("layout", "Layout"),
        ]

        for category, label in display_order:
            values = rules.get(category, {})
            if not isinstance(values, dict):
                continue
            lines.append(f"{label}:")
            for field, value in values.items():
                note = provenance.get(f"{category}.{field}", {})
                if note:
                    lines.append(
                        f"  {field}: {value} "
                        f"[{note.get('source', 'unknown')}, confidence {note.get('confidence', 0):.2f}]"
                    )
                    lines.append(f"    Evidence: {note.get('evidence', 'No evidence recorded')}")
                else:
                    lines.append(f"  {field}: {value}")
            lines.append("")

        return "\n".join(lines).strip()
